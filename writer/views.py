from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
import requests

# AI Quote Extraction & Summarization Endpoint
@csrf_exempt
@login_required
def ai_quote_extraction(request):
    if request.method == 'POST':
        import json
        data = json.loads(request.body)
        chapter_text = data.get('chapter_text', '')
        if not chapter_text:
            return JsonResponse({'success': False, 'error': 'No chapter text provided.'}, status=400)

        # Call OpenAI API (or other AI provider)
        OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
        if not OPENAI_API_KEY:
            return JsonResponse({'success': False, 'error': 'OpenAI API key not configured.'}, status=500)

        prompt = (
            "Extract the most important quotes from the following chapter. "
            "For each quote, provide a brief summary explaining its relevance to the central story. "
            "Return a JSON array of objects with 'quote' and 'summary'.\n\nChapter:\n" + chapter_text
        )

        headers = {
            'Authorization': f'Bearer {OPENAI_API_KEY}',
            'Content-Type': 'application/json',
        }
        payload = {
            'model': 'gpt-3.5-turbo',
            'messages': [
                {'role': 'system', 'content': 'You are a helpful literary assistant.'},
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 800,
            'temperature': 0.7,
        }
        try:
            response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            ai_text = result['choices'][0]['message']['content']
            # Try to parse JSON from AI response
            import json as pyjson
            try:
                quotes = pyjson.loads(ai_text)
            except Exception:
                quotes = ai_text  # Fallback: return raw text
            return JsonResponse({'success': True, 'quotes': quotes})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    return JsonResponse({'error': 'POST required'}, status=405)
# Writing Stats API endpoint
from django.http import JsonResponse
from django.http import JsonResponse

def api_endpoint(request):
    data = {"message": "Hello from Django!", "date": "2025-08-21"}
    return JsonResponse(data)

def stats_api(request):
    # Replace these with real stats from your models
    data = {
        "wordCount": 12450,
        "streak": 7,
        "progress": 62
    }
    return JsonResponse(data)
from django.urls import path
from . import views

urlpatterns = [
    # ...existing code...
    path('api/stats/', views.stats_api, name='stats_api'),
]

urlpatterns = [
    # ...existing code...
    path('api/stats/', views.stats_api, name='stats_api'),
]
# Shortcut: /writer/editor/ → latest project editor or project create
from django.contrib.auth.decorators import login_required
@login_required
def latest_editor_shortcut(request):
    from .models import Project
    latest_project = Project.objects.filter(author=request.user).order_by('-updated_at').first()
    if latest_project:
        return redirect('writer:chapter_editor', project_id=latest_project.pk)
    else:
        return redirect('writer:project_create')
# TinyMCE Editor View (for both documents and chapters)
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
@login_required
def tinymce_editor(request):
    # For demo: store in session, or extend to DB as needed
    if request.method == 'POST':
        action = request.POST.get('action')
        title = request.POST.get('title', '')
        content = request.POST.get('content', '')
        if action == 'save':
            request.session['tinymce_editor_title'] = title
            request.session['tinymce_editor_content'] = content
            return JsonResponse({'status': 'success'})
        elif action == 'save_to_library':
            # Save as a Document for now
            doc = Document.objects.create(
                title=title or 'Untitled',
                content=content,
                author=request.user
            )
            return JsonResponse({'status': 'success', 'message': 'Saved to library!'})
        return JsonResponse({'status': 'error', 'message': 'Invalid action'})
    title = request.session.get('tinymce_editor_title', '')
    content = request.session.get('tinymce_editor_content', '')
    return render(request, 'writer/tinymce_editor.html', {'title': title, 'content': content})

from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy, reverse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.models import User
from django.db.models import Q, Max, F
from django.db import models, transaction
import re
from .models import (Document, Project, Chapter, AIAssistanceRequest, Character, 
                    ImportedDocument, ProjectCollaborator, WritingTheme, 
                    PersonalLibrary, WritingSession, CreativeNotebook, CreativeNode, NodeConnection, UserProfile)
from .forms import (DocumentForm, ProjectForm, ChapterForm, AIAssistanceForm, 
                   CharacterForm, ImportDocumentForm, CollaboratorForm)
import json
import random
import os
import mimetypes
import logging
# PDF generation imports - optional
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
from io import BytesIO
from html.parser import HTMLParser
from django.conf import settings
from .document_parser import extract_text_from_file  # , is_google_docs_url
import tempfile

# Simple Editor View
@login_required
def simple_editor(request):
    # For demo: store in session, or extend to DB as needed
    if request.method == 'POST':
        action = request.POST.get('action')
        title = request.POST.get('title', '')
        content = request.POST.get('content', '')
        if action == 'save':
            request.session['simple_editor_title'] = title
            request.session['simple_editor_content'] = content
            return JsonResponse({'status': 'success'})
        elif action == 'save_to_library':
            # Save as a Document for now
            doc = Document.objects.create(
                title=title or 'Untitled',
                content=content,
                author=request.user
            )
            return JsonResponse({'status': 'success', 'message': 'Saved to library!'})
        return JsonResponse({'status': 'error', 'message': 'Invalid action'})
    title = request.session.get('simple_editor_title', '')
    content = request.session.get('simple_editor_content', '')
    return render(request, 'writer/simple_editor.html', {'title': title, 'content': content})

# Editor shortcut view
def editor_shortcut(request):
    from .models import Project
    project = Project.objects.order_by('id').first()
    if project:
        return HttpResponseRedirect(reverse('writer:chapter_editor', args=[project.pk]))
    else:
        return HttpResponse('<div style="text-align:center;margin-top:3em;font-size:1.5em;">No projects found. Please create a project first.</div>')


# Personal Library Views
@login_required
def personal_library(request):
    try:
        library, created = PersonalLibrary.objects.get_or_create(user=request.user)
        user_projects = Project.objects.filter(author=request.user).select_related('author').prefetch_related('chapters').order_by('-updated_at')
        imported_docs = ImportedDocument.objects.filter(user=request.user).order_by('-import_date')

        latest_project = user_projects.first() if user_projects else None

        context = {
            'library': library,
            'projects': user_projects,
            'imported_documents': imported_docs,
            'latest_project': latest_project,
        }
        return render(request, 'writer/personal_library.html', context)
    except Exception as e:
        import logging
        logging.error(f"Error loading personal library for user {request.user}: {e}")
        print(f"Error loading personal library for user {request.user}: {e}")
        context = {
            'library': None,
            'projects': [],
            'imported_documents': [],
            'latest_project': None,
            'error_message': 'There was an issue loading your library. Please try again.'
        }

@login_required
def vault_of_lagrimas(request):
    """
    Vault of Lágrimas - An empathetic AI writing sanctuary where users can pour out their hearts,
    share their deepest emotions, and receive understanding from an AI companion learning about humanity.
    """
    return render(request, 'writer/vault_of_lagrimas.html')


# Project Views
class ProjectListView(LoginRequiredMixin, ListView):
    model = Project
    template_name = 'writer/project_list.html'
    context_object_name = 'projects'
    
    def get_queryset(self):
        return Project.objects.filter(
            Q(author=self.request.user) | Q(collaborators=self.request.user)
        ).distinct()


class ProjectDetailView(LoginRequiredMixin, DetailView):
    model = Project
    template_name = 'writer/project_detail.html'
    context_object_name = 'project'
    
    def get_queryset(self):
        return Project.objects.filter(
            Q(author=self.request.user) | Q(collaborators=self.request.user)
        ).distinct()
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['chapters'] = self.object.chapters.all().order_by('order')
        context['characters'] = self.object.characters.all()
        context['collaborators'] = self.object.projectcollaborator_set.all()
        context['writing_sessions'] = self.object.writing_sessions.filter(
            is_active=True
        )[:5]
        return context


class ProjectCreateView(LoginRequiredMixin, CreateView):
    model = Project
    form_class = ProjectForm
    template_name = 'writer/project_form.html'
    
    def get_success_url(self):
        # Redirect back to ultimate dashboard after project creation
        return reverse('writer:ultimate_dashboard')
    
    def form_valid(self, form):
        form.instance.author = self.request.user
        response = super().form_valid(form)
        
        # Add success message
        messages.success(self.request, f'Project "{form.instance.title}" created successfully! Welcome to your ultimate dashboard.')
        
        # Add collaborators if specified
        collaborator_emails = form.cleaned_data.get('collaborator_emails')
        if collaborator_emails:
            self.add_collaborators(form.instance, collaborator_emails)
        
        # Create personal library if it doesn't exist
        PersonalLibrary.objects.get_or_create(user=self.request.user)
        
        return response
    
    def add_collaborators(self, project, email_string):
        emails = [email.strip() for email in email_string.split(',') if email.strip()]
        for email in emails:
            try:
                user = User.objects.get(email=email)
                ProjectCollaborator.objects.get_or_create(
                    project=project,
                    user=user,
                    defaults={'role': 'contributor'}
                )
            except User.DoesNotExist:
                messages.warning(self.request, f"User with email {email} not found.")


class ProjectUpdateView(LoginRequiredMixin, UpdateView):
    model = Project
    form_class = ProjectForm
    template_name = 'writer/project_form.html'
    
    def get_queryset(self):
        return Project.objects.filter(author=self.request.user)
    
    def get_success_url(self):
        return reverse('writer:project_detail', kwargs={'pk': self.object.pk})
    
    def form_valid(self, form):
        messages.success(self.request, f'Project "{form.instance.title}" updated successfully!')
        return super().form_valid(form)


class ProjectDeleteView(LoginRequiredMixin, DeleteView):
    model = Project
    template_name = 'writer/project_confirm_delete.html'
    success_url = reverse_lazy('writer:project_list')
    
    def get_queryset(self):
        return Project.objects.filter(author=self.request.user)


@login_required
def toggle_project_collaboration(request, project_id):
    """Toggle collaboration status for a project"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Only POST method allowed'}, status=405)
    
    try:
        project = get_object_or_404(Project, id=project_id, author=request.user)
        project.is_collaborative = not project.is_collaborative
        project.save()
        
        return JsonResponse({
            'success': True,
            'is_collaborative': project.is_collaborative,
            'message': f'Project is now {"collaborative" if project.is_collaborative else "private"}'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def get_project_share_link(request, project_id):
    """Generate share link for collaborative project"""
    try:
        project = get_object_or_404(Project, id=project_id)
        
        # Check if user has permission (author or collaborator with sharing rights)
        if not (project.author == request.user or 
                project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({
                'success': False,
                'error': 'You do not have permission to share this project'
            }, status=403)
        
        if not project.is_collaborative:
            return JsonResponse({
                'success': False,
                'error': 'This project is not collaborative'
            }, status=400)
        
        # Generate share URL
        share_url = request.build_absolute_uri(
            reverse('writer:project_collaborate', kwargs={'project_id': project.id})
        )
        
        return JsonResponse({
            'success': True,
            'share_url': share_url,
            'project_title': project.title
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def join_collaborative_project(request, project_id):
    """Allow users to join a collaborative project via shared link"""
    try:
        project = get_object_or_404(Project, id=project_id, is_collaborative=True)
        
        if request.method == 'GET':
            # Show collaboration join page
            context = {
                'project': project,
            }
            return render(request, 'writer/project_collaborate.html', context)
        
        elif request.method == 'POST':
            # Check if user is already a collaborator
            if project.collaborators.filter(id=request.user.id).exists() or project.author == request.user:
                messages.info(request, 'You are already part of this project.')
                return redirect('writer:chapter_editor', project_id=project.id)
            
            # Add user as collaborator
            ProjectCollaborator.objects.create(
                project=project,
                user=request.user,
                role='contributor',
                can_edit=True,
                can_delete=False,
                can_invite_others=False
            )
            
            messages.success(request, f'You have successfully joined the collaborative project "{project.title}"!')
            return redirect('writer:chapter_editor', project_id=project.id)
        
    except Exception as e:
        messages.error(request, f'Error joining project: {str(e)}')
        return redirect('writer:personal_library')


# Character Views
class CharacterListView(LoginRequiredMixin, ListView):
    model = Character
    template_name = 'writer/character_list.html'
    context_object_name = 'characters'
    
    def get_queryset(self):
        project_id = self.kwargs.get('project_id')
        return Character.objects.filter(
            project_id=project_id,
            project__author=self.request.user
        )
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id, author=self.request.user)
        return context


class GlobalCharacterListView(LoginRequiredMixin, ListView):
    """Global view showing all characters from all user's projects"""
    model = Character
    template_name = 'writer/all_characters.html'
    context_object_name = 'characters'
    
    def get_queryset(self):
        return Character.objects.filter(
            project__author=self.request.user
        ).select_related('project').order_by('name')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Get all user projects for the "Add Character" functionality
        context['user_projects'] = Project.objects.filter(author=self.request.user)
        return context


class CharacterDetailView(LoginRequiredMixin, DetailView):
    model = Character
    template_name = 'writer/character_detail.html'
    context_object_name = 'character'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)


class CharacterCreateView(LoginRequiredMixin, CreateView):
    model = Character
    form_class = CharacterForm
    template_name = 'writer/character_creation_portal.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id, author=self.request.user)
        return context
    
    def form_valid(self, form):
        project_id = self.kwargs.get('project_id')
        form.instance.project = get_object_or_404(Project, id=project_id, author=self.request.user)
        return super().form_valid(form)
    
    def get_success_url(self):
        project_id = self.kwargs.get('project_id')
        return reverse_lazy('writer:character_universe', kwargs={'project_id': project_id})


class CharacterUpdateView(LoginRequiredMixin, UpdateView):
    model = Character
    form_class = CharacterForm
    template_name = 'writer/character_creation_portal.html'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:character_universe', kwargs={'project_id': self.object.project.pk})


class CharacterDeleteView(LoginRequiredMixin, DeleteView):
    model = Character
    template_name = 'writer/character_confirm_delete.html'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:character_universe', kwargs={'project_id': self.object.project.pk})


@login_required
def character_universe(request, project_id):
    """
    Enhanced Character Universe View - Comic book inspired character management
    Makes you feel like a god creating a world of people
    """
    project = get_object_or_404(Project, id=project_id)
    
    # Check permissions
    if not (project.author == request.user or 
            project.collaborators.filter(id=request.user.id).exists()):
        messages.error(request, "You don't have permission to view this project's characters.")
        return redirect('writer:dashboard')
    
    characters = Character.objects.filter(project=project).order_by('name')
    
    # Calculate universe statistics
    protagonists_count = characters.filter(role__icontains='protagonist').count()
    antagonists_count = characters.filter(role__icontains='antagonist').count()
    relationships_count = characters.exclude(relationships__isnull=True).exclude(relationships__exact='').count()
    
    context = {
        'project': project,
        'characters': characters,
        'protagonists_count': protagonists_count,
        'antagonists_count': antagonists_count,
        'relationships_count': relationships_count,
        'total_characters': characters.count(),
    }
    
    return render(request, 'writer/character_universe.html', context)


# Chapter Views (Updated)
class ChapterDetailView(LoginRequiredMixin, DetailView):
    model = Chapter
    template_name = 'writer/chapter_detail.html'
    context_object_name = 'chapter'
    
    def get_queryset(self):
        return Chapter.objects.filter(
            Q(project__author=self.request.user) | Q(project__collaborators=self.request.user)
        ).distinct()


class ChapterCreateView(LoginRequiredMixin, CreateView):
    model = Chapter
    form_class = ChapterForm
    template_name = 'writer/chapter_form.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id)
        return context
    
    def form_valid(self, form):
        project_id = self.kwargs.get('project_id')
        project = get_object_or_404(Project, id=project_id)
        
        # Check permissions
        if not (project.author == self.request.user or 
                project.collaborators.filter(id=self.request.user.id).exists()):
            messages.error(self.request, "You don't have permission to add chapters to this project.")
            return redirect('writer:project_detail', pk=project.id)
        
        form.instance.project = project
        form.instance.last_edited_by = self.request.user
        
        if not form.instance.order:
            # Set order to be the next chapter
            last_chapter = Chapter.objects.filter(project=project).order_by('-order').first()
            form.instance.order = (last_chapter.order + 1) if last_chapter else 1
        
        return super().form_valid(form)


class ChapterUpdateView(LoginRequiredMixin, UpdateView):
    model = Chapter
    form_class = ChapterForm
    template_name = 'writer/chapter_form.html'
    
    def get_queryset(self):
        return Chapter.objects.filter(
            Q(project__author=self.request.user) | Q(project__collaborators=self.request.user)
        ).distinct()
    
    def form_valid(self, form):
        form.instance.last_edited_by = self.request.user
        return super().form_valid(form)


class ChapterDeleteView(LoginRequiredMixin, DeleteView):
    model = Chapter
    template_name = 'writer/chapter_confirm_delete.html'
    
    def get_queryset(self):
        return Chapter.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:project_detail', kwargs={'pk': self.object.project.pk})


@login_required
def chapter_editor(request, project_id):
    """Integrated chapter management and editor view"""
    # Allow both author and collaborators
    project = get_object_or_404(Project, id=project_id)
    if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
        return JsonResponse({'status': 'error', 'message': 'Permission denied'}, status=403)
    all_chapters = Chapter.objects.filter(project=project).order_by('order')

    # Ensure at least one chapter exists for this project
    if not all_chapters.exists():
        default_chapter = Chapter.objects.create(
            title="Untitled Chapter",
            content="",
            project=project,
            last_edited_by=request.user,
            order=0
        )
        all_chapters = Chapter.objects.filter(project=project).order_by('order')

    # Get the current chapter being edited (default to first or create one)
    current_chapter_id = request.GET.get('chapter_id')
    current_chapter = None

    if current_chapter_id:
        current_chapter = get_object_or_404(Chapter, id=current_chapter_id, project=project)
    else:
        current_chapter = all_chapters.first()
    
    if request.method == 'POST':
        action = request.POST.get('action')

        # Map TinyMCE actions to chapter actions for compatibility
        if action == 'save':
            action = 'save_chapter'
        elif action == 'save_to_library':
            action = 'publish_chapter'

        if action == 'reorder':
            # Handle drag-and-drop reordering
            chapter_ids = request.POST.getlist('chapter_ids[]')
            for index, chapter_id in enumerate(chapter_ids):
                Chapter.objects.filter(id=chapter_id, project=project).update(order=index)
            return JsonResponse({'status': 'success'})
        
        elif action == 'save_chapter':
            # Handle chapter content saving
            chapter_id = request.POST.get('chapter_id')
            title = request.POST.get('title', 'Untitled Chapter')
            content = request.POST.get('content', '')
            
            if chapter_id:
                # Update existing chapter
                chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                chapter.title = title
                chapter.content = content
                chapter.save()
                # After save, fix duplicate orders if any
                chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
                for idx, ch in enumerate(chapters):
                    if ch.order != idx:
                        ch.order = idx
                        ch.save(update_fields=['order'])
                return JsonResponse({
                    'status': 'success',
                    'word_count': chapter.word_count,
                    'message': 'Chapter saved successfully'
                })
            else:
                # Create new chapter if none exists
                # Always assign next available order
                used_orders = set(Chapter.objects.filter(project=project).values_list('order', flat=True))
                order = 0
                while order in used_orders:
                    order += 1
                chapter = Chapter.objects.create(
                    title=title,
                    content=content,
                    project=project,
                    order=order
                )
                # After create, fix duplicate orders if any
                chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
                for idx, ch in enumerate(chapters):
                    if ch.order != idx:
                        ch.order = idx
                        ch.save(update_fields=['order'])
                return JsonResponse({
                    'status': 'success',
                    'chapter_id': chapter.id,
                    'word_count': chapter.word_count,
                    'message': 'New chapter created and saved successfully',
                    'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={chapter.id}'
                })
        
        elif action == 'create_chapter':
            # Handle creating new chapter
            title = request.POST.get('title', 'Untitled Chapter')
            heading_level = request.POST.get('heading_level', 'h2')
            
            # Get the next order number
            max_order = Chapter.objects.filter(project=project).aggregate(
                max_order=models.Max('order')
            )['max_order'] or -1
            
            # Create heading content based on level
            heading_tag = f"<{heading_level}>{title}</{heading_level}>"
            content = f"{heading_tag}<p></p>"
            
            chapter = Chapter.objects.create(
                title=title,
                content=content,
                project=project,
                order=max_order + 1
            )
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': chapter.id,
                'chapter_title': chapter.title,
                'chapter_order': chapter.order + 1,
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={chapter.id}'
            })
        
        elif action == 'delete_chapter':
            # Handle chapter deletion
            chapter_id = request.POST.get('chapter_id')
            chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
            chapter.delete()
            
            # Reorder remaining chapters
            remaining_chapters = Chapter.objects.filter(project=project).order_by('order')
            for index, ch in enumerate(remaining_chapters):
                ch.order = index
                ch.save()
            
            return JsonResponse({'status': 'success'})
        
        elif action == 'publish_chapter':
            # Handle publishing chapter to library
            chapter_id = request.POST.get('chapter_id')
            title = request.POST.get('title', 'Untitled Chapter')
            content = request.POST.get('content', '')
            if chapter_id:
                chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                chapter.title = title
                chapter.content = content
                chapter.is_published = True
                chapter.save()
                return JsonResponse({
                    'status': 'success',
                    'message': 'Chapter published successfully',
                    'word_count': chapter.word_count
                })
            else:
                return JsonResponse({'status': 'error', 'message': 'No chapter selected to publish.'})
        elif action == 'split_long_chapter':
            # Handle automatic chapter splitting when content is too long
            try:
                import json
                data = json.loads(request.body)
                chapter_id = data.get('chapter_id')
                first_part = data.get('first_part')
                remaining_part = data.get('remaining_part')
                
                current_chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                
                # Update current chapter with first part
                current_chapter.content = first_part
                current_chapter.save()
                
                # Create new chapter with remaining content
                new_chapter = Chapter.objects.create(
                    title="Untitled Chapter (Auto-split)",
                    content=remaining_part,
                    project=project,
                    order=current_chapter.order + 1
                )
                
                # Reorder subsequent chapters
                subsequent_chapters = Chapter.objects.filter(
                    project=project, 
                    order__gt=current_chapter.order
                ).exclude(id=new_chapter.id)
                
                for chapter in subsequent_chapters:
                    chapter.order += 1
                    chapter.save()
                
                return JsonResponse({
                    'success': True,
                    'new_chapter_title': new_chapter.title,
                    'new_chapter_id': new_chapter.id
                })
                
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)})
        
        elif action == 'create_chapter_from_heading':
            # Handle creating new chapter from H2 heading
            title = request.POST.get('title', 'Untitled Chapter')
            current_chapter_id = request.POST.get('current_chapter_id')
            content_before = request.POST.get('content_before', '')
            content_after = request.POST.get('content_after', '')
            
            # Update current chapter content (remove the heading and content after)
            if current_chapter_id:
                current_chapter = get_object_or_404(Chapter, id=current_chapter_id, project=project)
                current_chapter.content = content_before
                current_chapter.save()
            
            # Get the next order number (insert after current chapter)
            current_order = current_chapter.order if current_chapter_id else -1
            
            # Shift all chapters after current chapter
            chapters_to_shift = Chapter.objects.filter(project=project, order__gt=current_order)
            for ch in chapters_to_shift:
                ch.order += 1
                ch.save()
            
            # Create new chapter
            new_chapter = Chapter.objects.create(
                title=title,
                content=f"<h2>{title}</h2>{content_after}",
                project=project,
                order=current_order + 1
            )
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': new_chapter.id,
                'chapter_title': new_chapter.title,
                'chapter_order': new_chapter.order + 1,
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={new_chapter.id}'
            })
    
    # Handle create_from_selection action
    if request.method == 'POST' and request.POST.get('action') == 'create_from_selection':
        try:
            selected_text = request.POST.get('selected_text', '').strip()
            new_title = request.POST.get('new_title', '').strip()
            remove_from_current = request.POST.get('remove_from_current', '').lower() == 'true'
            
            if not selected_text:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No text selected'
                })
            
            if not new_title:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Chapter title is required'
                })
            
            # Get the highest order number for new chapter positioning
            current_order = 0
            if current_chapter:
                current_order = current_chapter.order
            
            # Create new chapter with selected text
            new_chapter = Chapter.objects.create(
                title=new_title,
                content=selected_text,
                project=project,
                order=current_order + 1
            )
            
            # Update orders of subsequent chapters
            Chapter.objects.filter(
                project=project,
                order__gt=current_order
            ).exclude(id=new_chapter.id).update(order=F('order') + 1)
            
            # Optionally remove selected text from current chapter
            if remove_from_current and current_chapter:
                # Remove the selected text from current chapter's content
                current_content = current_chapter.content or ''
                if selected_text in current_content:
                    updated_content = current_content.replace(selected_text, '', 1)
                    current_chapter.content = updated_content
                    current_chapter.save()
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': new_chapter.id,
                'chapter_title': new_chapter.title,
                'chapter_order': new_chapter.order,
                'message': f'Chapter "{new_title}" created successfully',
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={new_chapter.id}'
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error creating chapter: {str(e)}'
            })
    
    # Handle direct_upload_to_editor action
    if request.method == 'POST' and request.POST.get('action') == 'direct_upload_to_editor':
        try:
            if 'file' not in request.FILES:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No file selected'
                })
            
            uploaded_file = request.FILES['file']
            insert_position = request.POST.get('insert_position', 'end')  # 'start', 'end', 'cursor'
            
            # Save the uploaded file temporarily
            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_uploads')
            os.makedirs(temp_dir, exist_ok=True)
            
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            # Save the file
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Extract text content
            content = extract_text_from_file(temp_file_path)
            
            if not content:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                
                return JsonResponse({
                    'status': 'error',
                    'message': 'Could not extract text from the uploaded file. Please try a different format.'
                })
            
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            
            # Return the extracted content for insertion
            return JsonResponse({
                'status': 'success',
                'content': content,
                'filename': uploaded_file.name,
                'message': f'File "{uploaded_file.name}" uploaded successfully'
            })
            
        except Exception as e:
            # Clean up temporary file on error
            if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            })
    
    context = {
        'project': project,
        'all_chapters': all_chapters,
        'current_chapter': current_chapter,
        'user': request.user,
        'document': current_chapter,  # Pass chapter as document for compatibility
    }
    return render(request, 'writer/google_docs_editor.html', context)

@login_required
def upload_test(request):
    """Simple upload test page."""
    return render(request, 'writer/upload_test.html')



# Document Views (Updated)
class DocumentListView(LoginRequiredMixin, ListView):
    model = Document
    template_name = 'writer/document_list.html'
    context_object_name = 'documents'
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentDetailView(LoginRequiredMixin, DetailView):
    model = Document
    template_name = 'writer/document_detail.html'
    context_object_name = 'document'
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentCreateView(LoginRequiredMixin, CreateView):
    model = Document
    form_class = DocumentForm
    template_name = 'writer/document_form.html'
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs
    
    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)


class DocumentUpdateView(LoginRequiredMixin, UpdateView):
    model = Document
    form_class = DocumentForm
    template_name = 'writer/document_form.html'
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentDeleteView(LoginRequiredMixin, DeleteView):
    model = Document
    template_name = 'writer/document_confirm_delete.html'
    success_url = reverse_lazy('writer:document_list')
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


# Document Import Views
@login_required
def import_document(request):
    if request.method == 'POST':
        form = ImportDocumentForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            imported_doc = form.save(commit=False)
            imported_doc.user = request.user
            
            # Handle Google Docs URL
            google_docs_url = form.cleaned_data.get('google_docs_url')
            if google_docs_url:
                # Extract content from Google Docs
                imported_doc.google_docs_url = google_docs_url
                extracted_content = extract_text_from_file(None, google_docs_url=google_docs_url)
                imported_doc.import_type = 'google_docs'
                imported_doc.file_size = 0  # No file uploaded
                
                # Save without file
                imported_doc.save()
                imported_doc.extracted_content = extracted_content
            else:
                # Handle file upload
                imported_doc.file_size = imported_doc.original_file.size
                
                # Determine file type
                file_name = imported_doc.original_file.name.lower()
                if file_name.endswith('.pdf'):
                    imported_doc.import_type = 'pdf'
                elif file_name.endswith(('.docx', '.doc')):
                    imported_doc.import_type = 'docx'
                elif file_name.endswith('.txt'):
                    imported_doc.import_type = 'txt'
                elif file_name.endswith('.rtf'):
                    imported_doc.import_type = 'rtf'
                elif file_name.endswith('.odt'):
                    imported_doc.import_type = 'odt'
                elif file_name.endswith(('.html', '.htm')):
                    imported_doc.import_type = 'html'
                
                # Save the file first to get the path
                imported_doc.save()
                
                # Extract text content from the saved file
                temp_file_path = imported_doc.original_file.path
                extracted_content = extract_text_from_file(temp_file_path)
            
            # Set target project if selected, otherwise create a new project
            target_project = form.cleaned_data.get('target_project')
            if not target_project:
                # Auto-create a new project based on the imported document
                project_title = imported_doc.title
                if len(project_title) > 50:  # Ensure title isn't too long
                    project_title = project_title[:47] + "..."
                
                # Create new project
                target_project = Project.objects.create(
                    title=project_title,
                    description=f"Project created from imported document: {imported_doc.title}",
                    author=request.user,
                    target_word_count=10000,  # Default word count
                    genre="General",  # Default genre
                )
                
                imported_doc.project = target_project
            
            # Apply formatting options
            if not form.cleaned_data.get('preserve_formatting', True):
                # Strip HTML formatting if user doesn't want it preserved
                import re
                extracted_content = re.sub(r'<[^>]+>', '', extracted_content)
                # Convert back to simple paragraphs
                paragraphs = [p.strip() for p in extracted_content.split('\n') if p.strip()]
                extracted_content = ''.join(f'<p>{p}</p>' for p in paragraphs)
            
            imported_doc.extracted_content = extracted_content
            
            # Save again with extracted content
            imported_doc.save()
            
            # Auto-create chapters if requested, or create a single chapter if no target project was selected
            if form.cleaned_data.get('auto_create_chapters', False) and target_project:
                create_chapters_from_content(imported_doc.extracted_content, target_project, imported_doc.title)
                messages.success(request, f"Document '{imported_doc.title}' imported and chapters created successfully!")
            elif target_project and not form.cleaned_data.get('auto_create_chapters', False):
                # Create a single chapter with all the content
                Chapter.objects.create(
                    title=imported_doc.title,
                    content=imported_doc.extracted_content,
                    project=target_project,
                    order=1
                )
                messages.success(request, f"Document '{imported_doc.title}' imported successfully and converted to a new project!")
            else:
                messages.success(request, f"Document '{imported_doc.title}' imported successfully!")
            
            return redirect('writer:import_detail', pk=imported_doc.pk)
    else:
        form = ImportDocumentForm(user=request.user)
    
    return render(request, 'writer/import_form.html', {'form': form})


def create_chapters_from_content(content, project, base_title):
    """Create chapters from content based on HTML headings"""
    import re
    from bs4 import BeautifulSoup
    
    try:
        soup = BeautifulSoup(content, 'html.parser')
        
        # Find all headings (h1, h2, h3)
        headings = soup.find_all(['h1', 'h2', 'h3'])
        
        if not headings:
            # No headings found, create a single chapter
            Chapter.objects.create(
                title=base_title,
                content=content,
                project=project,
                order=1
            )
            return
        
        current_content = ""
        chapter_order = 1
        
        for i, heading in enumerate(headings):
            # Get content between this heading and the next
            current_element = heading.next_sibling
            chapter_content = f"<{heading.name}>{heading.get_text()}</{heading.name}>"
            
            while current_element and current_element != headings[i + 1] if i + 1 < len(headings) else True:
                if hasattr(current_element, 'name') and current_element.name:
                    chapter_content += str(current_element)
                elif isinstance(current_element, str) and current_element.strip():
                    chapter_content += current_element
                current_element = current_element.next_sibling
                if i + 1 < len(headings) and current_element == headings[i + 1]:
                    break
            
            # Create chapter
            Chapter.objects.create(
                title=heading.get_text().strip() or f"{base_title} - Chapter {chapter_order}",
                content=chapter_content,
                project=project,
                order=chapter_order
            )
            chapter_order += 1
            
    except Exception as e:
        # Fallback: create single chapter
        Chapter.objects.create(
            title=base_title,
            content=content,
            project=project,
            order=1
        )


@login_required
def import_detail(request, pk):
    imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
    return render(request, 'writer/import_detail.html', {'imported_doc': imported_doc})


@login_required
def convert_import_to_project(request, pk):
    """Convert an imported document to a new project with chapters"""
    if request.method == 'POST':
        imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
        
        # Create a new project
        project = Project.objects.create(
            title=imported_doc.title,
            description=f"Project created from imported document: {imported_doc.title}",
            author=request.user,
            genre='Other'  # Default genre
        )
        
        # Create a chapter with the imported content
        chapter = Chapter.objects.create(
            title=f"Chapter 1: {imported_doc.title}",
            content=imported_doc.extracted_content or "No content available",
            project=project,
            order=1
        )
        
        messages.success(request, f"Successfully converted '{imported_doc.title}' to project '{project.title}'!")
        return JsonResponse({
            'status': 'success',
            'project_id': project.id,
            'redirect_url': reverse('writer:chapter_editor', kwargs={'project_id': project.id})
        })
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


@login_required
def add_import_to_project(request, pk, project_id):
    """Add imported document content as a new chapter to an existing project"""
    if request.method == 'POST':
        imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
        project = get_object_or_404(Project, id=project_id, author=request.user)
        
        # Get the next order for the new chapter
        last_chapter = Chapter.objects.filter(project=project).order_by('-order').first()
        next_order = (last_chapter.order + 1) if last_chapter else 1
        
        # Create a new chapter with the imported content
        chapter = Chapter.objects.create(
            title=f"Chapter {next_order}: {imported_doc.title}",
            content=imported_doc.extracted_content or "No content available",
            project=project,
            order=next_order
        )
        
        messages.success(request, f"Successfully added '{imported_doc.title}' as a new chapter to '{project.title}'!")
        return JsonResponse({
            'status': 'success',
            'project_id': project.id,
            'chapter_id': chapter.id,
            'redirect_url': reverse('writer:chapter_editor', kwargs={'project_id': project.id}) + f'?chapter_id={chapter.id}'
        })
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


@csrf_exempt
@login_required
def update_import_content(request, pk):
    """Update the content of an imported document"""
    if request.method == 'POST':
        try:
            imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
            data = json.loads(request.body)
            new_content = data.get('content', '')
            
            # Update the extracted content
            imported_doc.extracted_content = new_content
            imported_doc.save()
            
            return JsonResponse({
                'status': 'success',
                'message': 'Content updated successfully'
            })
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error updating content: {str(e)}'
            })
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.authentication import TokenAuthentication, SessionAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def projects_api_list(request):
    """API endpoint to list user's projects"""
    # Check if we should only return bookshelf projects
    bookshelf_only = request.GET.get('bookshelf_only', 'false').lower() == 'true'
    
    if bookshelf_only:
        projects = Project.objects.filter(author=request.user, show_on_dashboard=True).order_by('-updated_at')
    else:
        projects = Project.objects.filter(author=request.user).order_by('-updated_at')
    
    projects_data = [
        {
            'id': project.id,
            'title': project.title,
            'description': project.description or '',
            'updated_at': project.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            'created_at': project.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'chapter_count': project.chapters.count(),
            'word_count': project.total_word_count,
            'genre': project.genre or '',
            'target_word_count': project.target_word_count,
            'show_on_dashboard': project.show_on_dashboard,
            'progress_percentage': round(project.progress_percentage, 1),
        }
        for project in projects
    ]
    return Response({'projects': projects_data})


@login_required
@csrf_exempt
def chapters_api(request, project_id):
    """API endpoint for chapter operations"""
    try:
        project = Project.objects.get(id=project_id)
        # Check if user has access to this project
        if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({'error': 'Permission denied'}, status=403)
    except Project.DoesNotExist:
        return JsonResponse({'error': 'Project not found'}, status=404)
    
    if request.method == 'GET':
        chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
        chapters_data = [
            {
                'id': chapter.id,
                'title': chapter.title,
                'content': chapter.content or '',
                'word_count': chapter.word_count,
                'order': chapter.order,
                'created_at': chapter.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': chapter.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            }
            for chapter in chapters
        ]
        return JsonResponse({'chapters': chapters_data})
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            chapter = Chapter.objects.create(
                title=data.get('title', 'New Chapter'),
                content=data.get('content', ''),
                project=project,
                order=data.get('order', 0)
            )
            return JsonResponse({
                'success': True,
                'chapter': {
                    'id': chapter.id,
                    'title': chapter.title,
                    'content': chapter.content,
                    'word_count': chapter.word_count,
                    'order': chapter.order,
                }
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)


@api_view(['GET', 'PUT'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def chapter_api(request, chapter_id):
    """API endpoint for individual chapter operations"""
    try:
        chapter = Chapter.objects.get(id=chapter_id)
        # Check if user has access to this chapter's project
        if not (chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists()):
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)
    except Chapter.DoesNotExist:
        return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)
    
    if request.method == 'GET':
        return Response({
            'id': chapter.id,
            'title': chapter.title,
            'content': chapter.content or '',
            'word_count': chapter.word_count,
            'order': chapter.order,
            'project_id': chapter.project.id,
            'project_title': chapter.project.title,
            'created_at': chapter.created_at.isoformat(),
            'updated_at': chapter.updated_at.isoformat(),
        })
    
    elif request.method == 'PUT':
        try:
            # Handle both JSON data from DRF and direct request.data
            data = request.data if hasattr(request, 'data') else json.loads(request.body)
            if 'title' in data:
                chapter.title = data['title']
            if 'content' in data:
                chapter.content = data['content']
            if 'order' in data:
                chapter.order = data['order']
            
            chapter.save()
            
            return Response({
                'success': True,
                'chapter': {
                    'id': chapter.id,
                    'title': chapter.title,
                    'content': chapter.content,
                    'word_count': chapter.word_count,
                }
            })
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class ImportedDocumentDeleteView(LoginRequiredMixin, DeleteView):
    model = ImportedDocument
    template_name = 'writer/import_confirm_delete.html'
    success_url = reverse_lazy('writer:personal_library')
    
    def get_queryset(self):
        return ImportedDocument.objects.filter(user=self.request.user)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['object_type'] = 'Imported Document'
        return context


# AJAX Views
@login_required
@require_http_methods(["POST"])
def ajax_create_chapter(request):
    try:
        import json
        data = json.loads(request.body)
        
        project_id = data.get('project_id')
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        
        if not all([project_id, title, content]):
            return JsonResponse({'success': False, 'error': 'Missing required fields'})
        
        # Get the project and verify ownership
        try:
            project = Project.objects.get(id=project_id)
            # Check if user has access to this project
            if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'success': False, 'error': 'Permission denied'})
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        
        # Create the chapter
        chapter = Chapter.objects.create(
            project=project,
            title=title,
            content=content,
            order=project.chapters.count() + 1
        )
        
        # Update project word count
        project.update_word_count()
        
        return JsonResponse({
            'success': True, 
            'chapter_id': chapter.id,
            'message': f'Chapter "{title}" created successfully!'
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def dashboard(request):
    recent_documents = Document.objects.filter(author=request.user)[:5]
    
    total_documents = Document.objects.filter(author=request.user).count()
    total_projects = Project.objects.filter(
        Q(author=request.user) | Q(collaborators=request.user)
    ).distinct().count()
    
    # Calculate comprehensive writing statistics
    user_projects = Project.objects.filter(author=request.user).select_related('author')
    
    # Total words from both documents and projects
    doc_words = sum(doc.word_count for doc in Document.objects.filter(author=request.user))
    project_words = sum(project.total_word_count for project in user_projects)
    total_words = doc_words + project_words
    
    # Get user's characters
    user_characters = Character.objects.filter(
        project__in=Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct()
    ).count()
    
    # Create shelves with max 7 books each
    shelves = []
    projects_list = list(user_projects)
    for i in range(0, len(projects_list), 7):
        shelf = projects_list[i:i+7]
        shelves.append(shelf)
    
    # Get user profile for theme preferences
    profile = UserProfile.get_or_create_for_user(request.user)
    
    context = {
        'recent_documents': recent_documents,
        'total_documents': total_documents,
        'total_projects': total_projects,
        'total_words': total_words,
        'total_characters': user_characters,
        'user_projects': user_projects,
        'shelves': shelves,
        'show_bookshelf': True,
        'profile': profile,
    }
    return render(request, 'writer/dashboard.html', context)


# Format & Export Views
@login_required
def format_page(request):
    """Format page for styling and export options"""
    user_projects = Project.objects.filter(
        Q(author=request.user) | Q(collaborators=request.user)
    ).distinct()
    
    context = {
        'user_projects': user_projects,
    }
    return render(request, 'writer/format_page.html', context)


@login_required
def export_project(request, project_id, format_type):
    """Export project in various formats"""
    project = get_object_or_404(Project, id=project_id)
    
    # Check permissions
    if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
        messages.error(request, "You don't have permission to export this project.")
        return redirect('writer:format_page')
    
    # Get all chapters in order
    chapters = Chapter.objects.filter(project=project).order_by('order')
    
    if format_type == 'pdf':
        return export_as_pdf(project, chapters)
    elif format_type == 'epub':
        return export_as_epub(project, chapters)
    elif format_type == 'docx':
        return export_as_docx(project, chapters)
    elif format_type == 'gdoc':
        return export_as_google_doc(project, chapters)
    else:
        messages.error(request, "Unsupported export format.")
        return redirect('writer:format_page')


def export_as_pdf(project, chapters):
    """Export project as PDF"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    import io
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title page
    title = Paragraph(f"<para align=center><font size=24><b>{project.title}</b></font></para>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 72))
    
    if project.description:
        desc = Paragraph(f"<para align=center><i>{project.description}</i></para>", styles['Normal'])
        story.append(desc)
        story.append(Spacer(1, 36))
    
    author = Paragraph(f"<para align=center>by {project.author.get_full_name() or project.author.username}</para>", styles['Normal'])
    story.append(author)
    story.append(Spacer(1, 72))
    
    # Chapters
    for chapter in chapters:
        if chapter.title:
            chapter_title = Paragraph(f"<para><font size=18><b>{chapter.title}</b></font></para>", styles['Heading1'])
            story.append(chapter_title)
            story.append(Spacer(1, 12))
        
        if chapter.content:
            # Clean HTML content for PDF
            import re
            clean_content = re.sub('<[^<]+?>', '', chapter.content)
            content = Paragraph(clean_content, styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{project.title}.pdf"'
    return response


def export_as_epub(project, chapters):
    """Export project as EPUB"""
    from ebooklib import epub
    import io
    
    book = epub.EpubBook()
    book.set_identifier(f'project_{project.id}')
    book.set_title(project.title)
    book.set_language('en')
    book.add_author(project.author.get_full_name() or project.author.username)
    
    if project.description:
        book.add_metadata('DC', 'description', project.description)
    
    # Create chapters
    epub_chapters = []
    for i, chapter in enumerate(chapters):
        epub_chapter = epub.EpubHtml(
            title=chapter.title or f'Chapter {i+1}',
            file_name=f'chapter_{i+1}.xhtml',
            lang='en'
        )
        epub_chapter.content = f'''
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head><title>{chapter.title or f'Chapter {i+1}'}</title></head>
        <body>
        <h1>{chapter.title or f'Chapter {i+1}'}</h1>
        {chapter.content or ''}
        </body>
        </html>
        '''
        book.add_item(epub_chapter)
        epub_chapters.append(epub_chapter)
    
    # Create table of contents
    book.toc = epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Create spine
    book.spine = ['nav'] + epub_chapters
    
    # Generate EPUB
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/epub+zip')
    response['Content-Disposition'] = f'attachment; filename="{project.title}.epub"'
    return response


def export_as_docx(project, chapters):
    """Export project as Word document"""
    from docx import Document
    from docx.shared import Inches
    import io
    
    doc = Document()
    
    # Title page
    title = doc.add_heading(project.title, 0)
    title.alignment = 1  # Center alignment
    
    if project.description:
        desc_para = doc.add_paragraph(project.description)
        desc_para.alignment = 1
    
    author_para = doc.add_paragraph(f"by {project.author.get_full_name() or project.author.username}")
    author_para.alignment = 1
    
    doc.add_page_break()
    
    # Chapters
    for chapter in chapters:
        if chapter.title:
            doc.add_heading(chapter.title, 1)
        
        if chapter.content:
            # Clean HTML content
            import re
            clean_content = re.sub('<[^<]+?>', '', chapter.content)
            doc.add_paragraph(clean_content)
        
        doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{project.title}.docx"'
    return response


def export_as_google_doc(project, chapters):
    """Export project as Google Doc (HTML format for import)"""
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{project.title}</title>
        <meta charset="utf-8">
        <style>
            body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 1in; }}
            h1 {{ text-align: center; page-break-before: always; }}
            h2 {{ margin-top: 2em; }}
            .title-page {{ text-align: center; margin-bottom: 3em; }}
            .chapter {{ page-break-before: always; }}
        </style>
    </head>
    <body>
        <div class="title-page">
            <h1>{project.title}</h1>
            {f'<p><i>{project.description}</i></p>' if project.description else ''}
            <p>by {project.author.get_full_name() or project.author.username}</p>
        </div>
    """
    
    for chapter in chapters:
        html_content += f'<div class="chapter">'
        if chapter.title:
            html_content += f'<h2>{chapter.title}</h2>'
        if chapter.content:
            html_content += chapter.content
        html_content += '</div>'
    
    html_content += """
    </body>
    </html>
    """
    
    response = HttpResponse(html_content, content_type='text/html')
    response['Content-Disposition'] = f'attachment; filename="{project.title}_for_google_docs.html"'
    return response


# AI Assistant Views (Enhanced)
@login_required
def ai_assistant(request):
    if request.method == 'POST':
        form = AIAssistanceForm(request.POST)
        if form.is_valid():
            assistance = form.save(commit=False)
            assistance.user = request.user
            
            # Generate AI response
            assistance.response = generate_ai_response(
                assistance.content, 
                assistance.assistance_type, 
                assistance.prompt
            )
            assistance.save()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'response': assistance.response
                })
            else:
                messages.success(request, 'AI assistance generated successfully!')
                return redirect('writer:ai_assistant')
    else:
        form = AIAssistanceForm()
    
    recent_requests = AIAssistanceRequest.objects.filter(user=request.user)[:10]
    context = {
        'form': form,
        'recent_requests': recent_requests
    }
    return render(request, 'writer/ai_assistant.html', context)


def generate_ai_response(content, assistance_type, prompt):
    """Enhanced AI response generator with character and plot development"""
    responses = {
        'brainstorm': [
            "Here are some creative ideas to expand your story:\n\n• What if your protagonist discovers a hidden secret about their past?\n• Consider adding a subplot involving a secondary character\n• Explore the consequences of your main character's decisions\n• Add unexpected obstacles or complications\n• Introduce a mentor figure who challenges your protagonist's beliefs",
            "To develop this further, consider:\n\n• Character motivations and backstory\n• Setting details that affect the plot\n• Potential conflicts between characters\n• Themes you want to explore\n• How the environment shapes your characters",
        ],
        'character': [
            "Character development suggestions:\n\n• Give your character a specific fear or weakness\n• Create a contradiction in their personality\n• Develop their unique voice and speech patterns\n• Consider their relationships with family and friends\n• What drives them? What do they want most?\n• Add physical quirks or habits that reveal personality",
            "To make this character more compelling:\n\n• Create internal conflict that mirrors external conflict\n• Give them a secret they're hiding\n• Develop their backstory - what shaped them?\n• Consider how they've changed throughout the story\n• What would they never do? What would force them to do it?",
        ],
        'dialogue': [
            "Dialogue enhancement tips:\n\n• Each character should have a distinct voice\n• Use subtext - characters rarely say exactly what they mean\n• Add action beats between dialogue\n• Show character relationships through how they speak to each other\n• Cut unnecessary dialogue tags\n• Make every line serve a purpose",
            "To improve this dialogue:\n\n• Add more conflict or tension\n• Use contractions for natural speech\n• Show character emotions through word choice\n• Consider what characters aren't saying\n• Add interruptions and overlapping speech for realism",
        ],
        'plot': [
            "Plot development ideas:\n\n• Increase the stakes - what happens if your protagonist fails?\n• Add a ticking clock element\n• Create obstacles that force character growth\n• Consider the three-act structure\n• Plant seeds early that pay off later\n• Make sure each scene advances plot or character",
            "To strengthen your plot:\n\n• Give your antagonist clear, understandable motivations\n• Create moments where victory seems impossible\n• Add plot twists that feel inevitable in hindsight\n• Make sure cause and effect are clear\n• Consider multiple storylines that intersect",
        ],
        'edit': [
            "Here are some suggestions to improve your writing:\n\n• Consider varying your sentence structure for better flow\n• Some paragraphs could be shortened for better pacing\n• Look for opportunities to show rather than tell\n• Strong dialogue! Consider adding more action beats between speeches\n• Watch for repeated words or phrases",
            "Editorial suggestions:\n\n• The opening could be stronger with more immediate action\n• Character descriptions are vivid and engaging\n• Consider tightening some exposition\n• The dialogue feels natural and authentic\n• Look for places to add sensory details",
        ],
        'continue': [
            "Here's a possible continuation:\n\nAs the door creaked open, Sarah held her breath. The hallway beyond was darker than she'd expected, and the musty smell of old books filled her nostrils. Each step forward felt like crossing into another world, where the familiar rules no longer applied...",
            "To continue this scene, you might consider:\n\nHaving your character notice something unexpected, introduce a new character, or reveal important information that changes everything. The tension you've built is perfect for a major revelation or plot twist.",
        ],
        'rewrite': [
            "Here's a rewritten version with improved flow:\n\n[Your original text would be revised here with better sentence structure, clearer descriptions, and enhanced dialogue while maintaining your unique voice and style.]",
            "I've rewritten this section to be more concise and impactful:\n\n[The rewritten version would focus on stronger verbs, clearer imagery, and better pacing while preserving the essence of your original writing.]",
        ],
        'grammar': [
            "Grammar and style check complete:\n\n• Fixed several comma splices\n• Corrected subject-verb agreement in paragraph 2\n• Suggested alternatives for repetitive word usage\n• Overall, your writing is clear and well-structured!\n• Consider breaking up some longer sentences",
            "Grammar review:\n\n✓ Sentence structure is generally good\n• Watch for run-on sentences in longer paragraphs\n• Consider breaking up some complex sentences\n• Punctuation is mostly correct\n• Strong vocabulary choices throughout",
        ],
        'style': [
            "Style improvements:\n\n• Vary your sentence beginnings for better rhythm\n• Consider using more active voice\n• Your descriptive language is engaging\n• Look for opportunities to eliminate redundant words\n• Add more specific, concrete details",
            "To enhance your writing style:\n\n• Strong character voice comes through clearly\n• Consider adding more sensory details\n• Good use of dialogue to advance the plot\n• Pacing works well for this genre\n• Consider the emotional arc of your scenes",
        ]
    }
    
    return random.choice(responses.get(assistance_type, ["I'm here to help! Please provide more specific guidance for your request."]))


@csrf_exempt
@login_required
def auto_save(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        document_id = data.get('document_id')
        chapter_id = data.get('chapter_id')
        content = data.get('content')
        title = data.get('title', 'Untitled')
        
        if chapter_id:
            # Auto-save chapter
            chapter = get_object_or_404(Chapter, id=chapter_id)
            # Check permissions
            if not (chapter.project.author == request.user or 
                   chapter.project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'status': 'error', 'message': 'Permission denied'})
            
            chapter.content = content
            chapter.title = title
            chapter.last_edited_by = request.user
            chapter.save()
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': chapter.id,
                'word_count': chapter.word_count
            })
        elif document_id:
            # Auto-save document
            document = get_object_or_404(Document, id=document_id, author=request.user)
            document.content = content
            document.title = title
            document.save()
            return JsonResponse({
                'status': 'success',
                'document_id': document.id,
                'word_count': document.word_count
            })
        else:
            # Create new document
            document = Document.objects.create(
                title=title,
                content=content,
                author=request.user
            )
            return JsonResponse({
                'status': 'success',
                'document_id': document.id,
                'word_count': document.word_count
            })
    
    return JsonResponse({'status': 'error'})


@login_required
def upload_file(request):
    """Simple file upload for chapter editor."""
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        try:
            # Get file path for processing
            import tempfile
            import os
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            # Extract text content and identify chapters
            from .document_parser import extract_text_and_chapters_from_file
            result = extract_text_and_chapters_from_file(temp_file_path)
            
            # Clean up temporary file
            os.unlink(temp_file_path)
            
            # Return content and chapter information
            response_data = {
                'success': True,
                'content': result['full_text'],
                'title': uploaded_file.name.rsplit('.', 1)[0]  # Remove extension
            }
            
            # Add chapter information if detected
            if result['chapters'] and len(result['chapters']) > 1:
                response_data.update({
                    'has_chapters': True,
                    'chapters': result['chapters'],
                    'chapter_count': result['chapter_count'],
                    'total_words': result['total_words'],
                    'average_confidence': result['average_confidence']
                })
            else:
                response_data['has_chapters'] = False
            
            return JsonResponse(response_data)
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    return JsonResponse({
        'success': False,
        'error': 'No file uploaded'
    })


@login_required
@require_http_methods(["POST"])
def reorder_chapters(request, project_id):
    """Handle drag-and-drop reordering of chapters"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    
    try:
        data = json.loads(request.body)
        chapter_orders = data.get('chapter_orders', [])
        
        # Update chapter orders
        for item in chapter_orders:
            chapter_id = item.get('id')
            new_order = item.get('order')
            Chapter.objects.filter(id=chapter_id, project=project).update(order=new_order)
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


@csrf_exempt
@login_required
@require_http_methods(["POST"])
def update_chapter_order(request, project_id, chapter_id):
    """Manually update chapter order by number input"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
    
    try:
        # Handle both JSON and form data for mobile app compatibility
        if request.content_type == 'application/json' or 'application/json' in request.META.get('CONTENT_TYPE', ''):
            import json
            data = json.loads(request.body)
            new_order = int(data.get('order', 0))
            print(f"Reordering chapter {chapter_id} via JSON: new_order={new_order}")
        else:
            new_order = int(request.POST.get('order', 0))
            print(f"Reordering chapter {chapter_id} via form: new_order={new_order}")
            
        old_order = chapter.order
        
        if new_order != old_order:
            # Adjust other chapters' orders
            if new_order > old_order:
                # Moving down: decrease order of chapters between old and new position
                Chapter.objects.filter(
                    project=project,
                    order__gt=old_order,
                    order__lte=new_order
                ).update(order=F('order') - 1)
            else:
                # Moving up: increase order of chapters between new and old position
                Chapter.objects.filter(
                    project=project,
                    order__gte=new_order,
                    order__lt=old_order
                ).update(order=F('order') + 1)
            
            # Set the new order for the moved chapter
            chapter.order = new_order
            chapter.save()
        
        return JsonResponse({'status': 'success'})
    except (ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Invalid order number'})


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_chapter_list(request, project_id):
    """Get paginated list of chapters for the sidebar"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapters = Chapter.objects.filter(project=project).order_by('order')
    
    chapter_data = []
    for chapter in chapters:
        chapter_data.append({
            'id': chapter.id,
            'title': chapter.title,
            'order': chapter.order,
            'word_count': chapter.word_count,
            'content': chapter.content[:200] if chapter.content else '',  # Add content preview for mobile
            'created_at': chapter.created_at.isoformat(),  # Use ISO format for proper date parsing
            'updated_at': chapter.updated_at.isoformat()  # Use ISO format for proper date parsing
        })
    
    return Response({
        'chapters': chapter_data,
        'total_chapters': len(chapter_data),
        'project_word_count': project.total_word_count
    })


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def create_new_chapter(request, project_id):
    """Create a new chapter"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    
    try:
        # Handle both POST form data and JSON data
        if request.content_type == 'application/json':
            import json
            data = json.loads(request.body)
            title = data.get('title', 'Untitled Chapter')
        else:
            title = request.POST.get('title', 'Untitled Chapter')
        
        # Get the next order number safely
        with transaction.atomic():
            max_order = Chapter.objects.filter(project=project).aggregate(
                max_order=models.Max('order')
            )['max_order']
            
            # Fix: Handle when max_order is 0 (which is falsy in Python)
            next_order = 0 if max_order is None else max_order + 1
            
            chapter = Chapter.objects.create(
                title=title,
                content='',
                project=project,
                order=next_order,
                last_edited_by=request.user
            )
        
        return Response({
            'success': True,
            'chapter': {
                'id': chapter.id,
                'title': chapter.title,
                'order': chapter.order,
                'word_count': chapter.word_count,
                'content': chapter.content,
                'created_at': chapter.created_at.isoformat(),
                'updated_at': chapter.updated_at.isoformat()
            }
        })
    except Exception as e:
        return Response({'success': False, 'error': str(e)})


@login_required
@require_http_methods(["POST"])
def delete_chapter(request, project_id, chapter_id):
    """Delete a chapter and reorder remaining chapters"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
    
    try:
        deleted_order = chapter.order
        chapter.delete()
        
        # Reorder remaining chapters
        Chapter.objects.filter(
            project=project,
            order__gt=deleted_order
        ).update(order=F('order') - 1)
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


# Chapter Management API Endpoints

@login_required
@csrf_exempt 
def create_chapter_api(request):
    """Create a new chapter via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST
            
        project_id = data.get('project')
        title = data.get('title', '').strip()
        content = data.get('content', '')
        
        if not project_id or not title:
            return JsonResponse({'success': False, 'error': 'Project ID and title are required'})
        
        # Get the project and verify access
        try:
            project = Project.objects.get(id=project_id)
            if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'success': False, 'error': 'Permission denied'})
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        
        # Create the chapter
        chapter = Chapter.objects.create(
            project=project,
            title=title,
            content=content,
            order=Chapter.objects.filter(project=project).count() + 1
        )
        
        # Update project word count
        project.update_word_count()
        
        return JsonResponse({
            'success': True, 
            'chapter_id': chapter.id,
            'message': f'Chapter "{title}" created successfully!'
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@csrf_exempt
def delete_chapter_api(request, chapter_id):
    """Delete a chapter via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        chapter = get_object_or_404(Chapter, id=chapter_id)
        # Check if user has access to this chapter's project
        if not (chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
        chapter_title = chapter.title
        deleted_order = chapter.order
        project = chapter.project
        
        chapter.delete()
        
        # Reorder remaining chapters
        Chapter.objects.filter(
            project=project,
            order__gt=deleted_order
        ).update(order=F('order') - 1)
        
        return JsonResponse({
            'success': True,
            'message': f'Chapter "{chapter_title}" has been deleted.'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Error deleting chapter: {str(e)}'
        }, status=400)


@login_required
@csrf_exempt
def reorder_chapters_api(request):
    """Reorder chapters via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        chapters_data = data.get('chapters', [])
        
        # Update chapter orders
        for item in chapters_data:
            chapter_id = item.get('id')
            new_order = item.get('order')
            if chapter_id and new_order is not None:
                try:
                    chapter = Chapter.objects.get(id=chapter_id)
                    # Check if user has access to this chapter's project
                    if chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists():
                        chapter.order = new_order
                        chapter.save()
                except Chapter.DoesNotExist:
                    pass
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import Document, User
import json

@csrf_exempt
@login_required
def share_document_api(request, document_id):
    """Share a document with other users"""
    if request.method == 'POST':
        try:
            doc = Document.objects.get(id=document_id, author=request.user)
            data = json.loads(request.body)
            user_ids = data.get('user_ids', [])
            users = User.objects.filter(id__in=user_ids)
            doc.shared_with.set(users)
            doc.save()
            return JsonResponse({'success': True, 'shared_with': [u.username for u in users]})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
    return JsonResponse({'error': 'POST required'}, status=405)

@login_required
def get_document_collaborators_api(request, document_id):
    """Get list of collaborators for a document"""
    try:
        doc = Document.objects.get(id=document_id)
        collaborators = doc.shared_with.all()
        return JsonResponse({'collaborators': [u.username for u in collaborators]})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def get_project_collaborators_api(request, project_id):
    """Get list of collaborators for a project"""
    try:
        project = get_object_or_404(Project, id=project_id, author=request.user)
        collaborators = project.collaborators.all()
        collaborators_data = []
        for collaborator in collaborators:
            collaborators_data.append({
                'id': collaborator.id,
                'username': collaborator.username,
                'email': collaborator.email,
            })
        return JsonResponse({'collaborators': collaborators_data})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def users_api(request):
    users = User.objects.all()
    data = [
        {'id': user.id, 'username': user.username}
        for user in users
    ]
    return JsonResponse(data, safe=False)


# New Modern Editor Views
@login_required
def modern_editor(request):
    """Modern Google Docs-like editor"""
    context = {
        'user': request.user,
    }
    
    # Get current chapter or document if specified
    chapter_id = request.GET.get('chapter')
    project_id = request.GET.get('project')
    
    if chapter_id:
        try:
            chapter = Chapter.objects.get(id=chapter_id, project__author=request.user)
            context['chapter'] = chapter
            context['project'] = chapter.project
        except Chapter.DoesNotExist:
            messages.error(request, 'Chapter not found.')
    elif project_id:
        try:
            project = Project.objects.get(id=project_id, author=request.user)
            context['project'] = project
            # Get the latest chapter or create a new one
            latest_chapter = project.chapters.order_by('-created_at').first()
            if latest_chapter:
                context['chapter'] = latest_chapter
        except Project.DoesNotExist:
            messages.error(request, 'Project not found.')
    
    return render(request, 'writer/modern_editor.html', context)


@login_required  
def integrated_editor(request):
    """Integrated editor with all features"""
    context = {
        'user': request.user,
    }
    
    # Get current chapter or document if specified
    chapter_id = request.GET.get('chapter')
    project_id = request.GET.get('project')
    
    if chapter_id:
        try:
            chapter = Chapter.objects.get(id=chapter_id, project__author=request.user)
            context['chapter'] = chapter
            context['project'] = chapter.project
        except Chapter.DoesNotExist:
            messages.error(request, 'Chapter not found.')
    elif project_id:
        try:
            project = Project.objects.get(id=project_id, author=request.user)
            context['project'] = project
            # Get the latest chapter or create a new one
            latest_chapter = project.chapters.order_by('-created_at').first()
            if latest_chapter:
                context['chapter'] = latest_chapter
        except Project.DoesNotExist:
            messages.error(request, 'Project not found.')
    
    return render(request, 'writer/integrated_editor.html', context)


@login_required
def creativity_workshop(request):
    """AI-powered creativity workshop for writers"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creativity_workshop.html', context)


@login_required
def clean_dashboard(request):
    """Clean, writer-focused dashboard"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')[:6]
        
        # Get recent chapters
        recent_chapters = Chapter.objects.filter(
            project__author=request.user
        ).order_by('-updated_at')[:5]
        
        # Calculate stats
        total_words = sum(
            chapter.word_count or 0 
            for project in user_projects 
            for chapter in project.chapters.all()
        )
        
        projects_count = user_projects.count()
        chapters_count = Chapter.objects.filter(project__author=request.user).count()
        
        # Get personal library counts (if models exist)
        try:
            library = PersonalLibrary.objects.get(user=request.user)
            research_notes_count = library.imported_documents.count()
            characters_count = sum(project.characters.count() for project in user_projects)
        except:
            research_notes_count = 0
            characters_count = 0
        
        context.update({
            'recent_projects': user_projects,
            'recent_chapters': recent_chapters,
            'total_words': total_words,
            'projects_count': projects_count,
            'chapters_count': chapters_count,
            'research_notes_count': research_notes_count,
            'characters_count': characters_count,
            'worldbuilding_notes_count': 0,  # Placeholder
            'resources_count': research_notes_count,
        })
        
    except Exception as e:
        logging.error(f"Error loading clean dashboard for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your dashboard.'
    
    return render(request, 'writer/clean_dashboard.html', context)


# API Endpoints for Modern Editor
@csrf_exempt
@login_required
def api_save_document(request):
    """API endpoint to save document content from modern editor"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            title = data.get('title', '')
            content = data.get('content', '')
            chapters_data = data.get('chapters', [])
            
            # For now, save to the user's first project or create a new one
            project = Project.objects.filter(author=request.user).first()
            if not project:
                project = Project.objects.create(
                    title="My Writing Project",
                    description="Created from editor",
                    author=request.user
                )
            
            # Save or update the chapter
            if title and content:
                chapter, created = Chapter.objects.get_or_create(
                    project=project,
                    title=title,
                    defaults={'content': content, 'author': request.user}
                )
                if not created:
                    chapter.content = content
                    chapter.save()
                
                return JsonResponse({
                    'success': True,
                    'chapter_id': chapter.id,
                    'message': 'Document saved successfully'
                })
            
            return JsonResponse({'success': False, 'error': 'Title and content required'})
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})


@csrf_exempt
@login_required  
def api_ai_assistance(request):
    """API endpoint for AI writing assistance"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            action = data.get('action', '')
            text = data.get('text', '')
            
            # Simple mock responses for now (you can integrate real AI later)
            mock_responses = {
                'continue': f"Continuing from your text: '{text[:50]}...', here's a possible continuation: The story unfolds with unexpected twists as the protagonist discovers hidden truths about their world.",
                'improve': f"Here's an improved version: {text.replace('good', 'excellent').replace('bad', 'terrible').replace('nice', 'wonderful')}" if text else "Please select some text to improve.",
                'grammar': f"Grammar-corrected version: {text.replace(' i ', ' I ').replace("dont", "don't").replace("cant", "can't")}" if text else "No grammar issues found.",
                'character': "Character development suggestion: Consider adding internal conflict to make your character more relatable and three-dimensional.",
                'plot': "Plot suggestion: What if the protagonist's greatest strength becomes their weakness at the climax?"
            }
            
            suggestion = mock_responses.get(action, "AI assistance is not available for this action.")
            
            return JsonResponse({
                'success': True,
                'suggestion': suggestion,
                'action': action
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})


@login_required
def compact_library(request):
    """Compact, readable version of the personal library"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Get latest project for continue writing
        latest_project = user_projects.first()
        
        # Get imported documents if available
        try:
            library = PersonalLibrary.objects.get(user=request.user)
            imported_documents = library.imported_documents.all().order_by('-created_at')[:10]
        except PersonalLibrary.DoesNotExist:
            imported_documents = []
        
        context.update({
            'projects': user_projects,
            'latest_project': latest_project,
            'imported_documents': imported_documents,
        })
        
    except Exception as e:
        logging.error(f"Error loading compact library for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your library.'
    
    return render(request, 'writer/compact_library.html', context)


# Ultimate Templates - The Most Beautiful Writing Platform


@login_required 
def ultimate_dashboard(request):
    """The most beautiful, organized, and mobile-friendly dashboard"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects (with visibility control)
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Calculate statistics
        total_words = sum(project.total_word_count for project in user_projects)
        active_projects = user_projects.filter(is_public=True, author=request.user).count()
        
        context.update({
            'user_projects': user_projects,
            'total_words': total_words,
            'active_projects': active_projects,
        })
        
    except Exception as e:
        logging.error(f"Error loading ultimate dashboard for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your dashboard.'
    
    return render(request, 'writer/ultimate_dashboard.html', context)


@login_required
def ultimate_library(request):
    """The most beautifully organized library with perfect text visibility"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects with detailed information
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Add computed fields for library display
        projects_data = []
        for project in user_projects:
            project_data = {
                'id': project.id,
                'title': project.title,
                'description': project.description or 'No description available.',
                'genre': project.genre or 'Unspecified',
                'word_count': project.total_word_count,
                'target_words': project.target_word_count,
                'chapter_count': project.chapter_count,
                'progress': project.progress_percentage,
                'updated_at': project.updated_at,
                'created_at': project.created_at,
                'show_on_dashboard': getattr(project, 'show_on_dashboard', True),
            }
            projects_data.append(project_data)
        
        context.update({
            'user_projects': projects_data,
            'projects': projects_data,
            'total_projects': len(projects_data),
            'total_words': sum(p['word_count'] for p in projects_data),
        })
        
    except Exception as e:
        logging.error(f"Error loading ultimate library for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your library.'
    
    return render(request, 'writer/ultimate_library.html', context)


@login_required
def ultimate_workshop(request):
    """The ultimate AI-powered creativity workshop for writers"""
    context = {
        'user': request.user,
    }
    
    # In a real implementation, this would load writing exercises from the database
    # For now, we'll use the JavaScript-based exercises in the template
    
    return render(request, 'writer/ultimate_workshop.html', context)


@login_required
def ai_playground(request):
    """AI Creative Playground - Where writers play with AI tools"""
    context = {
        'user': request.user,
    }
    
    return render(request, 'writer/ai_playground.html', context)


@login_required
def creativity_workshop(request):
    """Creativity workshop for brainstorming and connecting ideas"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creativity_workshop.html', context)

@login_required
def workshop_sessions_api(request):
    """API endpoint for managing workshop sessions"""
    from django.http import JsonResponse
    import json
    
    if request.method == 'GET':
        # Return user's saved workshop sessions
        # For now, return empty list since we don't have a model yet
        return JsonResponse({
            'sessions': [],
            'status': 'success'
        })
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            # For now, just simulate saving and return success
            # In a real implementation, you'd save to database
            return JsonResponse({
                'status': 'success',
                'message': 'Session saved successfully',
                'session_id': 1  # Mock session ID
            })
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=400)

@login_required
def workshop_session_detail_api(request, session_id):
    """API endpoint for individual workshop session"""
    from django.http import JsonResponse
    
    if request.method == 'GET':
        # Return session details
        # Mock data for now
        mock_session = {
            'id': session_id,
            'title': 'Sample Workshop Session',
            'description': 'A sample session',
            'character_ideas': [],
            'plot_ideas': [],
            'general_notes': '',
            'created_at': '2024-01-01T00:00:00Z'
        }
        
        return JsonResponse({
            'status': 'success',
            'session': mock_session
        })

@login_required
def workshop_history(request):
    """Workshop session history page"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/workshop_history.html', context)


@login_required
def creative_notebook(request):
    """Creative notebook for mind mapping and connecting ideas"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creative_notebook.html', context)


@login_required
def philosophers_corner(request):
    """Philosopher's Corner - Wisdom and inspiration for writers"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/philosophers_corner.html', context)


@login_required
def user_profile(request, username=None):
    """View user profile (public or own profile)"""
    from .models import UserProfile
    from django.shortcuts import get_object_or_404
    
    if username:
        # Viewing someone else's profile
        profile_user = get_object_or_404(User, username=username)
    else:
        # Viewing own profile
        profile_user = request.user
    
    # Get or create profile
    profile = UserProfile.get_or_create_for_user(profile_user)
    
    # Get user's projects for display
    projects = profile_user.owned_projects.filter(is_public=True) if username else profile_user.owned_projects.all()
    
    context = {
        'profile_user': profile_user,
        'profile': profile,
        'projects': projects[:6],  # Show only 6 recent projects
        'is_own_profile': profile_user == request.user,
    }
    return render(request, 'writer/user_profile.html', context)


@login_required
def edit_profile(request):
    """Edit user profile"""
    from .models import UserProfile
    from django.contrib import messages
    from django.shortcuts import redirect
    
    profile = UserProfile.get_or_create_for_user(request.user)
    
    if request.method == 'POST':
        try:
            # Update profile fields
            profile.bio = request.POST.get('bio', '')
            profile.interests = request.POST.get('interests', '')
            profile.favorite_writers = request.POST.get('favorite_writers', '')
            profile.favorite_quotes = request.POST.get('favorite_quotes', '')
            profile.hopes_and_dreams = request.POST.get('hopes_and_dreams', '')
            profile.location = request.POST.get('location', '')
            profile.website = request.POST.get('website', '')
            
            # Handle profile picture upload
            if 'profile_picture' in request.FILES:
                profile.profile_picture = request.FILES['profile_picture']
            
            # Update theme and other preferences
            profile.theme = request.POST.get('theme', profile.theme)
            profile.font_size = request.POST.get('font_size', profile.font_size)
            profile.writing_goal_daily = int(request.POST.get('writing_goal_daily', profile.writing_goal_daily))
            profile.show_statistics = request.POST.get('show_statistics') == 'on'
            profile.show_bookshelf = request.POST.get('show_bookshelf') == 'on'
            
            profile.save()
            messages.success(request, 'Profile updated successfully!')
            
            # If this is an AJAX request, return JSON response
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                from django.http import JsonResponse
                return JsonResponse({
                    'status': 'success',
                    'message': 'Profile updated successfully!',
                    'profile_picture_url': profile.profile_picture.url if profile.profile_picture else None
                })
            
            return redirect('writer:user_profile')
            
        except Exception as e:
            messages.error(request, f'Error updating profile: {str(e)}')
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                from django.http import JsonResponse
                return JsonResponse({
                    'status': 'error',
                    'message': f'Error updating profile: {str(e)}'
                })
    
    context = {
        'profile': profile,
        'user': request.user,
    }
    return render(request, 'writer/edit_profile.html', context)


@login_required
def poetry_workshop(request):
    """Poetry Workshop - Create poems, haikus, and transform text into poetry"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/poetry_workshop.html', context)

def book_formatter(request):
    """Professional Book Formatter - 50+ themes for beautiful book formatting"""
    # Get user's projects for content preview
    user_projects = []
    if request.user.is_authenticated:
        user_projects = Project.objects.filter(author=request.user).order_by('-updated_at')[:5]
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
    }
    return render(request, 'writer/book_formatter.html', context)

def book_formatter_new(request):
    """New Clean Book Formatter - Settings-focused interface without preview"""
    user_projects = []
    if request.user.is_authenticated:
        user_projects = Project.objects.filter(author=request.user).order_by('-updated_at')
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
    }
    return render(request, 'writer/book_formatter_new.html', context)

def cover_designer(request):
    """Professional Cover Designer - Create stunning book covers"""
    # Get user's projects for cover customization
    user_projects = []
    if request.user.is_authenticated:
        user_projects = Project.objects.filter(author=request.user).order_by('-updated_at')[:10]
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
    }
    return render(request, 'writer/cover_designer.html', context)

@csrf_exempt
def export_formatted_book(request):
    """Export book with professional formatting"""
    if request.method == 'POST':
        import json
        from django.http import JsonResponse
        
        try:
            data = json.loads(request.body)
            format_type = data.get('format', 'pdf')  # pdf or epub
            theme = data.get('theme', 'heritage')
            settings = data.get('settings', {})
            content = data.get('content', '')
            
            # In a real implementation, you would:
            # 1. Apply the selected theme and settings
            # 2. Generate PDF using libraries like WeasyPrint or ReportLab
            # 3. Generate EPUB using ebooklib
            # 4. Return download URL
            
            # For now, return success message
            return JsonResponse({
                'success': True,
                'download_url': f'/media/exports/book_{theme}_{format_type}.{format_type}',
                'message': f'Book exported as {format_type.upper()} with {theme} theme'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@csrf_exempt
def export_book_cover(request):
    """Export book cover as high-resolution image"""
    if request.method == 'POST':
        import json
        from django.http import JsonResponse
        
        try:
            data = json.loads(request.body)
            quality = data.get('quality', 'high')  # high or print
            settings = data.get('settings', {})
            
            # In a real implementation, you would:
            # 1. Use PIL or similar to generate high-res cover
            # 2. Apply gradients, fonts, and text positioning
            # 3. Save as PNG/JPG at specified resolution
            # 4. Return download URL
            
            # For now, return success message
            cover_name = settings.get('title', 'book_cover').replace(' ', '_')
            return JsonResponse({
                'success': True,
                'download_url': f'/media/covers/{cover_name}_{quality}.png',
                'message': f'Cover exported as {quality} quality image'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required
def bookshelf_dashboard(request):
    """Bookshelf dashboard with visual book display"""
    user_projects = Project.objects.filter(author=request.user).order_by('-updated_at')
    user_documents = Document.objects.filter(author=request.user).order_by('-updated_at')
    
    # Calculate writing statistics
    total_words = sum(project.total_word_count for project in user_projects)
    total_projects = user_projects.count()
    total_documents = user_documents.count()
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
        'user_documents': user_documents,
        'total_words': total_words,
        'total_projects': total_projects,
        'total_documents': total_documents,
        'show_bookshelf': True,
    }
    return render(request, 'writer/dashboard_with_bookshelf.html', context)


# Legacy Creative Notebook models remain for data migration if needed


# Google Docs-like Editor View
@login_required
def google_docs_editor(request, document_id=None, project_id=None, chapter_id=None, media_id=None):
    """Google Docs-like editor for documents, projects, and media files"""
    document = None
    project = None
    current_chapter = None
    chapters = []
    media_file = None
    
    # Handle project-based editing
    if project_id:
        project = get_object_or_404(Project, id=project_id)
        if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
            from django.contrib import messages
            messages.error(request, 'You do not have permission to edit this project.')
            return redirect('writer:project_list')
        
        # Get chapters for this project
        chapters = Chapter.objects.filter(project=project).order_by('order', 'created_at')
        
        # Get specific chapter or first chapter
        if chapter_id:
            current_chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
        elif chapters.exists():
            current_chapter = chapters.first()
        else:
            # Create a default chapter if none exists
            current_chapter = Chapter.objects.create(
                project=project,
                title="Chapter 1",
                content="",
                order=1
            )
            chapters = Chapter.objects.filter(project=project).order_by('order', 'created_at')
    
    # Handle document-based editing
    elif document_id:
        document = get_object_or_404(Document, id=document_id, author=request.user)
    
    # Handle POST requests for saving
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'save_chapter' or action == 'save':
            title = request.POST.get('title', 'Untitled Document')
            content = request.POST.get('content', '')
            chapter_id_param = request.POST.get('chapter_id')
            
            # Handle project chapter saving
            if project:
                if chapter_id_param:
                    # Update existing chapter
                    chapter = get_object_or_404(Chapter, id=chapter_id_param, project=project)
                    chapter.title = title
                    chapter.content = content
                    chapter.save()
                    return JsonResponse({
                        'status': 'success',
                        'message': 'Chapter saved successfully',
                        'chapter_id': chapter.id
                    })
                elif current_chapter:
                    # Update current chapter
                    current_chapter.title = title
                    current_chapter.content = content
                    current_chapter.save()
                    return JsonResponse({
                        'status': 'success',
                        'message': 'Chapter saved successfully',
                        'chapter_id': current_chapter.id
                    })
                else:
                    # Create new chapter for project
                    chapter = Chapter.objects.create(
                        project=project,
                        title=title,
                        content=content,
                        order=chapters.count() + 1
                    )
                    return JsonResponse({
                        'status': 'success',
                        'message': 'Chapter created successfully',
                        'chapter_id': chapter.id
                    })
            
            # Handle document saving
            elif document:
                # Check if this is a media file document
                if hasattr(document, 'is_media_file') and document.is_media_file:
                    # Update the media file's extracted text
                    from .models import MediaFile
                    media_file = get_object_or_404(MediaFile, id=document.media_file_id, user=request.user)
                    media_file.title = title
                    media_file.extracted_text = content
                    media_file.save()
                    return JsonResponse({
                        'status': 'success',
                        'message': 'Media file content saved successfully',
                        'document_id': document.id
                    })
                else:
                    # Regular document saving
                    document.title = title
                    document.content = content
                    document.save()
                    return JsonResponse({
                        'status': 'success',
                        'message': 'Document saved successfully',
                        'document_id': document.id
                    })
            else:
                # Create a new document
                document = Document.objects.create(
                    title=title,
                    content=content,
                    author=request.user
                )
                return JsonResponse({
                    'status': 'success',
                    'message': 'Document created successfully',
                    'document_id': document.id
                })
        
        elif action == 'get_document':
            if document:
                return JsonResponse({
                    'status': 'success',
                    'title': document.title,
                    'content': document.content,
                    'document_id': document.id
                })
            else:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No document found'
                })
        
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid action'
        })
    
    # Handle media file editing
    if media_id:
        from .models import MediaFile
        media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
        
        # Create a temporary document-like object from the media file
        class MediaDocument:
            def __init__(self, media_file):
                self.id = f"media_{media_file.id}"
                self.title = media_file.title
                self.content = media_file.extracted_text or ""
                self.author = media_file.user
                self.created_at = media_file.upload_date
                self.updated_at = media_file.updated_at
                self.is_media_file = True
                self.media_file_id = media_file.id
                self.word_count = len(self.content.split()) if self.content else 0
        
        document = MediaDocument(media_file)
    
    # Build context for template
    import time
    context = {
        'user': request.user,
        'document': document if document else current_chapter,
        'project': project,
        'current_chapter': current_chapter,
        'all_chapters': chapters if project else [],
        'cache_bust': str(int(time.time())),  # Cache busting timestamp
        'is_media_file': media_id is not None,
        'media_file': media_file,
        'chapters': [
            {
                'id': ch.id,
                'title': ch.title or f'Chapter {ch.order}',
                'order': ch.order,
                'word_count': ch.word_count,
                'updated_at': ch.updated_at.isoformat() if ch.updated_at else None
            }
            for ch in chapters
        ] if project else []
    }
    return render(request, 'writer/google_docs_editor.html', context)


@login_required
def academic_editor(request):
    """Academic writing editor"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/academic_editor.html', context)


@login_required
def book_reader(request, project_id):
    """Book reader view - displays project content as a formatted ebook"""
    project = get_object_or_404(Project, pk=project_id)
    
    # Check if user can access this project
    if project.author != request.user and not project.collaborators.filter(id=request.user.id).exists():
        return render(request, '403.html')
    
    chapters = project.chapters.all().order_by('order')
    
    context = {
        'user': request.user,
        'project': project,
        'chapters': chapters,
    }
    return render(request, 'writer/book_reader.html', context)


@login_required
def user_preferences(request):
    """User preferences page"""
    profile = UserProfile.get_or_create_for_user(request.user)
    context = {
        'user': request.user,
        'profile': profile,
    }
    return render(request, 'writer/user_preferences.html', context)


@csrf_exempt
@login_required
def api_update_theme(request):
    """Update user theme preference"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            theme = data.get('theme', 'ethereal')
            
            # Get or create user profile
            profile = UserProfile.get_or_create_for_user(request.user)
            profile.theme = theme
            profile.save()
            
            return JsonResponse({'success': True, 'theme': theme})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required
def my_library(request):
    """My Library - Personal collection of saved books"""
    user_projects = Project.objects.filter(author=request.user).order_by('-created_at')
    user_documents = Document.objects.filter(author=request.user).order_by('-created_at')
    
    # Separate projects by bookshelf visibility
    bookshelf_projects = user_projects.filter(show_on_dashboard=True)
    library_projects = user_projects.filter(show_on_dashboard=False)
    
    # Calculate statistics
    total_books = user_projects.count() + user_documents.count()
    total_words = sum(p.total_word_count or 0 for p in user_projects) + sum(d.word_count or 0 for d in user_documents)
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
        'bookshelf_projects': bookshelf_projects,
        'library_projects': library_projects,
        'user_documents': user_documents,
        'total_books': total_books,
        'total_words': total_words,
    }
    
    return render(request, 'writer/my_library.html', context)


@csrf_exempt
@login_required
def toggle_bookshelf_visibility(request):
    """Toggle project visibility on dashboard bookshelf"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_id = data.get('project_id')
            
            project = Project.objects.get(id=project_id, author=request.user)
            project.show_on_dashboard = not project.show_on_dashboard
            project.save()
            
            return JsonResponse({
                'success': True,
                'show_on_dashboard': project.show_on_dashboard,
                'message': f"Project {'added to' if project.show_on_dashboard else 'removed from'} bookshelf"
            })
            
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})

@csrf_exempt
@login_required
def detect_chapters_from_text(request):
    """
    Django backend function to detect chapters from document text using regex patterns
    """
    if request.method == 'POST':
        try:
            import json
            import re
            
            data = json.loads(request.body)
            text = data.get('text', '')
            
            if not text or len(text.strip()) < 100:
                return JsonResponse({
                    'success': False,
                    'error': 'Text too short for chapter detection'
                })
            
            # Django backend chapter detection function
            def detect_chapters(document_text):
                chapters = []
                
                # Define regex patterns for chapter detection
                patterns = [
                    # Pattern 1: "Chapter X" or "CHAPTER X" with optional title
                    {
                        'regex': r'(?:^|\n)\s*(CHAPTER|Chapter|chapter)\s+(\d+|[IVXLCDM]+|One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten)(?:[:\-\s]+([^\n]+))?\s*(?=\n)',
                        'weight': 10,
                        'type': 'chapter'
                    },
                    # Pattern 2: Markdown headers (# Chapter Title)
                    {
                        'regex': r'(?:^|\n)(#{1,3})\s+([^\n]+)\s*(?=\n)',
                        'weight': 8,
                        'type': 'markdown'
                    },
                    # Pattern 3: HTML headings
                    {
                        'regex': r'<(h[1-3])[^>]*>([^<]+)</\1>',
                        'weight': 9,
                        'type': 'html'
                    },
                    # Pattern 4: Numbered sections (1. Title, 2. Title)
                    {
                        'regex': r'(?:^|\n)\s*(\d+)[\.):]\s+([A-Z][^\n]{5,80})\s*(?=\n)',
                        'weight': 6,
                        'type': 'numbered'
                    },
                    # Pattern 5: All caps titles (standalone lines)
                    {
                        'regex': r'(?:^|\n)\s*([A-Z][A-Z\s]{8,50})\s*(?=\n)',
                        'weight': 5,
                        'type': 'caps'
                    },
                    # Pattern 6: Part/Section markers
                    {
                        'regex': r'(?:^|\n)\s*(PART|Part|part|SECTION|Section|section)\s+(\d+|[IVXLCDM]+|One|Two|Three|Four|Five)(?:[:\-\s]+([^\n]+))?\s*(?=\n)',
                        'weight': 7,
                        'type': 'section'
                    }
                ]
                
                matches = []
                
                # Find all matches
                for pattern_info in patterns:
                    for match in re.finditer(pattern_info['regex'], document_text, re.MULTILINE | re.IGNORECASE):
                        start_pos = match.start()
                        end_pos = match.end()
                        
                        # Extract title based on pattern type
                        title = extract_title_from_match(match, pattern_info['type'])
                        
                        matches.append({
                            'start': start_pos,
                            'end': end_pos,
                            'title': title,
                            'weight': pattern_info['weight'],
                            'type': pattern_info['type'],
                            'text': match.group(0).strip()
                        })
                
                # Sort matches by position
                matches.sort(key=lambda x: x['start'])
                
                # Filter overlapping matches (keep highest weight)
                filtered_matches = []
                for match in matches:
                    # Check for overlaps
                    overlapping = [m for m in filtered_matches if abs(m['start'] - match['start']) < 100]
                    
                    if not overlapping:
                        filtered_matches.append(match)
                    elif match['weight'] > max(overlapping, key=lambda x: x['weight'])['weight']:
                        # Remove overlapping matches with lower weight
                        for overlap in overlapping:
                            filtered_matches.remove(overlap)
                        filtered_matches.append(match)
                
                # Create chapter objects
                for i, match in enumerate(filtered_matches):
                    start_pos = match['start']
                    next_match = filtered_matches[i + 1] if i + 1 < len(filtered_matches) else None
                    end_pos = next_match['start'] if next_match else len(document_text)
                    
                    # Extract chapter content
                    chapter_content = document_text[start_pos:end_pos].strip()
                    
                    # Remove the title from content if it appears at the beginning
                    title_text = match['text']
                    if chapter_content.startswith(title_text):
                        chapter_content = chapter_content[len(title_text):].strip()
                    
                    word_count = len(chapter_content.split()) if chapter_content else 0
                    
                    # Skip very short chapters
                    if word_count < 20:
                        continue
                    
                    chapters.append({
                        'id': f'chapter_{i + 1}',
                        'title': match['title'],
                        'content': chapter_content,
                        'word_count': word_count,
                        'start_position': start_pos,
                        'end_position': end_pos,
                        'confidence': calculate_confidence_score(match, chapter_content),
                        'type': match['type'],
                        'order': i + 1
                    })
                
                return chapters
            
            def extract_title_from_match(match, pattern_type):
                """Extract clean title from regex match based on pattern type"""
                if pattern_type == 'chapter':
                    number = match.group(2)
                    subtitle = match.group(3) if len(match.groups()) > 2 and match.group(3) else None
                    if subtitle:
                        return f"Chapter {number}: {subtitle.strip()}"
                    return f"Chapter {number}"
                    
                elif pattern_type == 'markdown':
                    return match.group(2).strip()
                    
                elif pattern_type == 'html':
                    return match.group(2).strip()
                    
                elif pattern_type == 'numbered':
                    number = match.group(1)
                    title = match.group(2).strip()
                    return title
                    
                elif pattern_type == 'caps':
                    title = match.group(1).strip()
                    # Convert to title case
                    return ' '.join(word.capitalize() for word in title.split())
                    
                elif pattern_type == 'section':
                    section_type = match.group(1).title()
                    number = match.group(2)
                    subtitle = match.group(3) if len(match.groups()) > 2 and match.group(3) else None
                    if subtitle:
                        return f"{section_type} {number}: {subtitle.strip()}"
                    return f"{section_type} {number}"
                    
                return "Untitled Chapter"
            
            def calculate_confidence_score(match, content):
                """Calculate confidence score based on various factors"""
                base_score = {
                    'chapter': 0.95,
                    'html': 0.9,
                    'markdown': 0.85,
                    'section': 0.8,
                    'numbered': 0.7,
                    'caps': 0.6
                }.get(match['type'], 0.5)
                
                # Adjust based on content length
                word_count = len(content.split())
                if word_count > 500:
                    base_score += 0.1
                elif word_count < 100:
                    base_score -= 0.2
                
                # Adjust based on title quality
                title_length = len(match['title'])
                if 5 <= title_length <= 50:
                    base_score += 0.05
                elif title_length < 5:
                    base_score -= 0.1
                
                return max(0.1, min(1.0, base_score))
            
            # Detect chapters
            detected_chapters = detect_chapters(text)
            
            return JsonResponse({
                'success': True,
                'chapters': detected_chapters,
                'chapter_count': len(detected_chapters),
                'total_words': sum(ch['word_count'] for ch in detected_chapters),
                'average_confidence': sum(ch['confidence'] for ch in detected_chapters) / len(detected_chapters) if detected_chapters else 0
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Error detecting chapters: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'})

@csrf_exempt
@login_required
def create_chapters_from_document(request):
    """
    Create chapters automatically from an uploaded document
    """
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('file')
            project_id = request.POST.get('project_id')
            
            if not uploaded_file:
                return JsonResponse({
                    'success': False,
                    'error': 'No file provided'
                }, status=400)
            
            if not project_id:
                return JsonResponse({
                    'success': False,
                    'error': 'No project ID provided'
                }, status=400)
            
            # Verify project ownership
            try:
                project = Project.objects.get(id=project_id, author=request.user)
            except Project.DoesNotExist:
                return JsonResponse({
                    'success': False,
                    'error': 'Project not found or access denied'
                }, status=404)
            
            # Create temporary file
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            try:
                # Extract text and identify chapters
                from .document_parser import extract_text_and_chapters_from_file
                result = extract_text_and_chapters_from_file(temp_file_path)
                
                if not result['chapters']:
                    return JsonResponse({
                        'success': False,
                        'error': 'No chapters could be identified in this document'
                    })
                
                # Create chapters in the project
                created_chapters = []
                
                for chapter_data in result['chapters']:
                    chapter = Chapter.objects.create(
                        title=chapter_data['title'],
                        content=chapter_data['content'],
                        project=project,
                        order=chapter_data['order'],
                        author=request.user
                    )
                    
                    created_chapters.append({
                        'id': chapter.id,
                        'title': chapter.title,
                        'order': chapter.order,
                        'word_count': chapter_data['word_count'],
                        'confidence': chapter_data['confidence'],
                        'detected_method': chapter_data['detected_method']
                    })
                
                return JsonResponse({
                    'success': True,
                    'message': f'Successfully created {len(created_chapters)} chapters',
                    'chapters': created_chapters,
                    'total_words': result['total_words'],
                    'average_confidence': result['average_confidence'],
                    'project_id': project.id,
                    'project_title': project.title
                })
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'})

@csrf_exempt
@login_required  
def analyze_document_chapters(request):
    """
    Analyze an uploaded document for chapter structure without creating chapters
    """
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('file')
            
            if not uploaded_file:
                return JsonResponse({
                    'success': False,
                    'error': 'No file provided'
                }, status=400)
            
            # Create temporary file
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            try:
                # Extract text and identify chapters
                from .document_parser import extract_text_and_chapters_from_file
                result = extract_text_and_chapters_from_file(temp_file_path)
                
                # Prepare response with chapter preview
                chapter_preview = []
                for chapter in result['chapters']:
                    preview_content = chapter['content'][:200] + '...' if len(chapter['content']) > 200 else chapter['content']
                    chapter_preview.append({
                        'title': chapter['title'],
                        'word_count': chapter['word_count'],
                        'confidence': chapter['confidence'],
                        'detected_method': chapter['detected_method'],
                        'preview': preview_content.strip()
                    })
                
                return JsonResponse({
                    'success': True,
                    'filename': uploaded_file.name,
                    'total_words': result['total_words'],
                    'chapter_count': result['chapter_count'],
                    'average_confidence': result['average_confidence'],
                    'chapters': chapter_preview,
                    'has_multiple_chapters': result['chapter_count'] > 1
                })
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Error analyzing document: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'})

@csrf_exempt
@login_required
def api_upload_document(request):
    """API endpoint for uploading documents from project detail page"""
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('document')
            project_id = request.POST.get('project_id')
            
            if not uploaded_file:
                return JsonResponse({
                    'success': False, 
                    'error': 'No document provided'
                }, status=400)
            
            # Create ImportedDocument instance
            from .models import ImportedDocument
            imported_doc = ImportedDocument(
                user=request.user,
                original_file=uploaded_file,
                file_size=uploaded_file.size,
                title=uploaded_file.name.rsplit('.', 1)[0]  # Remove file extension
            )
            
            # Determine file type
            file_name = uploaded_file.name.lower()
            if file_name.endswith('.pdf'):
                imported_doc.import_type = 'pdf'
            elif file_name.endswith(('.docx', '.doc')):
                imported_doc.import_type = 'docx'
            elif file_name.endswith('.txt'):
                imported_doc.import_type = 'txt'
            elif file_name.endswith('.epub'):
                imported_doc.import_type = 'epub'
            else:
                imported_doc.import_type = 'other'
            
            # Save the document
            imported_doc.save()
            
            # Extract content
            try:
                from .utils import extract_text_from_file
                extracted_content = extract_text_from_file(imported_doc.original_file)
                imported_doc.extracted_content = extracted_content
                imported_doc.save()
                
                # If project_id provided, optionally associate with project
                if project_id:
                    imported_doc.notes = f"Uploaded from project {project_id}"
                    imported_doc.save()
                
                return JsonResponse({
                    'success': True,
                    'message': f'Document "{uploaded_file.name}" uploaded successfully!',
                    'import_id': imported_doc.id,
                    'redirect_url': f'/writer/import/{imported_doc.id}/'
                })
                
            except Exception as e:
                # Document saved but content extraction failed
                return JsonResponse({
                    'success': True,
                    'message': f'Document "{uploaded_file.name}" uploaded successfully, but content extraction failed.',
                    'import_id': imported_doc.id,
                    'redirect_url': f'/writer/import/{imported_doc.id}/',
                    'warning': str(e)
                })
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Upload failed: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'}, status=405)

@csrf_exempt
def api_register(request):
    """API endpoint for user registration"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')
            
            if not username or not email or not password:
                return JsonResponse({
                    'success': False, 
                    'error': 'Username, email, and password are required'
                }, status=400)
            
            # Check if username already exists
            if User.objects.filter(username=username).exists():
                return JsonResponse({
                    'success': False, 
                    'error': 'Username already exists'
                }, status=400)
                
            # Check if email already exists
            if User.objects.filter(email=email).exists():
                return JsonResponse({
                    'success': False, 
                    'error': 'Email already exists'
                }, status=400)
            
            # Create user
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password
            )
            
            # Create auth token
            from rest_framework.authtoken.models import Token
            token, created = Token.objects.get_or_create(user=user)
            
            return JsonResponse({
                'success': True,
                'message': 'Account created successfully',
                'token': token.key,
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': user.email
                }
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'error': str(e)
            }, status=500)
    
@login_required
def upload_to_editor(request):
    """Handle file upload to the document editor"""
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('file')
            if not uploaded_file:
                return JsonResponse({'success': False, 'message': 'No file provided'}, status=400)
            
            # Get file extension
            file_name = uploaded_file.name
            file_ext = file_name.split('.')[-1].lower()
            
            # Initialize content and title
            content = ""
            title = file_name.rsplit('.', 1)[0]  # Remove extension for title
            
            # Handle different file types
            if file_ext in ['txt', 'text']:
                # Plain text file
                content = uploaded_file.read().decode('utf-8', errors='ignore')
                
            elif file_ext in ['doc', 'docx']:
                # Microsoft Word documents
                try:
                    import docx
                    from io import BytesIO
                    doc = docx.Document(BytesIO(uploaded_file.read()))
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    content = '\n\n'.join(paragraphs)
                except ImportError:
                    return JsonResponse({
                        'success': False, 
                        'message': 'Word document support not available. Please install python-docx.'
                    }, status=500)
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Error reading Word document: {str(e)}'
                    }, status=500)
                    
            elif file_ext == 'pdf':
                # PDF files
                try:
                    import PyPDF2
                    from io import BytesIO
                    pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
                    pages_text = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            pages_text.append(text)
                    content = '\n\n'.join(pages_text)
                except ImportError:
                    return JsonResponse({
                        'success': False, 
                        'message': 'PDF support not available. Please install PyPDF2.'
                    }, status=500)
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Error reading PDF: {str(e)}'
                    }, status=500)
                    
            elif file_ext == 'rtf':
                # RTF files
                try:
                    from striprtf.striprtf import rtf_to_text
                    content = rtf_to_text(uploaded_file.read().decode('utf-8', errors='ignore'))
                except ImportError:
                    # Fallback: just read as text and try to clean it up
                    raw_content = uploaded_file.read().decode('utf-8', errors='ignore')
                    # Basic RTF stripping (very basic)
                    import re
                    content = re.sub(r'\\[a-z]+\d?\s?', '', raw_content)
                    content = re.sub(r'[{}]', '', content)
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Error reading RTF: {str(e)}'
                    }, status=500)
                    
            elif file_ext == 'odt':
                # OpenDocument Text files
                try:
                    from odf import text, teletype
                    from odf.opendocument import load
                    from io import BytesIO
                    doc = load(BytesIO(uploaded_file.read()))
                    paragraphs = doc.getElementsByType(text.P)
                    content = '\n\n'.join([teletype.extractText(p) for p in paragraphs])
                except ImportError:
                    return JsonResponse({
                        'success': False, 
                        'message': 'ODT support not available. Please install odfpy.'
                    }, status=500)
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Error reading ODT: {str(e)}'
                    }, status=500)
            else:
                return JsonResponse({
                    'success': False, 
                    'message': f'Unsupported file type: {file_ext}'
                }, status=400)
            
            # Clean up the content
            content = content.strip()
            
            # Return the processed content
            return JsonResponse({
                'success': True,
                'title': title,
                'content': content,
                'message': 'File uploaded successfully'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'message': f'Upload error: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'message': 'POST request required'}, status=405)


# ============ MEDIA LIBRARY VIEWS ============

@login_required
def media_library(request):
    """Media library view for multiple file types"""
    from .models import MediaFile
    
    # Get user's media files
    media_files = MediaFile.objects.filter(user=request.user).order_by('-upload_date')
    
    # Filter by type if requested
    file_type = request.GET.get('type')
    if file_type and file_type in ['image', 'video', 'audio']:
        media_files = media_files.filter(file_type=file_type)
    
    # Search functionality
    search = request.GET.get('search')
    if search:
        media_files = media_files.filter(title__icontains=search)
    
    # Get statistics (only for allowed file types)
    stats = {
        'total_files': MediaFile.objects.filter(user=request.user, file_type__in=['image', 'video', 'audio']).count(),
        'images': MediaFile.objects.filter(user=request.user, file_type='image').count(),
        'videos': MediaFile.objects.filter(user=request.user, file_type='video').count(),
        'audio': MediaFile.objects.filter(user=request.user, file_type='audio').count(),
        'total_size': sum(MediaFile.objects.filter(user=request.user, file_type__in=['image', 'video', 'audio']).values_list('file_size', flat=True) or [0])
    }
    
    # Convert total size to human readable
    size = stats['total_size']
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            stats['total_size_human'] = f"{size:.1f} {unit}"
            break
        size /= 1024
    else:
        stats['total_size_human'] = f"{size:.1f} TB"
    
    context = {
        'media_files': media_files,
        'stats': stats,
        'total_files': stats['total_files'],
        'image_count': stats['images'],
        'video_count': stats['videos'],
        'audio_count': stats['audio'],
        'current_filter': file_type,
        'search_query': search,
    }
    
    return render(request, 'writer/media_library.html', context)


@csrf_exempt
@login_required
def media_upload(request):
    """Handle media file upload"""
    if request.method == 'POST':
        from .models import MediaFile
        import json
        
        try:
            # Handle both 'files' (multiple) and 'file' (single) upload
            uploaded_files = request.FILES.getlist('files') or [request.FILES.get('file')]
            uploaded_files = [f for f in uploaded_files if f is not None]  # Remove None values
            project_id = request.POST.get('project')
            
            if not uploaded_files:
                return JsonResponse({'success': False, 'error': 'No files provided'}, status=400)
            
            uploaded_file_info = []
            project = None
            
            if project_id:
                try:
                    project = Project.objects.get(id=project_id, author=request.user)
                except Project.DoesNotExist:
                    project = None
            
            for uploaded_file in uploaded_files:
                try:
                    # Get file extension
                    file_extension = uploaded_file.name.split('.')[-1].lower() if '.' in uploaded_file.name else ''
                    
                    # Define allowed file types (only images, audio, video)
                    allowed_extensions = {
                        # Image formats
                        'jpg', 'jpeg', 'png', 'gif', 'webp', 'svg', 'bmp', 'tiff', 'ico',
                        # Video formats  
                        'mp4', 'avi', 'mov', 'wmv', 'flv', 'webm', 'mkv', 'm4v', '3gp',
                        # Audio formats
                        'mp3', 'wav', 'flac', 'aac', 'ogg', 'm4a', 'wma', 'opus'
                    }
                    
                    # Check if file type is allowed
                    if file_extension not in allowed_extensions:
                        return JsonResponse({
                            'success': False, 
                            'error': f'File type "{file_extension}" not allowed. Only images, audio, and video files are supported.'
                        }, status=400)
                    
                    # Create media file
                    media_file = MediaFile.objects.create(
                        title=uploaded_file.name.rsplit('.', 1)[0],
                        user=request.user,
                        project=project,
                        file=uploaded_file,
                        description=f"Uploaded from media library"
                    )
                    
                    # No text extraction needed for media files (images, audio, video)
                    
                    uploaded_file_info.append({
                        'id': media_file.id,
                        'title': media_file.title,
                        'file_type': media_file.file_type,
                        'file_format': media_file.file_format,
                        'file_size': media_file.file_size_human,
                        'thumbnail_url': media_file.thumbnail_url,
                        'upload_date': media_file.upload_date.strftime('%Y-%m-%d %H:%M')
                    })
                    
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'error': f'Failed to upload {uploaded_file.name}: {str(e)}'
                    }, status=500)
            
            # For single file upload, return 'file' instead of 'files' array for JS compatibility
            response_data = {
                'success': True,
                'message': f'{len(uploaded_file_info)} file(s) uploaded successfully',
                'files': uploaded_file_info
            }
            
            # Add 'file' key for single file uploads (JS expects this)
            if len(uploaded_file_info) == 1:
                response_data['file'] = uploaded_file_info[0]
            
            return JsonResponse(response_data)
            
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'error': f'Upload failed: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'}, status=405)


def extract_document_content(uploaded_file, file_format):
    """Extract text content from various document formats"""
    content = ""
    
    try:
        # Save current position and reset to beginning
        original_position = uploaded_file.tell()
        uploaded_file.seek(0)
        
        if file_format in ['txt', 'text']:
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            
        elif file_format in ['doc', 'docx']:
            try:
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(uploaded_file.read()))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                content = '\n\n'.join(paragraphs)
            except ImportError:
                content = "Word document content extraction requires python-docx"
                
        elif file_format == 'pdf':
            try:
                import PyPDF2
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
                pages_text = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                content = '\n\n'.join(pages_text)
            except ImportError:
                content = "PDF content extraction requires PyPDF2"
                
        elif file_format == 'rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                content = rtf_to_text(uploaded_file.read().decode('utf-8', errors='ignore'))
            except ImportError:
                raw_content = uploaded_file.read().decode('utf-8', errors='ignore')
                import re
                content = re.sub(r'\\[a-z]+\d?\s?', '', raw_content)
                content = re.sub(r'[{}]', '', content)
                
        elif file_format == 'odt':
            try:
                from odf import text, teletype
                from odf.opendocument import load
                from io import BytesIO
                doc = load(BytesIO(uploaded_file.read()))
                paragraphs = doc.getElementsByType(text.P)
                content = '\n\n'.join([teletype.extractText(p) for p in paragraphs])
            except ImportError:
                content = "ODT content extraction requires odfpy"
                
        elif file_format in ['html', 'htm']:
            try:
                from bs4 import BeautifulSoup
                html_content = uploaded_file.read().decode('utf-8', errors='ignore')
                soup = BeautifulSoup(html_content, 'html.parser')
                content = soup.get_text(separator='\n\n')
            except ImportError:
                content = "HTML content extraction requires beautifulsoup4"
                
    except Exception as e:
        content = f"Content extraction failed: {str(e)}"
    finally:
        # Restore original file position
        try:
            uploaded_file.seek(original_position)
        except:
            uploaded_file.seek(0)
    
    return content.strip()


@login_required
def media_detail(request, media_id):
    """View media file details"""
    from .models import MediaFile
    
    media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
    
    context = {
        'media_file': media_file,
    }
    
    return render(request, 'writer/media_detail.html', context)


@login_required
def media_delete(request, media_id):
    """Delete media file"""
    from .models import MediaFile
    
    if request.method == 'POST':
        media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
        
        try:
            # Delete the file from storage
            if media_file.file:
                media_file.file.delete()
            
            # Delete the database record
            media_file.delete()
            
            return JsonResponse({
                'success': True,
                'message': 'File deleted successfully'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Delete failed: {str(e)}'
            }, status=500)
    
    return JsonResponse({'success': False, 'error': 'POST method required'}, status=405)


@login_required
def media_view(request, media_id):
    """View media file in browser"""
    from .models import MediaFile
    from django.http import HttpResponse, Http404
    import os
    import mimetypes
    
    media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
    
    try:
        if not media_file.file or not os.path.exists(media_file.file.path):
            raise Http404("File not found")
        
        # Get the MIME type
        content_type, _ = mimetypes.guess_type(media_file.file.path)
        if not content_type:
            # Default content types based on file type
            if media_file.file_type == 'image':
                content_type = 'image/jpeg'
            elif media_file.file_type == 'video':
                content_type = 'video/mp4'
            elif media_file.file_type == 'audio':
                content_type = 'audio/mpeg'
            elif media_file.file_format == 'pdf':
                content_type = 'application/pdf'
            else:
                content_type = 'application/octet-stream'
        
        with open(media_file.file.path, 'rb') as f:
            response = HttpResponse(f.read(), content_type=content_type)
            # For inline viewing instead of download
            response['Content-Disposition'] = f'inline; filename="{media_file.title}.{media_file.file_format}"'
            return response
            
    except Exception as e:
        raise Http404(f"File view failed: {str(e)}")


@login_required
def media_download(request, media_id):
    """Download media file"""
    from .models import MediaFile
    from django.http import HttpResponse, Http404
    import os
    
    media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
    
    try:
        if not media_file.file or not os.path.exists(media_file.file.path):
            raise Http404("File not found")
        
        with open(media_file.file.path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename="{media_file.title}.{media_file.file_format}"'
            return response
            
    except Exception as e:
        raise Http404(f"File download failed: {str(e)}")


@login_required
def media_update(request, media_id):
    """Update media file details - handles both GET (show form) and POST (save changes)"""
    from .models import MediaFile
    import json
    
    media_file = get_object_or_404(MediaFile, id=media_id, user=request.user)
    
    if request.method == 'POST':
        # Handle both JSON API requests and form submissions
        if request.content_type == 'application/json':
            try:
                data = json.loads(request.body)
                
                # Update fields from JSON
                if 'title' in data:
                    media_file.title = data['title']
                if 'description' in data:
                    media_file.description = data['description']
                if 'tags' in data:
                    media_file.tags = data['tags']
                if 'project' in data:
                    project_id = data['project']
                    if project_id:
                        try:
                            project = Project.objects.get(id=project_id, author=request.user)
                            media_file.project = project
                        except Project.DoesNotExist:
                            pass
                    else:
                        media_file.project = None
                
                media_file.save()
                
                return JsonResponse({
                    'success': True,
                    'message': 'File updated successfully'
                })
                
            except Exception as e:
                return JsonResponse({
                    'success': False,
                    'error': f'Update failed: {str(e)}'
                }, status=500)
        else:
            # Handle regular form submission
            try:
                media_file.title = request.POST.get('title', media_file.title)
                media_file.description = request.POST.get('description', media_file.description)
                
                # Handle tags (convert comma-separated string to list)
                tags_string = request.POST.get('tags', '')
                if tags_string:
                    media_file.tags = [tag.strip() for tag in tags_string.split(',') if tag.strip()]
                else:
                    media_file.tags = []
                
                # Handle project assignment
                project_id = request.POST.get('project')
                if project_id:
                    try:
                        project = Project.objects.get(id=project_id, author=request.user)
                        media_file.project = project
                    except Project.DoesNotExist:
                        pass
                else:
                    media_file.project = None
                
                media_file.save()
                messages.success(request, 'Media file updated successfully!')
                return redirect('writer:media_detail', media_id=media_file.id)
                
            except Exception as e:
                messages.error(request, f'Failed to update media file: {str(e)}')
    
    # GET request - show edit form
    user_projects = Project.objects.filter(author=request.user).order_by('title')
    
    context = {
        'media_file': media_file,
        'projects': user_projects,
        'tags_string': ', '.join(media_file.tags) if media_file.tags else '',
    }
    return render(request, 'writer/media_edit.html', context)


# User Registration Views
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
from django.shortcuts import render, redirect
from django.contrib import messages
from django.urls import reverse
from django import forms
from django.contrib.auth.models import User

class CustomUserCreationForm(UserCreationForm):
    """Custom user creation form with email field"""
    email = forms.EmailField(required=True, help_text="We'll never share your email with anyone else.")
    
    class Meta:
        model = User
        fields = ("username", "email", "password1", "password2")
    
    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data["email"]
        if commit:
            user.save()
        return user

def signup(request):
    """User registration view"""
    if request.user.is_authenticated:
        return redirect('writer:dashboard')
    
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Log the user in immediately after registration
            login(request, user)
            messages.success(request, f'Welcome to Writer\'s Web Dream, {user.username}! Your account has been created successfully.')
            return redirect('writer:dashboard')
    else:
        form = CustomUserCreationForm()
    
    return render(request, 'registration/signup.html', {'form': form})


class MLStripper(HTMLParser):
    """HTML tag stripper for converting HTML content to plain text"""
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []
    
    def handle_data(self, d):
        self.text.append(d)
    
    def get_data(self):
        return ''.join(self.text)


def strip_html_tags(html):
    """Remove HTML tags from content"""
    if not html:
        return ""
    s = MLStripper()
    s.feed(html)
    return s.get_data()


@login_required
def export_project_pdf(request, pk):
    """Export entire project as PDF"""
    project = get_object_or_404(Project, pk=pk)
    
    # Check permissions
    if project.author != request.user and request.user not in project.collaborators.all():
        messages.error(request, "You don't have permission to export this project.")
        return redirect('writer:project_detail', pk=pk)
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='BookTitle',
                             parent=styles['Heading1'],
                             fontSize=24,
                             alignment=TA_CENTER,
                             spaceAfter=30))
    styles.add(ParagraphStyle(name='ChapterTitle',
                             parent=styles['Heading2'],
                             fontSize=18,
                             spaceAfter=12))
    styles.add(ParagraphStyle(name='ChapterContent',
                             parent=styles['Normal'],
                             fontSize=12,
                             alignment=TA_JUSTIFY,
                             spaceAfter=12))
    
    # Add title page
    elements.append(Paragraph(project.title, styles['BookTitle']))
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph(f"by {project.author.get_full_name() or project.author.username}", styles['Title']))
    elements.append(Spacer(1, 0.5*inch))
    
    if project.description:
        elements.append(Paragraph(strip_html_tags(project.description), styles['Normal']))
    
    elements.append(PageBreak())
    
    # Add chapters
    chapters = project.chapters.all().order_by('order')
    has_content = False
    
    if not chapters.exists():
        # No chapters - add placeholder
        elements.append(Paragraph("This project doesn't have any chapters yet.", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph("Add chapters to your project to see content in the export.", styles['Normal']))
    else:
        for chapter in chapters:
            # Chapter title
            elements.append(Paragraph(chapter.title, styles['ChapterTitle']))
            elements.append(Spacer(1, 0.2*inch))
            
            # Chapter content
            if chapter.content:
                # Convert HTML content to plain text
                plain_content = strip_html_tags(chapter.content)
                # Split into paragraphs
                paragraphs = plain_content.split('\n')
                chapter_has_content = False
                for para in paragraphs:
                    if para.strip():
                        elements.append(Paragraph(para, styles['ChapterContent']))
                        elements.append(Spacer(1, 0.1*inch))
                        has_content = True
                        chapter_has_content = True
                
                if not chapter_has_content:
                    elements.append(Paragraph("(This chapter is empty)", styles['Italic']))
                    elements.append(Spacer(1, 0.2*inch))
            else:
                elements.append(Paragraph("(This chapter is empty)", styles['Italic']))
                elements.append(Spacer(1, 0.2*inch))
            
            elements.append(PageBreak())
    
    if not has_content and chapters.exists():
        # Project has chapters but no content
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph("Note: This project contains empty chapters. Add content to your chapters to see it in the export.", styles['Italic']))
    
    # Build PDF
    doc.build(elements)
    
    # FileResponse sets the Content-Disposition header
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{project.title}.pdf"'
    
    return response


@login_required
def export_chapter_pdf(request, pk):
    """Export single chapter as PDF"""
    chapter = get_object_or_404(Chapter, pk=pk)
    project = chapter.project
    
    # Check permissions
    if project.author != request.user and request.user not in project.collaborators.all():
        messages.error(request, "You don't have permission to export this chapter.")
        return redirect('writer:project_detail', pk=project.pk)
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    elements = []
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ChapterTitle',
                             parent=styles['Heading1'],
                             fontSize=20,
                             alignment=TA_CENTER,
                             spaceAfter=30))
    styles.add(ParagraphStyle(name='ChapterContent',
                             parent=styles['Normal'],
                             fontSize=12,
                             alignment=TA_JUSTIFY,
                             spaceAfter=12))
    
    # Add chapter title
    elements.append(Paragraph(chapter.title, styles['ChapterTitle']))
    elements.append(Paragraph(f"from {project.title}", styles['Italic']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Add chapter content
    if chapter.content:
        plain_content = strip_html_tags(chapter.content)
        paragraphs = plain_content.split('\n')
        for para in paragraphs:
            if para.strip():
                elements.append(Paragraph(para, styles['ChapterContent']))
                elements.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{project.title} - {chapter.title}.pdf"'
    
    return response


@login_required
def export_document_pdf(request, pk):
    """Export document as PDF"""
    document = get_object_or_404(Document, pk=pk)
    
    # Check permissions
    if document.author != request.user and request.user not in document.shared_with.all():
        messages.error(request, "You don't have permission to export this document.")
        return redirect('writer:document_list')
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    elements = []
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='DocTitle',
                             parent=styles['Heading1'],
                             fontSize=20,
                             alignment=TA_CENTER,
                             spaceAfter=30))
    styles.add(ParagraphStyle(name='DocContent',
                             parent=styles['Normal'],
                             fontSize=12,
                             alignment=TA_JUSTIFY,
                             spaceAfter=12))
    
    # Add document title
    elements.append(Paragraph(document.title, styles['DocTitle']))
    elements.append(Paragraph(f"by {document.author.get_full_name() or document.author.username}", styles['Italic']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Add document content
    if document.content:
        plain_content = strip_html_tags(document.content)
        paragraphs = plain_content.split('\n')
        for para in paragraphs:
            if para.strip():
                elements.append(Paragraph(para, styles['DocContent']))
                elements.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.pdf"'
    
    return response


@login_required
@csrf_exempt
def export_cover_pdf(request):
    """Export book cover as PDF from canvas image data"""
    if request.method == 'POST':
        try:
            import json
            import base64
            from io import BytesIO
            from PIL import Image
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Image as RLImage
            from reportlab.lib.units import inch
            
            data = json.loads(request.body)
            image_data = data.get('image_data')
            filename = data.get('filename', 'book-cover.pdf')
            size = data.get('size', 'kindle')
            
            if not image_data:
                return JsonResponse({'error': 'No image data provided'}, status=400)
            
            # Remove data URL prefix
            if ',' in image_data:
                image_data = image_data.split(',')[1]
            
            # Decode base64 image
            image_bytes = base64.b64decode(image_data)
            image_buffer = BytesIO(image_bytes)
            
            # Open with PIL to get dimensions
            pil_image = Image.open(image_buffer)
            img_width, img_height = pil_image.size
            
            # Create PDF buffer
            pdf_buffer = BytesIO()
            
            # Calculate PDF dimensions based on cover size
            if size == 'kindle':
                page_width, page_height = 6*inch, 9*inch
            elif size == 'paperback':
                page_width, page_height = 6*inch, 9*inch  
            else:  # hardcover
                page_width, page_height = 6.5*inch, 9.5*inch
            
            # Create PDF document with custom page size
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=(page_width, page_height),
                rightMargin=0, leftMargin=0,
                topMargin=0, bottomMargin=0
            )
            
            # Reset image buffer position
            image_buffer.seek(0)
            
            # Create ReportLab image
            rl_image = RLImage(image_buffer, width=page_width, height=page_height)
            
            # Build PDF with single image page
            doc.build([rl_image])
            
            # Return PDF response
            pdf_buffer.seek(0)
            response = HttpResponse(pdf_buffer.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'POST request required'}, status=405)


@login_required
def export_project_epub(request, pk):
    """Export entire project as EPUB"""
    project = get_object_or_404(Project, pk=pk)
    
    # Check permissions
    if project.author != request.user and request.user not in project.collaborators.all():
        messages.error(request, "You don't have permission to export this project.")
        return redirect('writer:project_detail', pk=pk)
    
    try:
        from ebooklib import epub
        import io
        import re
        
        # Create EPUB book
        book = epub.EpubBook()
        book.set_identifier(f'project-{project.id}')
        book.set_title(project.title)
        book.set_language('en')
        book.add_author(project.author.get_full_name() or project.author.username)
        
        if project.description:
            book.add_metadata('DC', 'description', strip_html_tags(project.description))
        
        if project.genre:
            book.add_metadata('DC', 'subject', project.genre)
        
        # Add CSS styles
        style = """
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
        }
        h1 {
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin: 40px 0 30px 0;
            page-break-before: always;
        }
        h2 {
            font-size: 18pt;
            font-weight: bold;
            margin: 30px 0 20px 0;
        }
        p {
            text-indent: 1.5em;
            margin: 0 0 1em 0;
            text-align: justify;
        }
        p.first-paragraph {
            text-indent: 0;
        }
        .scene-break {
            text-align: center;
            margin: 2em 0;
            font-size: 14pt;
        }
        """
        
        nav_css = epub.EpubItem(uid="nav_css", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Get chapters ordered by their sequence
        chapters = project.chapters.all().order_by('order')
        
        epub_chapters = []
        spine_items = ['nav']
        
        # Add title page
        title_page = epub.EpubHtml(
            title='Title Page',
            file_name='title.xhtml',
            lang='en'
        )
        
        title_content = f"""
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>{project.title}</title>
            <link rel="stylesheet" type="text/css" href="style/nav.css"/>
        </head>
        <body>
            <div style="text-align: center; margin-top: 30%; page-break-after: always;">
                <h1 style="font-size: 36pt; margin-bottom: 30px;">{project.title}</h1>
                {"<h2>" + strip_html_tags(project.description) + "</h2>" if project.description else ""}
                <h3 style="margin-top: 60px;">by {project.author.get_full_name() or project.author.username}</h3>
            </div>
        </body>
        </html>
        """
        
        title_page.content = title_content
        book.add_item(title_page)
        spine_items.append(title_page)
        
        # Check if project has any chapters
        if not chapters.exists():
            # Add a placeholder chapter for empty projects
            placeholder_content = f"""
            <html xmlns="http://www.w3.org/1999/xhtml">
            <head>
                <title>Empty Project</title>
                <link rel="stylesheet" type="text/css" href="style/nav.css"/>
            </head>
            <body>
                <h1>Getting Started</h1>
                <p class="first-paragraph">This project doesn't have any chapters yet.</p>
                <p>To add content to your book:</p>
                <ol>
                    <li>Go to your project page</li>
                    <li>Click "Write Now" or "Create First Chapter"</li>
                    <li>Start writing your story</li>
                    <li>Export again to see your content</li>
                </ol>
            </body>
            </html>
            """
            
            placeholder_chapter = epub.EpubHtml(
                title='Getting Started',
                file_name='placeholder.xhtml',
                lang='en'
            )
            placeholder_chapter.content = placeholder_content
            book.add_item(placeholder_chapter)
            spine_items.append(placeholder_chapter)
            epub_chapters = [(epub.Link('placeholder.xhtml', 'Getting Started', 'placeholder'))]
        else:
            # Add each chapter
            has_content = False
        for i, chapter in enumerate(chapters, 1):
            # Clean and process chapter content
            if chapter.content:
                content = strip_html_tags(chapter.content)
                # Check if content is just empty HTML tags
                content_stripped = re.sub(r'<[^>]*>', '', content).strip()
                if not content_stripped:
                    content = ""
            else:
                content = ""
                
            # Split content into paragraphs
            paragraphs = content.split('\n') if content else []
            processed_paragraphs = []
            chapter_has_content = False
            
            for j, para in enumerate(paragraphs):
                para = para.strip()
                if para:
                    # Mark first paragraph of chapter
                    if j == 0:
                        processed_paragraphs.append(f'<p class="first-paragraph">{para}</p>')
                    # Handle scene breaks (lines with just * or # symbols)
                    elif re.match(r'^[\*#\-\s]+$', para):
                        processed_paragraphs.append(f'<div class="scene-break">{para}</div>')
                    else:
                        processed_paragraphs.append(f'<p>{para}</p>')
                    chapter_has_content = True
            
            # If chapter is empty, add placeholder
            if not chapter_has_content:
                processed_paragraphs.append('<p><em>(This chapter is empty)</em></p>')
            else:
                has_content = True
            
            chapter_content = f"""
            <html xmlns="http://www.w3.org/1999/xhtml">
            <head>
                <title>{chapter.title}</title>
                <link rel="stylesheet" type="text/css" href="style/nav.css"/>
            </head>
            <body>
                <h1>{chapter.title}</h1>
                {"".join(processed_paragraphs)}
            </body>
            </html>
            """
            
            epub_chapter = epub.EpubHtml(
                title=chapter.title,
                file_name=f'chapter_{i}.xhtml',
                lang='en'
            )
            epub_chapter.content = chapter_content
            epub_chapter.add_item(nav_css)
            
            book.add_item(epub_chapter)
            epub_chapters.append(epub_chapter)
            spine_items.append(epub_chapter)
        
        # Add table of contents
        book.toc = [(epub.Section('Chapters'), epub_chapters)]
        
        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Set spine (reading order)
        book.spine = spine_items
        
        # Generate EPUB file
        epub_buffer = io.BytesIO()
        epub.write_epub(epub_buffer, book)
        
        # Return as download
        epub_buffer.seek(0)
        response = HttpResponse(epub_buffer.read(), content_type='application/epub+zip')
        filename = f"{project.title.replace(' ', '_')}.epub"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except ImportError:
        messages.error(request, "EPUB export not available. Please contact administrator.")
        return redirect('writer:project_detail', pk=pk)
    except Exception as e:
        messages.error(request, f"Error generating EPUB: {str(e)}")
        return redirect('writer:project_detail', pk=pk)


@login_required 
def export_chapter_epub(request, pk):
    """Export single chapter as EPUB"""
    chapter = get_object_or_404(Chapter, pk=pk)
    project = chapter.project
    
    # Check permissions
    if project.author != request.user and request.user not in project.collaborators.all():
        messages.error(request, "You don't have permission to export this chapter.")
        return redirect('writer:project_detail', pk=project.pk)
    
    try:
        from ebooklib import epub
        import io
        
        # Create EPUB book
        book = epub.EpubBook()
        book.set_identifier(f'chapter-{chapter.id}')
        book.set_title(f"{project.title}: {chapter.title}")
        book.set_language('en')
        book.add_author(project.author.get_full_name() or project.author.username)
        
        # Add CSS styles
        style = """
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
        }
        h1 {
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin: 40px 0 30px 0;
        }
        p {
            text-indent: 1.5em;
            margin: 0 0 1em 0;
            text-align: justify;
        }
        p.first-paragraph {
            text-indent: 0;
        }
        """
        
        nav_css = epub.EpubItem(uid="nav_css", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Process chapter content
        content = strip_html_tags(chapter.content) if chapter.content else ""
        paragraphs = content.split('\n')
        processed_paragraphs = []
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if para:
                if i == 0:
                    processed_paragraphs.append(f'<p class="first-paragraph">{para}</p>')
                else:
                    processed_paragraphs.append(f'<p>{para}</p>')
        
        chapter_content = f"""
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>{chapter.title}</title>
            <link rel="stylesheet" type="text/css" href="style/nav.css"/>
        </head>
        <body>
            <h1>{chapter.title}</h1>
            <p style="text-align: center; font-style: italic; margin-bottom: 30px;">from {project.title}</p>
            {"".join(processed_paragraphs)}
        </body>
        </html>
        """
        
        epub_chapter = epub.EpubHtml(
            title=chapter.title,
            file_name='chapter.xhtml',
            lang='en'
        )
        epub_chapter.content = chapter_content
        epub_chapter.add_item(nav_css)
        
        book.add_item(epub_chapter)
        book.spine = ['nav', epub_chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Generate EPUB
        epub_buffer = io.BytesIO()
        epub.write_epub(epub_buffer, book)
        
        epub_buffer.seek(0)
        response = HttpResponse(epub_buffer.read(), content_type='application/epub+zip')
        filename = f"{project.title}_{chapter.title}".replace(' ', '_') + '.epub'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except ImportError:
        messages.error(request, "EPUB export not available. Please contact administrator.")
        return redirect('writer:project_detail', pk=project.pk)
    except Exception as e:
        messages.error(request, f"Error generating EPUB: {str(e)}")
        return redirect('writer:project_detail', pk=project.pk)


@login_required
def export_document_epub_view(request, pk):
    """Export document as EPUB (view version)"""
    document = get_object_or_404(Document, pk=pk)
    
    # Check permissions
    if document.author != request.user and request.user not in document.shared_with.all():
        messages.error(request, "You don't have permission to export this document.")
        return redirect('writer:document_list')
    
    try:
        from ebooklib import epub
        import io
        
        # Create EPUB book
        book = epub.EpubBook()
        book.set_identifier(f'document-{document.id}')
        book.set_title(document.title)
        book.set_language('en')
        book.add_author(document.author.get_full_name() or document.author.username)
        
        # Add CSS
        style = """
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
        }
        h1 {
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin: 40px 0 30px 0;
        }
        p {
            text-indent: 1.5em;
            margin: 0 0 1em 0;
            text-align: justify;
        }
        p.first-paragraph {
            text-indent: 0;
        }
        """
        
        nav_css = epub.EpubItem(uid="nav_css", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Process document content
        content = strip_html_tags(document.content) if document.content else ""
        paragraphs = content.split('\n')
        processed_paragraphs = []
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if para:
                if i == 0:
                    processed_paragraphs.append(f'<p class="first-paragraph">{para}</p>')
                else:
                    processed_paragraphs.append(f'<p>{para}</p>')
        
        doc_content = f"""
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>{document.title}</title>
            <link rel="stylesheet" type="text/css" href="style/nav.css"/>
        </head>
        <body>
            <h1>{document.title}</h1>
            <p style="text-align: center; font-style: italic; margin-bottom: 30px;">by {document.author.get_full_name() or document.author.username}</p>
            {"".join(processed_paragraphs)}
        </body>
        </html>
        """
        
        epub_doc = epub.EpubHtml(
            title=document.title,
            file_name='document.xhtml',
            lang='en'
        )
        epub_doc.content = doc_content
        epub_doc.add_item(nav_css)
        
        book.add_item(epub_doc)
        book.spine = ['nav', epub_doc]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Generate EPUB
        epub_buffer = io.BytesIO()
        epub.write_epub(epub_buffer, book)
        
        epub_buffer.seek(0)
        response = HttpResponse(epub_buffer.read(), content_type='application/epub+zip')
        filename = f"{document.title}".replace(' ', '_') + '.epub'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except ImportError:
        messages.error(request, "EPUB export not available. Please contact administrator.")
        return redirect('writer:document_list')
    except Exception as e:
        messages.error(request, f"Error generating EPUB: {str(e)}")
        return redirect('writer:document_list')


# ================================
# VERSION CONTROL VIEWS
# ================================

from .models import DocumentVersion, ChapterVersion, VersionControlSettings

def create_document_version(document, user, save_reason=None, is_major=False):
    """Helper function to create a new document version"""
    try:
        # Get the latest version number
        latest_version = DocumentVersion.objects.filter(document=document).first()
        version_number = (latest_version.version_number + 1) if latest_version else 1
        
        # Create new version
        version = DocumentVersion.objects.create(
            document=document,
            version_number=version_number,
            title=document.title,
            content=document.content or '',
            saved_by=user,
            save_reason=save_reason,
            is_major_version=is_major
        )
        
        # Clean up old versions if needed
        settings = VersionControlSettings.get_or_create_for_user(user)
        if settings.max_versions_to_keep > 0:
            old_versions = DocumentVersion.objects.filter(
                document=document
            )[settings.max_versions_to_keep:]
            
            for old_version in old_versions:
                if not old_version.is_major_version and not old_version.is_published_version:
                    old_version.delete()
        
        return version
    except Exception as e:
        print(f"Error creating document version: {e}")
        return None

def create_chapter_version(chapter, user, save_reason=None, is_major=False):
    """Helper function to create a new chapter version"""
    try:
        # Get the latest version number
        latest_version = ChapterVersion.objects.filter(chapter=chapter).first()
        version_number = (latest_version.version_number + 1) if latest_version else 1
        
        # Create new version
        version = ChapterVersion.objects.create(
            chapter=chapter,
            version_number=version_number,
            title=chapter.title,
            content=chapter.content or '',
            saved_by=user,
            save_reason=save_reason,
            is_major_version=is_major
        )
        
        # Clean up old versions if needed
        settings = VersionControlSettings.get_or_create_for_user(user)
        if settings.max_versions_to_keep > 0:
            old_versions = ChapterVersion.objects.filter(
                chapter=chapter
            )[settings.max_versions_to_keep:]
            
            for old_version in old_versions:
                if not old_version.is_major_version and not old_version.is_published_version:
                    old_version.delete()
        
        return version
    except Exception as e:
        print(f"Error creating chapter version: {e}")
        return None

@login_required
def document_versions(request, document_id):
    """View all versions of a document"""
    document = get_object_or_404(Document, id=document_id, author=request.user)
    versions = DocumentVersion.objects.filter(document=document).order_by('-version_number')
    
    # Return JSON for AJAX requests
    if request.headers.get('Content-Type') == 'application/json' or request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest' or 'json' in request.headers.get('Accept', ''):
        versions_data = [{
            'id': version.id,
            'version_name': version.version_name,
            'created_at': version.created_at.isoformat(),
            'version_number': version.version_number
        } for version in versions]
        
        return JsonResponse({
            'success': True,
            'versions': versions_data
        })
    
    context = {
        'document': document,
        'versions': versions,
        'current_version_number': versions.first().version_number if versions else 0
    }
    return render(request, 'writer/document_versions.html', context)

@login_required
def chapter_versions(request, chapter_id):
    """View all versions of a chapter"""
    chapter = get_object_or_404(Chapter, id=chapter_id)
    
    # Check permissions
    if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
        if request.headers.get('Content-Type') == 'application/json' or request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest' or 'json' in request.headers.get('Accept', ''):
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        messages.error(request, "You don't have permission to view this chapter's versions.")
        return redirect('writer:project_list')
    
    versions = ChapterVersion.objects.filter(chapter=chapter).order_by('-version_number')
    
    # Return JSON for AJAX requests
    if request.headers.get('Content-Type') == 'application/json' or request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest' or 'json' in request.headers.get('Accept', ''):
        versions_data = [{
            'id': version.id,
            'version_name': version.version_name,
            'created_at': version.created_at.isoformat(),
            'version_number': version.version_number
        } for version in versions]
        
        return JsonResponse({
            'success': True,
            'versions': versions_data
        })
    
    context = {
        'chapter': chapter,
        'versions': versions,
        'current_version_number': versions.first().version_number if versions else 0
    }
    return render(request, 'writer/chapter_versions.html', context)

@login_required
def save_document_version(request, document_id):
    """Manually save a version of a document"""
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, author=request.user)
        save_reason = request.POST.get('save_reason', '')
        is_major = request.POST.get('is_major') == 'on'
        
        version = create_document_version(document, request.user, save_reason, is_major)
        
        if version:
            messages.success(request, f'Version {version.version_number} saved successfully.')
        else:
            messages.error(request, 'Failed to save version.')
        
        return redirect('writer:document_versions', document_id=document_id)
    
    return redirect('writer:document_detail', pk=document_id)

@login_required  
def save_chapter_version(request, chapter_id):
    """Manually save a version of a chapter"""
    if request.method == 'POST':
        chapter = get_object_or_404(Chapter, id=chapter_id)
        
        # Check permissions
        if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
            messages.error(request, "You don't have permission to save this chapter.")
            return redirect('writer:project_list')
        
        save_reason = request.POST.get('save_reason', '')
        is_major = request.POST.get('is_major') == 'on'
        
        version = create_chapter_version(chapter, request.user, save_reason, is_major)
        
        if version:
            messages.success(request, f'Version {version.version_number} saved successfully.')
        else:
            messages.error(request, 'Failed to save version.')
        
        return redirect('writer:chapter_versions', chapter_id=chapter_id)
    
    return redirect('writer:chapter_detail', pk=chapter_id)

@login_required
def restore_document_version(request, document_id, version_id):
    """Restore a document to a specific version"""
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, author=request.user)
        version = get_object_or_404(DocumentVersion, id=version_id, document=document)
        
        # Create a backup of current version first
        create_document_version(document, request.user, f"Backup before restoring to version {version.version_number}")
        
        # Restore the content
        document.title = version.title
        document.content = version.content
        document.save()
        
        # Create a new version to mark the restoration
        restore_version = create_document_version(
            document, request.user, 
            f"Restored to version {version.version_number}", 
            is_major=True
        )
        
        messages.success(request, f'Document restored to version {version.version_number}.')
        return redirect('writer:document_detail', pk=document_id)
    
    return redirect('writer:document_versions', document_id=document_id)

@login_required
def restore_chapter_version(request, chapter_id, version_id):
    """Restore a chapter to a specific version"""
    if request.method == 'POST':
        chapter = get_object_or_404(Chapter, id=chapter_id)
        version = get_object_or_404(ChapterVersion, id=version_id, chapter=chapter)
        
        # Check permissions
        if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
            messages.error(request, "You don't have permission to restore this chapter.")
            return redirect('writer:project_list')
        
        # Create a backup of current version first
        create_chapter_version(chapter, request.user, f"Backup before restoring to version {version.version_number}")
        
        # Restore the content
        chapter.title = version.title
        chapter.content = version.content
        chapter.save()
        
        # Create a new version to mark the restoration
        restore_version = create_chapter_version(
            chapter, request.user, 
            f"Restored to version {version.version_number}", 
            is_major=True
        )
        
        messages.success(request, f'Chapter restored to version {version.version_number}.')
        return redirect('writer:chapter_detail', pk=chapter_id)
    
    return redirect('writer:chapter_versions', chapter_id=chapter_id)

@login_required
def compare_document_versions(request, document_id, version1_id, version2_id):
    """Compare two versions of a document"""
    document = get_object_or_404(Document, id=document_id, author=request.user)
    version1 = get_object_or_404(DocumentVersion, id=version1_id, document=document)
    version2 = get_object_or_404(DocumentVersion, id=version2_id, document=document)
    
    context = {
        'document': document,
        'version1': version1,
        'version2': version2,
    }
    return render(request, 'writer/compare_document_versions.html', context)

@login_required
def compare_chapter_versions(request, chapter_id, version1_id, version2_id):
    """Compare two versions of a chapter"""
    chapter = get_object_or_404(Chapter, id=chapter_id)
    
    # Check permissions
    if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
        messages.error(request, "You don't have permission to view this chapter.")
        return redirect('writer:project_list')
    
    version1 = get_object_or_404(ChapterVersion, id=version1_id, chapter=chapter)
    version2 = get_object_or_404(ChapterVersion, id=version2_id, chapter=chapter)
    
    context = {
        'chapter': chapter,
        'version1': version1,
        'version2': version2,
    }
    return render(request, 'writer/compare_chapter_versions.html', context)

@login_required
def version_control_settings(request):
    """Manage version control settings"""
    settings = VersionControlSettings.get_or_create_for_user(request.user)
    
    if request.method == 'POST':
        settings.auto_save_enabled = request.POST.get('auto_save_enabled') == 'on'
        settings.auto_save_interval = int(request.POST.get('auto_save_interval', 10))
        settings.max_versions_to_keep = int(request.POST.get('max_versions_to_keep', 50))
        settings.notify_on_version_save = request.POST.get('notify_on_version_save') == 'on'
        settings.require_save_reason = request.POST.get('require_save_reason') == 'on'
        settings.save()
        
        messages.success(request, 'Version control settings updated successfully.')
        return redirect('writer:version_control_settings')
    
    context = {'settings': settings}
    return render(request, 'writer/version_control_settings.html', context)

@csrf_exempt
@login_required
def auto_save_version_api(request):
    """API endpoint for auto-saving versions"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content_type = data.get('type')  # 'document' or 'chapter'
            content_id = data.get('id')
            save_reason = data.get('reason', 'Auto-save')
            
            if content_type == 'document':
                document = get_object_or_404(Document, id=content_id, author=request.user)
                version = create_document_version(document, request.user, save_reason)
                
                if version:
                    return JsonResponse({
                        'success': True,
                        'version_number': version.version_number,
                        'message': f'Auto-saved as version {version.version_number}'
                    })
            
            elif content_type == 'chapter':
                chapter = get_object_or_404(Chapter, id=content_id)
                
                # Check permissions
                if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
                    return JsonResponse({'success': False, 'error': 'Permission denied'})
                
                version = create_chapter_version(chapter, request.user, save_reason)
                
                if version:
                    return JsonResponse({
                        'success': True,
                        'version_number': version.version_number,
                        'message': f'Auto-saved as version {version.version_number}'
                    })
            
            return JsonResponse({'success': False, 'error': 'Failed to create version'})
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@login_required
def delete_document_version(request, document_id, version_id):
    """Delete a specific document version"""
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, author=request.user)
        version = get_object_or_404(DocumentVersion, id=version_id, document=document)
        
        # Don't allow deletion of published or major versions
        if version.is_published_version or version.is_major_version:
            messages.error(request, 'Cannot delete published or major versions.')
        else:
            version_number = version.version_number
            version.delete()
            messages.success(request, f'Version {version_number} deleted.')
        
        return redirect('writer:document_versions', document_id=document_id)
    
    return redirect('writer:document_versions', document_id=document_id)

@login_required  
def delete_chapter_version(request, chapter_id, version_id):
    """Delete a specific chapter version"""
    if request.method == 'POST':
        chapter = get_object_or_404(Chapter, id=chapter_id)
        version = get_object_or_404(ChapterVersion, id=version_id, chapter=chapter)
        
        # Check permissions
        if chapter.project.author != request.user and request.user not in chapter.project.collaborators.all():
            messages.error(request, "You don't have permission to delete this version.")
            return redirect('writer:project_list')
        
        # Don't allow deletion of published or major versions
        if version.is_published_version or version.is_major_version:
            messages.error(request, 'Cannot delete published or major versions.')
        else:
            version_number = version.version_number
            version.delete()
            messages.success(request, f'Version {version_number} deleted.')
        
        return redirect('writer:chapter_versions', chapter_id=chapter_id)
    
    return redirect('writer:chapter_versions', chapter_id=chapter_id)


# Chapter Auto-Detection Functions
def detect_chapters_from_text(text):
    """
    Detect chapters from text using common patterns:
    - Lines starting with "Chapter", "CHAPTER", "#" (Markdown)
    - Numbered headings (e.g., "Chapter 1", "Chapter 2")
    - HTML heading tags (<h1>, <h2>, etc.)
    """
    if not text:
        return []
    
    chapters = []
    lines = text.split('\n')
    current_chapter = None
    current_content = []
    chapter_number = 0
    
    # Regex patterns for chapter detection
    chapter_patterns = [
        # Standard chapter formats
        r'^\s*Chapter\s+(\d+|[IVXLCDM]+|One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten)\s*:?\s*(.*)$',
        r'^\s*CHAPTER\s+(\d+|[IVXLCDM]+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*:?\s*(.*)$',
        # Markdown headers
        r'^\s*#{1,3}\s+(.+)$',
        # Numbered sections
        r'^\s*(\d+)\.\s+(.+)$',
        # HTML headers
        r'<h[1-3][^>]*>(.*?)</h[1-3]>',
        # Alternative chapter formats
        r'^\s*(Part|PART)\s+(\d+|[IVXLCDM]+|One|Two|Three|Four|Five)\s*:?\s*(.*)$',
        r'^\s*-+\s*Chapter\s+(.+?)\s*-+$',
        r'^\s*=+\s*Chapter\s+(.+?)\s*=+$',
    ]
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            if current_content:
                current_content.append('')
            continue
            
        is_chapter_header = False
        chapter_title = None
        
        # Check each pattern
        for pattern in chapter_patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                is_chapter_header = True
                # Extract title from different match groups
                if len(match.groups()) >= 2:
                    chapter_title = match.group(2).strip() if match.group(2) else f"Chapter {match.group(1)}"
                elif len(match.groups()) == 1:
                    chapter_title = match.group(1).strip()
                else:
                    chapter_title = line.strip()
                break
        
        # Special handling for simple chapter numbers or roman numerals
        if not is_chapter_header:
            # Check for standalone chapter indicators
            if re.match(r'^\s*\d+\s*$', line) and len(line.strip()) <= 3:
                is_chapter_header = True
                chapter_title = f"Chapter {line.strip()}"
            elif re.match(r'^\s*[IVXLCDM]+\s*$', line):
                is_chapter_header = True
                chapter_title = f"Chapter {line.strip()}"
        
        if is_chapter_header:
            # Save previous chapter if it exists
            if current_chapter is not None and current_content:
                chapters.append({
                    'number': current_chapter['number'],
                    'title': current_chapter['title'],
                    'content': '\n'.join(current_content).strip(),
                    'start_line': current_chapter['start_line'],
                    'end_line': i - 1
                })
            
            # Start new chapter
            chapter_number += 1
            current_chapter = {
                'number': chapter_number,
                'title': chapter_title or f"Chapter {chapter_number}",
                'start_line': i
            }
            current_content = []
        else:
            # Add content to current chapter
            if current_content or line:  # Don't add leading empty lines
                current_content.append(line)
    
    # Add the last chapter
    if current_chapter is not None and current_content:
        chapters.append({
            'number': current_chapter['number'],
            'title': current_chapter['title'],
            'content': '\n'.join(current_content).strip(),
            'start_line': current_chapter['start_line'],
            'end_line': len(lines) - 1
        })
    
    # If no chapters were detected, treat the entire text as one chapter
    if not chapters and text.strip():
        chapters.append({
            'number': 1,
            'title': 'Chapter 1',
            'content': text.strip(),
            'start_line': 0,
            'end_line': len(lines) - 1
        })
    
    return chapters


@csrf_exempt
@login_required
def detect_chapters_api(request):
    """API endpoint for chapter detection"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            text = data.get('text', '')
            project_id = data.get('project_id')
            
            if not text:
                return JsonResponse({'success': False, 'error': 'No text provided'})
            
            # Detect chapters
            chapters = detect_chapters_from_text(text)
            
            return JsonResponse({
                'success': True,
                'chapters': chapters,
                'total_chapters': len(chapters)
            })
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'})


# ============================================================================
# CLASSICS LIBRARY VIEWS - Public Domain Books Feature
# ============================================================================

@login_required
def classics_library(request):
    """Main classics library view - displays public domain books"""
    from .models import ClassicBook, BorrowedBook
    
    # Get user's borrowed books for status tracking
    user_borrowed_books = set(
        BorrowedBook.objects.filter(user=request.user)
        .values_list('book_id', flat=True)
    )
    
    # Get featured and recent books for initial display
    featured_books = ClassicBook.objects.filter(is_featured=True, is_active=True)[:12]
    recent_books = ClassicBook.objects.filter(is_active=True).order_by('-created_at')[:12]
    
    # Get statistics
    total_books = ClassicBook.objects.filter(is_active=True).count()
    total_authors = ClassicBook.objects.filter(is_active=True).values('author').distinct().count()
    user_borrowed_count = BorrowedBook.objects.filter(user=request.user).count()
    
    context = {
        'featured_books': featured_books,
        'recent_books': recent_books,
        'total_books': total_books,
        'total_authors': total_authors,
        'user_borrowed_count': user_borrowed_count,
        'user_borrowed_books': user_borrowed_books,
    }
    
    return render(request, 'writer/classics_library.html', context)


@csrf_exempt
def classics_search_api(request):
    """API endpoint for searching classic books"""
    if request.method == 'GET':
        from .models import ClassicBook
        from django.db.models import Q
        
        # Get search parameters
        query = request.GET.get('query', '').strip()
        genre = request.GET.get('genre', '')
        language = request.GET.get('language', '')
        era = request.GET.get('era', '')
        page = int(request.GET.get('page', 1))
        per_page = int(request.GET.get('per_page', 24))
        
        # Start with active books
        books = ClassicBook.objects.filter(is_active=True)
        
        # Apply search query
        if query:
            books = books.filter(
                Q(title__icontains=query) |
                Q(author__icontains=query) |
                Q(description__icontains=query) |
                Q(subjects__icontains=query)
            )
        
        # Apply filters
        if genre:
            books = books.filter(genre=genre)
        if language:
            books = books.filter(language=language)
        if era:
            books = books.filter(era=era)
        
        # Order by rating and download count
        books = books.order_by('-is_featured', '-rating', '-download_count')
        
        # Get user's borrowed books (only if authenticated)
        from .models import BorrowedBook
        user_borrowed_books = set()
        if request.user.is_authenticated:
            user_borrowed_books = set(
                BorrowedBook.objects.filter(user=request.user)
                .values_list('book_id', flat=True)
            )
        
        # Pagination
        total_count = books.count()
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        books_page = books[start_index:end_index]
        
        # Serialize books
        books_data = []
        for book in books_page:
            books_data.append({
                'id': book.id,
                'title': book.title,
                'author': book.author,
                'genre': book.genre,
                'language': book.language,
                'era': book.era,
                'publication_year': book.publication_year,
                'description': book.description,
                'page_count': book.page_count,
                'rating': float(book.rating) if book.rating else 0,
                'download_count': book.download_count,
                'cover_image_url': book.cover_image_url,
                'available_formats': book.available_formats,
                'is_borrowed': book.id in user_borrowed_books,
                'epub_url': book.epub_url,
                'pdf_url': book.pdf_url,
                'txt_url': book.txt_url,
                'html_url': book.html_url,
                'gutenberg_id': book.gutenberg_id,
            })
        
        return JsonResponse({
            'success': True,
            'books': books_data,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total_count': total_count,
                'total_pages': (total_count + per_page - 1) // per_page,
            }
        })
    
    return JsonResponse({'error': 'GET method required'}, status=405)


@csrf_exempt
@login_required
def borrow_classic_book(request, book_id):
    """API endpoint to borrow a classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id, is_active=True)
            
            # Check if already borrowed
            borrowed_book, created = BorrowedBook.objects.get_or_create(
                user=request.user,
                book=book,
                defaults={
                    'reading_progress': 0,
                    'current_page': 1,
                    'reading_status': 'not_started'
                }
            )
            
            if created:
                # Increment download count
                book.download_count += 1
                book.save()
                
                return JsonResponse({
                    'success': True,
                    'message': f'"{book.title}" has been added to your library!',
                    'book_id': book.id,
                    'borrowed_at': borrowed_book.borrowed_at.isoformat()
                })
            else:
                return JsonResponse({
                    'success': False,
                    'message': f'"{book.title}" is already in your library.',
                    'book_id': book.id
                })
                
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)


@csrf_exempt
@login_required
def return_classic_book(request, book_id):
    """API endpoint to return a borrowed classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id)
            borrowed_book = get_object_or_404(BorrowedBook, user=request.user, book=book)
            
            book_title = book.title
            borrowed_book.delete()
            
            return JsonResponse({
                'success': True,
                'message': f'"{book_title}" has been returned.',
                'book_id': book.id
            })
            
        except BorrowedBook.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Book is not in your library.'
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)


@csrf_exempt
def read_classic_book(request, book_id):
    """API endpoint to start reading a classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        import json
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id, is_active=True)
            
            # Check if book is borrowed (only if user is authenticated)
            borrowed_book = None
            if request.user.is_authenticated:
                try:
                    borrowed_book = BorrowedBook.objects.get(user=request.user, book=book)
                except BorrowedBook.DoesNotExist:
                    # Auto-borrow if not already borrowed
                    borrowed_book = BorrowedBook.objects.create(
                        user=request.user,
                        book=book,
                        reading_progress=0,
                        current_page=1,
                        reading_status='reading'
                    )
                    book.download_count += 1
                    book.save()
            else:
                # For anonymous users, just track download count
                book.download_count += 1
                book.save()
            
            # Update reading status (only for authenticated users)
            if borrowed_book:
                if borrowed_book.reading_status == 'not_started':
                    borrowed_book.reading_status = 'reading'
                    borrowed_book.save()
                
                # Get reading data
                if request.content_type == 'application/json':
                    data = json.loads(request.body)
                    if 'progress' in data:
                        borrowed_book.reading_progress = min(100, max(0, int(data['progress'])))
                    if 'page' in data:
                        borrowed_book.current_page = max(1, int(data['page']))
                    borrowed_book.save()
            
            return JsonResponse({
                'success': True,
                'message': f'Opening "{book.title}" for reading...',
                'book': {
                    'id': book.id,
                    'title': book.title,
                    'author': book.author,
                    'available_formats': book.available_formats,
                    'primary_download_url': book.primary_download_url,
                    'epub_url': book.epub_url,
                    'pdf_url': book.pdf_url,
                    'txt_url': book.txt_url,
                    'html_url': book.html_url,
                    'gutenberg_id': book.gutenberg_id,
                },
                'reading_progress': borrowed_book.reading_progress if borrowed_book else 0,
                'current_page': borrowed_book.current_page if borrowed_book else 1,
                'reading_status': borrowed_book.reading_status if borrowed_book else 'reading'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)


@csrf_exempt
@login_required
def create_chapters_from_detection(request):
    """Create actual chapter objects from detected chapters"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_id = data.get('project_id')
            chapters_data = data.get('chapters', [])
            
            if not project_id:
                return JsonResponse({'success': False, 'error': 'No project ID provided'})
            
            if not chapters_data:
                return JsonResponse({'success': False, 'error': 'No chapters data provided'})
            
            # Get the project
            project = get_object_or_404(Project, id=project_id, author=request.user)
            
            created_chapters = []
            
            with transaction.atomic():
                # Delete existing chapters if requested
                if data.get('replace_existing', False):
                    Chapter.objects.filter(project=project).delete()
                
                # Create new chapters
                for chapter_data in chapters_data:
                    chapter = Chapter.objects.create(
                        title=chapter_data.get('title', f"Chapter {chapter_data.get('number', 1)}"),
                        content=chapter_data.get('content', ''),
                        project=project,
                        order=chapter_data.get('number', 1),
                        word_count=len(chapter_data.get('content', '').split())
                    )
                    created_chapters.append({
                        'id': chapter.id,
                        'title': chapter.title,
                        'order': chapter.order,
                        'word_count': chapter.word_count
                    })
            
            return JsonResponse({
                'success': True,
                'chapters': created_chapters,
                'total_created': len(created_chapters)
            })
            
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Method not allowed'})


# ============================================================================
# CLASSICS LIBRARY VIEWS - Public Domain Books Feature
# ============================================================================

@login_required
def classics_library(request):
    """Main classics library view - displays public domain books"""
    from .models import ClassicBook, BorrowedBook
    
    # Get user's borrowed books for status tracking
    user_borrowed_books = set(
        BorrowedBook.objects.filter(user=request.user)
        .values_list('book_id', flat=True)
    )
    
    # Get featured and recent books for initial display
    featured_books = ClassicBook.objects.filter(is_featured=True, is_active=True)[:12]
    recent_books = ClassicBook.objects.filter(is_active=True).order_by('-created_at')[:12]
    
    # Get statistics
    total_books = ClassicBook.objects.filter(is_active=True).count()
    total_authors = ClassicBook.objects.filter(is_active=True).values('author').distinct().count()
    user_borrowed_count = BorrowedBook.objects.filter(user=request.user).count()
    
    context = {
        'featured_books': featured_books,
        'recent_books': recent_books,
        'total_books': total_books,
        'total_authors': total_authors,
        'user_borrowed_count': user_borrowed_count,
        'user_borrowed_books': user_borrowed_books,
    }
    
    return render(request, 'writer/classics_library.html', context)


@csrf_exempt
def classics_search_api(request):
    """API endpoint for searching classic books"""
    if request.method == 'GET':
        from .models import ClassicBook
        from django.db.models import Q
        
        # Get search parameters
        query = request.GET.get('query', '').strip()
        genre = request.GET.get('genre', '')
        language = request.GET.get('language', '')
        era = request.GET.get('era', '')
        page = int(request.GET.get('page', 1))
        per_page = int(request.GET.get('per_page', 24))
        
        # Start with active books
        books = ClassicBook.objects.filter(is_active=True)
        
        # Apply search query
        if query:
            books = books.filter(
                Q(title__icontains=query) |
                Q(author__icontains=query) |
                Q(description__icontains=query) |
                Q(subjects__icontains=query)
            )
        
        # Apply filters
        if genre:
            books = books.filter(genre=genre)
        if language:
            books = books.filter(language=language)
        if era:
            books = books.filter(era=era)
        
        # Order by rating and download count
        books = books.order_by('-is_featured', '-rating', '-download_count')
        
        # Get user's borrowed books (only if authenticated)
        from .models import BorrowedBook
        user_borrowed_books = set()
        if request.user.is_authenticated:
            user_borrowed_books = set(
                BorrowedBook.objects.filter(user=request.user)
                .values_list('book_id', flat=True)
            )
        
        # Pagination
        total_count = books.count()
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        books_page = books[start_index:end_index]
        
        # Serialize books
        books_data = []
        for book in books_page:
            books_data.append({
                'id': book.id,
                'title': book.title,
                'author': book.author,
                'genre': book.genre,
                'language': book.language,
                'era': book.era,
                'publication_year': book.publication_year,
                'description': book.description,
                'page_count': book.page_count,
                'rating': float(book.rating) if book.rating else 0,
                'download_count': book.download_count,
                'cover_image_url': book.cover_image_url,
                'available_formats': book.available_formats,
                'is_borrowed': book.id in user_borrowed_books,
                'epub_url': book.epub_url,
                'pdf_url': book.pdf_url,
                'txt_url': book.txt_url,
                'html_url': book.html_url,
                'gutenberg_id': book.gutenberg_id,
            })
        
        return JsonResponse({
            'success': True,
            'books': books_data,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total_count': total_count,
                'total_pages': (total_count + per_page - 1) // per_page,
            }
        })
    
    return JsonResponse({'error': 'GET method required'}, status=405)


@csrf_exempt
@login_required
def borrow_classic_book(request, book_id):
    """API endpoint to borrow a classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id, is_active=True)
            
            # Check if already borrowed
            borrowed_book, created = BorrowedBook.objects.get_or_create(
                user=request.user,
                book=book,
                defaults={
                    'reading_progress': 0,
                    'current_page': 1,
                    'reading_status': 'not_started'
                }
            )
            
            if created:
                # Increment download count
                book.download_count += 1
                book.save()
                
                return JsonResponse({
                    'success': True,
                    'message': f'"{book.title}" has been added to your library!',
                    'book_id': book.id,
                    'borrowed_at': borrowed_book.borrowed_at.isoformat()
                })
            else:
                return JsonResponse({
                    'success': False,
                    'message': f'"{book.title}" is already in your library.',
                    'book_id': book.id
                })
                
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)


@csrf_exempt
@login_required
def return_classic_book(request, book_id):
    """API endpoint to return a borrowed classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id)
            borrowed_book = get_object_or_404(BorrowedBook, user=request.user, book=book)
            
            book_title = book.title
            borrowed_book.delete()
            
            return JsonResponse({
                'success': True,
                'message': f'"{book_title}" has been returned.',
                'book_id': book.id
            })
            
        except BorrowedBook.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Book is not in your library.'
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)


@csrf_exempt
def read_classic_book(request, book_id):
    """API endpoint to start reading a classic book"""
    if request.method == 'POST':
        from .models import ClassicBook, BorrowedBook
        import json
        
        try:
            book = get_object_or_404(ClassicBook, id=book_id, is_active=True)
            
            # Check if book is borrowed (only if user is authenticated)
            borrowed_book = None
            if request.user.is_authenticated:
                try:
                    borrowed_book = BorrowedBook.objects.get(user=request.user, book=book)
                except BorrowedBook.DoesNotExist:
                    # Auto-borrow if not already borrowed
                    borrowed_book = BorrowedBook.objects.create(
                        user=request.user,
                        book=book,
                        reading_progress=0,
                        current_page=1,
                        reading_status='reading'
                    )
                    book.download_count += 1
                    book.save()
            else:
                # For anonymous users, just track download count
                book.download_count += 1
                book.save()
            
            # Update reading status (only for authenticated users)
            if borrowed_book:
                if borrowed_book.reading_status == 'not_started':
                    borrowed_book.reading_status = 'reading'
                    borrowed_book.save()
                
                # Get reading data
                if request.content_type == 'application/json':
                    data = json.loads(request.body)
                    if 'progress' in data:
                        borrowed_book.reading_progress = min(100, max(0, int(data['progress'])))
                    if 'page' in data:
                        borrowed_book.current_page = max(1, int(data['page']))
                    borrowed_book.save()
            
            return JsonResponse({
                'success': True,
                'message': f'Opening "{book.title}" for reading...',
                'book': {
                    'id': book.id,
                    'title': book.title,
                    'author': book.author,
                    'available_formats': book.available_formats,
                    'primary_download_url': book.primary_download_url,
                    'epub_url': book.epub_url,
                    'pdf_url': book.pdf_url,
                    'txt_url': book.txt_url,
                    'html_url': book.html_url,
                    'gutenberg_id': book.gutenberg_id,
                },
                'reading_progress': borrowed_book.reading_progress if borrowed_book else 0,
                'current_page': borrowed_book.current_page if borrowed_book else 1,
                'reading_status': borrowed_book.reading_status if borrowed_book else 'reading'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'POST method required'}, status=405)



@csrf_exempt
def fetch_book_content(request):
    """API endpoint to fetch book content from external URLs"""
    if request.method == "POST":
        import json
        import requests
        
        try:
            data = json.loads(request.body)
            url = data.get("url")
            title = data.get("title", "Unknown")
            
            if not url:
                return JsonResponse({"success": False, "error": "No URL provided"})
            
            # Fetch content from the URL
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            # Get the text content
            content = response.text
            
            # For HTML content, try to extract just the text
            if "html" in url.lower() or "htm" in url.lower():
                # Simple HTML tag removal (in production, use BeautifulSoup)
                import re
                # Remove script and style elements
                content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL)
                content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL)
                # Remove HTML tags
                content = re.sub(r"<[^>]+>", "", content)
                # Clean up whitespace
                content = re.sub(r"\s+", " ", content)
            
            # Limit content size to prevent huge responses
            max_length = 500000  # 500KB of text
            if len(content) > max_length:
                content = content[:max_length] + "\n\n... [Content truncated due to size]"
            
            return JsonResponse({
                "success": True,
                "content": content,
                "title": title
            })
            
        except requests.RequestException as e:
            return JsonResponse({
                "success": False,
                "error": f"Failed to fetch content: {str(e)}"
            })
        except Exception as e:
            return JsonResponse({
                "success": False,
                "error": str(e)
            })
    
    return JsonResponse({"error": "POST method required"}, status=405)
