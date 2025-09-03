"""
Document parsing utilities for different file formats.
Handles PDF, DOCX, TXT, HTML, and Google Docs files for rich text editing.
"""

import os
import re
from typing import Dict, List, Tuple, Optional
import mimetypes
# import requests
# from urllib.parse import urlparse, parse_qs

def extract_text_from_file(file_path: str, google_docs_url: Optional[str] = None) -> str:
    """
    Extract text content from various file formats with rich formatting preservation.
    
    Args:
        file_path: Path to the file to parse
        google_docs_url: Optional Google Docs sharing URL for direct import
        
    Returns:
        Extracted text content as HTML string with formatting preserved
    """
    # Handle Google Docs URL
    if google_docs_url:
        return "<p>Google Docs import temporarily disabled. Please export as DOCX and upload the file instead.</p>"
        # return extract_from_google_docs(google_docs_url)
    
    if not os.path.exists(file_path):
        return ""
    
    mime_type, _ = mimetypes.guess_type(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt' or mime_type == 'text/plain':
            return extract_from_txt(file_path)
        elif file_extension == '.html' or mime_type == 'text/html':
            return extract_from_html(file_path)
        elif file_extension == '.pdf':
            return extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_from_docx(file_path)
        elif file_extension == '.rtf':
            return extract_from_rtf(file_path)
        elif file_extension == '.odt':
            return extract_from_odt(file_path)
        else:
            # Try to read as plain text
            return extract_from_txt(file_path)
    except Exception as e:
        print(f"Error parsing file {file_path}: {e}")
        return ""

def extract_from_txt(file_path: str) -> str:
    """Extract text from plain text file and convert to HTML."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        # Convert plain text to HTML with paragraph breaks
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        paragraphs = content.split('\n\n')
        html_content = ''
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Preserve line breaks within paragraphs
                paragraph = paragraph.replace('\n', '<br>')
                html_content += f'<p>{paragraph}</p>\n'
        content_with_br = content.replace("\n", "<br>")
        return html_content if html_content else f'<p>{content_with_br}</p>'
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as file:
            content = file.read()
        # Convert to HTML format
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        paragraphs = content.split('\n\n')
        html_content = ''
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                paragraph = paragraph.replace('\n', '<br>')
                html_content += f'<p>{paragraph}</p>\n'
        content_with_br = content.replace("\n", "<br>")
        return html_content if html_content else f'<p>{content_with_br}</p>'

def extract_from_html(file_path: str) -> str:
    """Extract text from HTML file while preserving structure."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Remove script and style elements
        content = re.sub(r'<script.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'<style.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        return content
    except Exception as e:
        print(f"Error parsing HTML: {e}")
        return ""

def extract_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file with formatting preservation.
    """
    try:
        import PyPDF2
        import pdfplumber
        
        # Try pdfplumber first for better formatting
        try:
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n\n"
                
                if text_content.strip():
                    # Convert to HTML with proper paragraph formatting
                    paragraphs = text_content.split('\n\n')
                    html_content = ''
                    
                    for paragraph in paragraphs:
                        paragraph = paragraph.strip()
                        if paragraph:
                            # Clean up excessive whitespace
                            paragraph = re.sub(r'\s+', ' ', paragraph)
                            # Preserve intentional line breaks
                            paragraph = paragraph.replace('\n', '<br>')
                            html_content += f'<p>{paragraph}</p>\n'
                    
                    return html_content
        except Exception as e:
            print(f"pdfplumber failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
            
            if text_content.strip():
                # Convert to HTML
                paragraphs = text_content.split('\n\n')
                html_content = ''
                
                for paragraph in paragraphs:
                    paragraph = paragraph.strip()
                    if paragraph:
                        paragraph = re.sub(r'\s+', ' ', paragraph)
                        paragraph = paragraph.replace('\n', '<br>')
                        html_content += f'<p>{paragraph}</p>\n'
                
                return html_content
            
        return "<p>Could not extract text from PDF file.</p>"
        
    except ImportError:
        return "<p>PDF parsing libraries not installed. Please install PyPDF2 and pdfplumber.</p>"
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return f"<p>Error reading PDF file: {str(e)}</p>"

def extract_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file with rich formatting preservation.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document(file_path)
        html_content = ""
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            # Start paragraph tag
            p_style = ""
            
            # Handle alignment
            if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                p_style += "text-align: center; "
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                p_style += "text-align: right; "
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                p_style += "text-align: justify; "
            
            # Check if it's a heading
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    if level_num <= 6:
                        tag = f"h{level_num}"
                        style_attr = f' style="{p_style}"' if p_style else ''
                        html_content += f"<{tag}{style_attr}>"
                    else:
                        tag = "p"
                        p_style += "font-weight: bold; font-size: 1.2em; "
                        style_attr = f' style="{p_style}"' if p_style else ''
                        html_content += f"<{tag}{style_attr}>"
                except ValueError:
                    tag = "p"
                    p_style += "font-weight: bold; "
                    style_attr = f' style="{p_style}"' if p_style else ''
                    html_content += f"<{tag}{style_attr}>"
            else:
                tag = "p"
                style_attr = f' style="{p_style}"' if p_style else ''
                html_content += f"<{tag}{style_attr}>"
            
            # Process text with formatting
            text_content = ""
            for run in paragraph.runs:
                run_text = run.text
                if not run_text:
                    continue
                
                # Apply formatting
                if run.bold and run.italic:
                    run_text = f"<strong><em>{run_text}</em></strong>"
                elif run.bold:
                    run_text = f"<strong>{run_text}</strong>"
                elif run.italic:
                    run_text = f"<em>{run_text}</em>"
                
                if run.underline:
                    run_text = f"<u>{run_text}</u>"
                
                text_content += run_text
            
            html_content += text_content + f"</{tag}>\n"
        
        # Handle tables
        for table in doc.tables:
            html_content += "<table border='1' style='border-collapse: collapse; width: 100%; margin: 10px 0;'>\n"
            for row in table.rows:
                html_content += "<tr>\n"
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        html_content += f"<td style='padding: 8px; border: 1px solid #ccc;'>{cell_text}</td>\n"
                    else:
                        html_content += "<td style='padding: 8px; border: 1px solid #ccc;'>&nbsp;</td>\n"
                html_content += "</tr>\n"
            html_content += "</table>\n"
        
        return html_content if html_content.strip() else "<p>No readable content found in document.</p>"
        
    except ImportError:
        return "<p>DOCX parsing library not installed. Please install python-docx.</p>"
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return f"<p>Error reading DOCX file: {str(e)}</p>"


def extract_from_rtf(file_path: str) -> str:
    """
    Extract text from RTF file with basic formatting preservation.
    """
    try:
        from striprtf.striprtf import rtf_to_text
        
        with open(file_path, 'rb') as file:
            rtf_content = file.read().decode('utf-8', errors='ignore')
        
        # Extract plain text
        plain_text = rtf_to_text(rtf_content)
        
        # Convert to HTML with paragraph formatting
        if plain_text.strip():
            paragraphs = plain_text.split('\n\n')
            html_content = ''
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if paragraph:
                    # Clean up excessive whitespace
                    paragraph = re.sub(r'\s+', ' ', paragraph)
                    # Preserve intentional line breaks
                    paragraph = paragraph.replace('\n', '<br>')
                    html_content += f'<p>{paragraph}</p>\n'
            
            return html_content
        
        return "<p>No readable content found in RTF file.</p>"
        
    except ImportError:
        return "<p>RTF parsing library not installed. Please install striprtf.</p>"
    except Exception as e:
        print(f"Error parsing RTF: {e}")
        return f"<p>Error reading RTF file: {str(e)}</p>"


def extract_from_odt(file_path: str) -> str:
    """
    Extract text from ODT file with formatting preservation.
    """
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        with zipfile.ZipFile(file_path, 'r') as odt_zip:
            # Read the content.xml file which contains the text
            try:
                content_xml = odt_zip.read('content.xml')
                root = ET.fromstring(content_xml)
                
                # Extract text from all text nodes
                text_content = ""
                
                # Define ODT namespaces
                namespaces = {
                    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
                }
                
                # Find all paragraphs
                for paragraph in root.findall('.//text:p', namespaces):
                    para_text = ''.join(paragraph.itertext()).strip()
                    if para_text:
                        text_content += para_text + '\n\n'
                
                # Find all headings
                for heading in root.findall('.//text:h', namespaces):
                    level = heading.get('{urn:oasis:names:tc:opendocument:xmlns:text:1.0}outline-level', '1')
                    heading_text = ''.join(heading.itertext()).strip()
                    if heading_text:
                        text_content = text_content.replace(heading_text + '\n\n', f'<h{level}>{heading_text}</h{level}>\n\n')
                
                if text_content.strip():
                    # Convert to HTML with paragraph formatting
                    paragraphs = text_content.split('\n\n')
                    html_content = ''
                    
                    for paragraph in paragraphs:
                        paragraph = paragraph.strip()
                        if paragraph:
                            # Check if it's already a heading
                            if paragraph.startswith('<h') and paragraph.endswith('>'):
                                html_content += paragraph + '\n'
                            else:
                                # Clean up excessive whitespace
                                paragraph = re.sub(r'\s+', ' ', paragraph)
                                html_content += f'<p>{paragraph}</p>\n'
                    
                    return html_content
                
                return "<p>No readable content found in ODT file.</p>"
                
            except KeyError:
                return "<p>Invalid ODT file structure.</p>"
        
    except Exception as e:
        print(f"Error parsing ODT: {e}")
        return f"<p>Error reading ODT file: {str(e)}</p>"


def extract_from_google_docs(docs_url: str) -> str:
    """
    Extract text from Google Docs using the public sharing URL.
    
    Args:
        docs_url: The Google Docs sharing URL
        
    Returns:
        Extracted HTML content from the Google Doc
    """
    try:
        # Extract document ID from various Google Docs URL formats
        doc_id = extract_google_docs_id(docs_url)
        if not doc_id:
            return "<p>Invalid Google Docs URL. Please ensure the document is publicly accessible.</p>"
        
        # Use the export URL to get HTML version
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"
        
        # Make request to get the document
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(export_url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Parse the HTML content
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Find the main content div
        content_div = soup.find('div', {'id': 'contents'}) or soup.find('body')
        
        if not content_div:
            return "<p>Could not extract content from Google Docs. Please check the sharing permissions.</p>"
        
        # Clean up the HTML
        # Remove Google Docs specific elements
        for element in content_div.find_all(['script', 'style', 'meta', 'title']):
            element.decompose()
        
        # Remove Google Docs specific classes and IDs
        for element in content_div.find_all():
            if element.get('class'):
                element['class'] = []
            if element.get('id'):
                del element['id']
            if element.get('style'):
                del element['style']
        
        # Convert to clean HTML
        html_content = str(content_div)
        
        # Clean up extra whitespace and empty elements
        html_content = re.sub(r'<p[^>]*>\s*</p>', '', html_content)
        html_content = re.sub(r'<div[^>]*>\s*</div>', '', html_content)
        html_content = re.sub(r'\s+', ' ', html_content)
        
        return html_content if html_content.strip() else "<p>No readable content found in the Google Doc.</p>"
        
    except requests.exceptions.RequestException as e:
        return f"<p>Error accessing Google Docs: {str(e)}. Please ensure the document is publicly accessible.</p>"
    except Exception as e:
        print(f"Error parsing Google Docs: {e}")
        return f"<p>Error reading Google Docs: {str(e)}</p>"


def extract_google_docs_id(url: str) -> Optional[str]:
    """
    Extract the document ID from various Google Docs URL formats.
    
    Args:
        url: The Google Docs URL
        
    Returns:
        The document ID if found, None otherwise
    """
    try:
        # Common Google Docs URL patterns
        patterns = [
            r'/document/d/([a-zA-Z0-9-_]+)',  # Standard sharing URL
            r'id=([a-zA-Z0-9-_]+)',          # Some export URLs
            r'/d/([a-zA-Z0-9-_]+)/edit',     # Edit URLs
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    except Exception:
        return None


def is_google_docs_url(url: str) -> bool:
    """
    Check if a URL is a Google Docs URL.
    
    Args:
        url: The URL to check
        
    Returns:
        True if it's a Google Docs URL, False otherwise
    """
    try:
        parsed = urlparse(url)
        return (
            'docs.google.com' in parsed.netloc and 
            '/document/' in parsed.path
        )
    except Exception:
        return False


def validate_document_access(url: str) -> Tuple[bool, str]:
    """
    Validate that a Google Docs URL is accessible for import.
    
    Args:
        url: The Google Docs URL to validate
        
    Returns:
        Tuple of (is_valid, message)
    """
    try:
        if not is_google_docs_url(url):
            return False, "Not a valid Google Docs URL"
        
        doc_id = extract_google_docs_id(url)
        if not doc_id:
            return False, "Could not extract document ID from URL"
        
        # Test access to the document
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.head(export_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            return True, "Document is accessible"
        elif response.status_code == 403:
            return False, "Document is not publicly accessible. Please change sharing settings to 'Anyone with the link can view'"
        elif response.status_code == 404:
            return False, "Document not found. Please check the URL"
        else:
            return False, f"Unable to access document (Status: {response.status_code})"
            
    except requests.exceptions.Timeout:
        return False, "Request timed out. Please try again"
    except requests.exceptions.RequestException as e:
        return False, f"Network error: {str(e)}"
    except Exception as e:
        return False, f"Validation error: {str(e)}"


def identify_chapters_in_text(content: str) -> List[Dict[str, any]]:
    """
    Identify and extract chapters from document content.
    
    Args:
        content: The extracted text content from a document
        
    Returns:
        List of chapter dictionaries with title, content, and metadata
    """
    chapters = []
    
    # Clean the content first
    content = clean_content_for_chapter_detection(content)
    
    # Try different chapter detection patterns
    chapter_patterns = [
        # Pattern 1: Standard chapter headings
        r'(?i)(?:^|\n)\s*(?:chapter|ch\.?)\s+([0-9]+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)(?:\s*[:-]?\s*(.+?))?\s*(?:\n|$)',
        # Pattern 2: Part/Section headings  
        r'(?i)(?:^|\n)\s*(?:part|section)\s+([0-9]+|[ivxlcdm]+|one|two|three|four|five)(?:\s*[:-]?\s*(.+?))?\s*(?:\n|$)',
        # Pattern 3: Numbered sections
        r'(?i)(?:^|\n)\s*([0-9]+)\.\s*(.+?)\s*(?:\n|$)',
        # Pattern 4: All caps headings
        r'(?:^|\n)\s*([A-Z][A-Z\s,.:!?-]{10,80})\s*(?:\n|$)',
        # Pattern 5: HTML headings
        r'<h[1-6][^>]*>([^<]+)</h[1-6]>',
        # Pattern 6: Markdown headings
        r'(?:^|\n)#{1,6}\s*(.+?)\s*(?:\n|$)'
    ]
    
    detected_breaks = []
    
    # Find all potential chapter breaks
    for i, pattern in enumerate(chapter_patterns):
        matches = re.finditer(pattern, content, re.MULTILINE | re.IGNORECASE)
        for match in matches:
            chapter_info = {
                'position': match.start(),
                'end_position': match.end(),
                'pattern_type': i,
                'raw_match': match.group(0).strip(),
                'title': extract_chapter_title(match, i),
                'confidence': calculate_confidence(match, i, content)
            }
            detected_breaks.append(chapter_info)
    
    # Sort by position and filter by confidence
    detected_breaks.sort(key=lambda x: x['position'])
    filtered_breaks = filter_chapter_breaks(detected_breaks, content)
    
    if not filtered_breaks:
        # No chapters detected, return entire content as single chapter
        return [{
            'title': 'Chapter 1',
            'content': content.strip(),
            'order': 1,
            'word_count': len(content.split()),
            'confidence': 0.5,
            'detected_method': 'fallback'
        }]
    
    # Split content into chapters
    for i, break_info in enumerate(filtered_breaks):
        start_pos = break_info['position']
        end_pos = filtered_breaks[i + 1]['position'] if i + 1 < len(filtered_breaks) else len(content)
        
        chapter_content = content[start_pos:end_pos].strip()
        
        # Remove the title from content if it's at the beginning
        title = break_info['title']
        if chapter_content.startswith(break_info['raw_match']):
            chapter_content = chapter_content[len(break_info['raw_match']):].strip()
        
        if chapter_content:  # Only add non-empty chapters
            chapters.append({
                'title': title,
                'content': chapter_content,
                'order': i + 1,
                'word_count': len(chapter_content.split()),
                'confidence': break_info['confidence'],
                'detected_method': f'pattern_{break_info["pattern_type"]}'
            })
    
    return chapters if chapters else [{
        'title': 'Chapter 1',
        'content': content.strip(),
        'order': 1,
        'word_count': len(content.split()),
        'confidence': 0.5,
        'detected_method': 'fallback'
    }]


def clean_content_for_chapter_detection(content: str) -> str:
    """
    Clean content for better chapter detection.
    """
    # Remove HTML tags but preserve structure
    content = re.sub(r'<script.*?</script>', '', content, flags=re.DOTALL)
    content = re.sub(r'<style.*?</style>', '', content, flags=re.DOTALL)
    
    # Convert common HTML elements to text markers
    content = re.sub(r'<h([1-6])[^>]*>', r'\n\nH\1: ', content)
    content = re.sub(r'</h[1-6]>', '\n\n', content)
    content = re.sub(r'<p[^>]*>', '\n\n', content)
    content = re.sub(r'</p>', '', content)
    content = re.sub(r'<br[^>]*>', '\n', content)
    content = re.sub(r'<[^>]+>', '', content)  # Remove remaining HTML
    
    # Clean up whitespace
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = re.sub(r'[ \t]+', ' ', content)
    
    return content.strip()


def extract_chapter_title(match, pattern_type: int) -> str:
    """
    Extract a clean chapter title from regex match.
    """
    if pattern_type == 0:  # Chapter X
        number = match.group(1)
        title_part = match.group(2) if len(match.groups()) > 1 and match.group(2) else ''
        if title_part:
            return f"Chapter {number}: {title_part.strip()}"
        else:
            return f"Chapter {number}"
    
    elif pattern_type == 1:  # Part/Section X
        number = match.group(1)
        title_part = match.group(2) if len(match.groups()) > 1 and match.group(2) else ''
        section_type = 'Part' if 'part' in match.group(0).lower() else 'Section'
        if title_part:
            return f"{section_type} {number}: {title_part.strip()}"
        else:
            return f"{section_type} {number}"
    
    elif pattern_type == 2:  # Numbered (1. Title)
        number = match.group(1)
        title = match.group(2).strip() if match.group(2) else f"Section {number}"
        return title
    
    elif pattern_type == 3:  # All caps
        title = match.group(1).strip()
        # Convert to title case
        return title.title()
    
    elif pattern_type == 4:  # HTML headings
        return match.group(1).strip()
    
    elif pattern_type == 5:  # Markdown headings
        return match.group(1).strip()
    
    return "Untitled Chapter"


def calculate_confidence(match, pattern_type: int, content: str) -> float:
    """
    Calculate confidence score for chapter detection.
    """
    base_confidence = {
        0: 0.9,  # Chapter X - highest confidence
        1: 0.8,  # Part/Section X
        2: 0.7,  # Numbered sections
        3: 0.6,  # All caps
        4: 0.8,  # HTML headings
        5: 0.7   # Markdown headings
    }
    
    confidence = base_confidence.get(pattern_type, 0.5)
    
    # Adjust based on context
    match_text = match.group(0).lower()
    
    # Boost confidence for common chapter words
    if any(word in match_text for word in ['chapter', 'prologue', 'epilogue', 'introduction', 'conclusion']):
        confidence += 0.1
    
    # Reduce confidence for very short titles
    if len(match.group(0).strip()) < 5:
        confidence -= 0.2
    
    # Boost confidence if title is on its own line
    start_pos = max(0, match.start() - 10)
    end_pos = min(len(content), match.end() + 10)
    context = content[start_pos:end_pos]
    
    if '\n\n' in context[:10] and '\n\n' in context[-10:]:
        confidence += 0.1
    
    return max(0.1, min(1.0, confidence))


def filter_chapter_breaks(breaks: List[Dict], content: str) -> List[Dict]:
    """
    Filter and deduplicate chapter breaks.
    """
    if not breaks:
        return []
    
    # Remove duplicates (breaks very close to each other)
    filtered = []
    for break_info in breaks:
        # Check if too close to existing break
        too_close = False
        for existing in filtered:
            if abs(break_info['position'] - existing['position']) < 50:
                # Keep the one with higher confidence
                if break_info['confidence'] > existing['confidence']:
                    filtered.remove(existing)
                else:
                    too_close = True
                break
        
        if not too_close:
            filtered.append(break_info)
    
    # Filter by minimum confidence and content length
    final_breaks = []
    for i, break_info in enumerate(filtered):
        if break_info['confidence'] < 0.5:
            continue
            
        # Check if there's enough content for a chapter
        start_pos = break_info['position']
        end_pos = filtered[i + 1]['position'] if i + 1 < len(filtered) else len(content)
        
        chapter_content = content[start_pos:end_pos].strip()
        word_count = len(chapter_content.split())
        
        # Skip very short chapters (unless it's the last one)
        if word_count < 50 and i < len(filtered) - 1:
            continue
            
        final_breaks.append(break_info)
    
    return final_breaks


def extract_text_and_chapters_from_file(file_path: str, google_docs_url: Optional[str] = None) -> Dict[str, any]:
    """
    Extract text and automatically identify chapters from a document file.
    
    Args:
        file_path: Path to the file to parse
        google_docs_url: Optional Google Docs sharing URL
        
    Returns:
        Dictionary containing full_text, chapters list, and metadata
    """
    # First extract the text
    full_text = extract_text_from_file(file_path, google_docs_url)
    
    if not full_text or full_text.strip() == "":
        return {
            'full_text': '',
            'chapters': [],
            'chapter_count': 0,
            'total_words': 0,
            'detection_method': 'failed'
        }
    
    # Then identify chapters
    chapters = identify_chapters_in_text(full_text)
    
    return {
        'full_text': full_text,
        'chapters': chapters,
        'chapter_count': len(chapters),
        'total_words': sum(chapter['word_count'] for chapter in chapters),
        'detection_method': 'automatic',
        'average_confidence': sum(chapter['confidence'] for chapter in chapters) / len(chapters) if chapters else 0
    }
