from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
import requests

# AI Quote Extraction & Summarization Endpoint
@csrf_exempt
@login_required
def ai_quote_extraction(request):
    if request.method == 'POST':
        import json
        data = json.loads(request.body)
        chapter_text = data.get('chapter_text', '')
        if not chapter_text:
            return JsonResponse({'success': False, 'error': 'No chapter text provided.'}, status=400)

        # Call OpenAI API (or other AI provider)
        OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
        if not OPENAI_API_KEY:
            return JsonResponse({'success': False, 'error': 'OpenAI API key not configured.'}, status=500)

        prompt = (
            "Extract the most important quotes from the following chapter. "
            "For each quote, provide a brief summary explaining its relevance to the central story. "
            "Return a JSON array of objects with 'quote' and 'summary'.\n\nChapter:\n" + chapter_text
        )

        headers = {
            'Authorization': f'Bearer {OPENAI_API_KEY}',
            'Content-Type': 'application/json',
        }
        payload = {
            'model': 'gpt-3.5-turbo',
            'messages': [
                {'role': 'system', 'content': 'You are a helpful literary assistant.'},
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 800,
            'temperature': 0.7,
        }
        try:
            response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            ai_text = result['choices'][0]['message']['content']
            # Try to parse JSON from AI response
            import json as pyjson
            try:
                quotes = pyjson.loads(ai_text)
            except Exception:
                quotes = ai_text  # Fallback: return raw text
            return JsonResponse({'success': True, 'quotes': quotes})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    return JsonResponse({'error': 'POST required'}, status=405)
# Writing Stats API endpoint
from django.http import JsonResponse
from django.http import JsonResponse

def api_endpoint(request):
    data = {"message": "Hello from Django!", "date": "2025-08-21"}
    return JsonResponse(data)

def stats_api(request):
    # Replace these with real stats from your models
    data = {
        "wordCount": 12450,
        "streak": 7,
        "progress": 62
    }
    return JsonResponse(data)
from django.urls import path
from . import views

urlpatterns = [
    # ...existing code...
    path('api/stats/', views.stats_api, name='stats_api'),
]

urlpatterns = [
    # ...existing code...
    path('api/stats/', views.stats_api, name='stats_api'),
]
# Shortcut: /writer/editor/ → latest project editor or project create
from django.contrib.auth.decorators import login_required
@login_required
def latest_editor_shortcut(request):
    from .models import Project
    latest_project = Project.objects.filter(author=request.user).order_by('-updated_at').first()
    if latest_project:
        return redirect('writer:chapter_editor', project_id=latest_project.pk)
    else:
        return redirect('writer:project_create')
# TinyMCE Editor View (for both documents and chapters)
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
@login_required
def tinymce_editor(request):
    # For demo: store in session, or extend to DB as needed
    if request.method == 'POST':
        action = request.POST.get('action')
        title = request.POST.get('title', '')
        content = request.POST.get('content', '')
        if action == 'save':
            request.session['tinymce_editor_title'] = title
            request.session['tinymce_editor_content'] = content
            return JsonResponse({'status': 'success'})
        elif action == 'save_to_library':
            # Save as a Document for now
            doc = Document.objects.create(
                title=title or 'Untitled',
                content=content,
                author=request.user
            )
            return JsonResponse({'status': 'success', 'message': 'Saved to library!'})
        return JsonResponse({'status': 'error', 'message': 'Invalid action'})
    title = request.session.get('tinymce_editor_title', '')
    content = request.session.get('tinymce_editor_content', '')
    return render(request, 'writer/tinymce_editor.html', {'title': title, 'content': content})

from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy, reverse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.models import User
from django.db.models import Q, Max, F
from django.db import models, transaction
import re
from .models import (Document, Project, Chapter, AIAssistanceRequest, Character, 
                    ImportedDocument, ProjectCollaborator, WritingTheme, 
                    PersonalLibrary, WritingSession, CreativeNotebook, CreativeNode, NodeConnection, UserProfile)
from .forms import (DocumentForm, ProjectForm, ChapterForm, AIAssistanceForm, 
                   CharacterForm, ImportDocumentForm, CollaboratorForm)
import json
import random
import os
import mimetypes
import logging
from django.conf import settings
from .document_parser import extract_text_from_file  # , is_google_docs_url
import tempfile

# Simple Editor View
@login_required
def simple_editor(request):
    # For demo: store in session, or extend to DB as needed
    if request.method == 'POST':
        action = request.POST.get('action')
        title = request.POST.get('title', '')
        content = request.POST.get('content', '')
        if action == 'save':
            request.session['simple_editor_title'] = title
            request.session['simple_editor_content'] = content
            return JsonResponse({'status': 'success'})
        elif action == 'save_to_library':
            # Save as a Document for now
            doc = Document.objects.create(
                title=title or 'Untitled',
                content=content,
                author=request.user
            )
            return JsonResponse({'status': 'success', 'message': 'Saved to library!'})
        return JsonResponse({'status': 'error', 'message': 'Invalid action'})
    title = request.session.get('simple_editor_title', '')
    content = request.session.get('simple_editor_content', '')
    return render(request, 'writer/simple_editor.html', {'title': title, 'content': content})

# Editor shortcut view
def editor_shortcut(request):
    from .models import Project
    project = Project.objects.order_by('id').first()
    if project:
        return HttpResponseRedirect(reverse('writer:chapter_editor', args=[project.pk]))
    else:
        return HttpResponse('<div style="text-align:center;margin-top:3em;font-size:1.5em;">No projects found. Please create a project first.</div>')


# Personal Library Views
@login_required
def personal_library(request):
    try:
        library, created = PersonalLibrary.objects.get_or_create(user=request.user)
        user_projects = Project.objects.filter(author=request.user).select_related('author').prefetch_related('chapters').order_by('-updated_at')
        imported_docs = ImportedDocument.objects.filter(user=request.user).order_by('-import_date')

        latest_project = user_projects.first() if user_projects else None

        context = {
            'library': library,
            'projects': user_projects,
            'imported_documents': imported_docs,
            'latest_project': latest_project,
        }
        return render(request, 'writer/personal_library.html', context)
    except Exception as e:
        import logging
        logging.error(f"Error loading personal library for user {request.user}: {e}")
        print(f"Error loading personal library for user {request.user}: {e}")
        context = {
            'library': None,
            'projects': [],
            'imported_documents': [],
            'latest_project': None,
            'error_message': 'There was an issue loading your library. Please try again.'
        }
        return render(request, 'writer/personal_library.html', context)


# Project Views
class ProjectListView(LoginRequiredMixin, ListView):
    model = Project
    template_name = 'writer/project_list.html'
    context_object_name = 'projects'
    paginate_by = 10
    
    def get_queryset(self):
        return Project.objects.filter(
            Q(author=self.request.user) | Q(collaborators=self.request.user)
        ).distinct()


class ProjectDetailView(LoginRequiredMixin, DetailView):
    model = Project
    template_name = 'writer/project_detail.html'
    context_object_name = 'project'
    
    def get_queryset(self):
        return Project.objects.filter(
            Q(author=self.request.user) | Q(collaborators=self.request.user)
        ).distinct()
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['chapters'] = self.object.chapters.all().order_by('order')
        context['characters'] = self.object.characters.all()
        context['collaborators'] = self.object.projectcollaborator_set.all()
        context['writing_sessions'] = self.object.writing_sessions.filter(
            is_active=True
        )[:5]
        return context


class ProjectCreateView(LoginRequiredMixin, CreateView):
    model = Project
    form_class = ProjectForm
    template_name = 'writer/project_form.html'
    
    def form_valid(self, form):
        form.instance.author = self.request.user
        response = super().form_valid(form)
        
        # Add collaborators if specified
        collaborator_emails = form.cleaned_data.get('collaborator_emails')
        if collaborator_emails:
            self.add_collaborators(form.instance, collaborator_emails)
        
        # Create personal library if it doesn't exist
        PersonalLibrary.objects.get_or_create(user=self.request.user)
        
        return response
    
    def add_collaborators(self, project, email_string):
        emails = [email.strip() for email in email_string.split(',') if email.strip()]
        for email in emails:
            try:
                user = User.objects.get(email=email)
                ProjectCollaborator.objects.get_or_create(
                    project=project,
                    user=user,
                    defaults={'role': 'contributor'}
                )
            except User.DoesNotExist:
                messages.warning(self.request, f"User with email {email} not found.")


class ProjectUpdateView(LoginRequiredMixin, UpdateView):
    model = Project
    form_class = ProjectForm
    template_name = 'writer/project_form.html'
    
    def get_queryset(self):
        return Project.objects.filter(author=self.request.user)
    
    def get_success_url(self):
        return reverse('writer:project_detail', kwargs={'pk': self.object.pk})
    
    def form_valid(self, form):
        messages.success(self.request, f'Project "{form.instance.title}" updated successfully!')
        return super().form_valid(form)


class ProjectDeleteView(LoginRequiredMixin, DeleteView):
    model = Project
    template_name = 'writer/project_confirm_delete.html'
    success_url = reverse_lazy('writer:project_list')
    
    def get_queryset(self):
        return Project.objects.filter(author=self.request.user)


@login_required
def toggle_project_collaboration(request, project_id):
    """Toggle collaboration status for a project"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Only POST method allowed'}, status=405)
    
    try:
        project = get_object_or_404(Project, id=project_id, author=request.user)
        project.is_collaborative = not project.is_collaborative
        project.save()
        
        return JsonResponse({
            'success': True,
            'is_collaborative': project.is_collaborative,
            'message': f'Project is now {"collaborative" if project.is_collaborative else "private"}'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def get_project_share_link(request, project_id):
    """Generate share link for collaborative project"""
    try:
        project = get_object_or_404(Project, id=project_id)
        
        # Check if user has permission (author or collaborator with sharing rights)
        if not (project.author == request.user or 
                project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({
                'success': False,
                'error': 'You do not have permission to share this project'
            }, status=403)
        
        if not project.is_collaborative:
            return JsonResponse({
                'success': False,
                'error': 'This project is not collaborative'
            }, status=400)
        
        # Generate share URL
        share_url = request.build_absolute_uri(
            reverse('writer:project_collaborate', kwargs={'project_id': project.id})
        )
        
        return JsonResponse({
            'success': True,
            'share_url': share_url,
            'project_title': project.title
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@login_required
def join_collaborative_project(request, project_id):
    """Allow users to join a collaborative project via shared link"""
    try:
        project = get_object_or_404(Project, id=project_id, is_collaborative=True)
        
        if request.method == 'GET':
            # Show collaboration join page
            context = {
                'project': project,
            }
            return render(request, 'writer/project_collaborate.html', context)
        
        elif request.method == 'POST':
            # Check if user is already a collaborator
            if project.collaborators.filter(id=request.user.id).exists() or project.author == request.user:
                messages.info(request, 'You are already part of this project.')
                return redirect('writer:chapter_editor', project_id=project.id)
            
            # Add user as collaborator
            ProjectCollaborator.objects.create(
                project=project,
                user=request.user,
                role='contributor',
                can_edit=True,
                can_delete=False,
                can_invite_others=False
            )
            
            messages.success(request, f'You have successfully joined the collaborative project "{project.title}"!')
            return redirect('writer:chapter_editor', project_id=project.id)
        
    except Exception as e:
        messages.error(request, f'Error joining project: {str(e)}')
        return redirect('writer:personal_library')


# Character Views
class CharacterListView(LoginRequiredMixin, ListView):
    model = Character
    template_name = 'writer/character_list.html'
    context_object_name = 'characters'
    
    def get_queryset(self):
        project_id = self.kwargs.get('project_id')
        return Character.objects.filter(
            project_id=project_id,
            project__author=self.request.user
        )
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id, author=self.request.user)
        return context


class CharacterDetailView(LoginRequiredMixin, DetailView):
    model = Character
    template_name = 'writer/character_detail.html'
    context_object_name = 'character'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)


class CharacterCreateView(LoginRequiredMixin, CreateView):
    model = Character
    form_class = CharacterForm
    template_name = 'writer/character_form.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id, author=self.request.user)
        return context
    
    def form_valid(self, form):
        project_id = self.kwargs.get('project_id')
        form.instance.project = get_object_or_404(Project, id=project_id, author=self.request.user)
        return super().form_valid(form)
    
    def get_success_url(self):
        project_id = self.kwargs.get('project_id')
        return reverse_lazy('writer:character_list', kwargs={'project_id': project_id})


class CharacterUpdateView(LoginRequiredMixin, UpdateView):
    model = Character
    form_class = CharacterForm
    template_name = 'writer/character_form.html'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:character_list', kwargs={'project_id': self.object.project.pk})


class CharacterDeleteView(LoginRequiredMixin, DeleteView):
    model = Character
    template_name = 'writer/character_confirm_delete.html'
    
    def get_queryset(self):
        return Character.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:character_list', kwargs={'project_id': self.object.project.pk})


# Chapter Views (Updated)
class ChapterDetailView(LoginRequiredMixin, DetailView):
    model = Chapter
    template_name = 'writer/chapter_detail.html'
    context_object_name = 'chapter'
    
    def get_queryset(self):
        return Chapter.objects.filter(
            Q(project__author=self.request.user) | Q(project__collaborators=self.request.user)
        ).distinct()


class ChapterCreateView(LoginRequiredMixin, CreateView):
    model = Chapter
    form_class = ChapterForm
    template_name = 'writer/chapter_form.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        project_id = self.kwargs.get('project_id')
        context['project'] = get_object_or_404(Project, id=project_id)
        return context
    
    def form_valid(self, form):
        project_id = self.kwargs.get('project_id')
        project = get_object_or_404(Project, id=project_id)
        
        # Check permissions
        if not (project.author == self.request.user or 
                project.collaborators.filter(id=self.request.user.id).exists()):
            messages.error(self.request, "You don't have permission to add chapters to this project.")
            return redirect('writer:project_detail', pk=project.id)
        
        form.instance.project = project
        form.instance.last_edited_by = self.request.user
        
        if not form.instance.order:
            # Set order to be the next chapter
            last_chapter = Chapter.objects.filter(project=project).order_by('-order').first()
            form.instance.order = (last_chapter.order + 1) if last_chapter else 1
        
        return super().form_valid(form)


class ChapterUpdateView(LoginRequiredMixin, UpdateView):
    model = Chapter
    form_class = ChapterForm
    template_name = 'writer/chapter_form.html'
    
    def get_queryset(self):
        return Chapter.objects.filter(
            Q(project__author=self.request.user) | Q(project__collaborators=self.request.user)
        ).distinct()
    
    def form_valid(self, form):
        form.instance.last_edited_by = self.request.user
        return super().form_valid(form)


class ChapterDeleteView(LoginRequiredMixin, DeleteView):
    model = Chapter
    template_name = 'writer/chapter_confirm_delete.html'
    
    def get_queryset(self):
        return Chapter.objects.filter(project__author=self.request.user)
    
    def get_success_url(self):
        return reverse_lazy('writer:project_detail', kwargs={'pk': self.object.project.pk})


@login_required
def chapter_editor(request, project_id):
    """Integrated chapter management and editor view"""
    # Allow both author and collaborators
    project = get_object_or_404(Project, id=project_id)
    if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
        return JsonResponse({'status': 'error', 'message': 'Permission denied'}, status=403)
    all_chapters = Chapter.objects.filter(project=project).order_by('order')

    # Ensure at least one chapter exists for this project
    if not all_chapters.exists():
        default_chapter = Chapter.objects.create(
            title="Untitled Chapter",
            content="",
            project=project,
            last_edited_by=request.user,
            order=0
        )
        all_chapters = Chapter.objects.filter(project=project).order_by('order')

    # Get the current chapter being edited (default to first or create one)
    current_chapter_id = request.GET.get('chapter_id')
    current_chapter = None

    if current_chapter_id:
        current_chapter = get_object_or_404(Chapter, id=current_chapter_id, project=project)
    else:
        current_chapter = all_chapters.first()
    
    if request.method == 'POST':
        action = request.POST.get('action')

        # Map TinyMCE actions to chapter actions for compatibility
        if action == 'save':
            action = 'save_chapter'
        elif action == 'save_to_library':
            action = 'publish_chapter'

        if action == 'reorder':
            # Handle drag-and-drop reordering
            chapter_ids = request.POST.getlist('chapter_ids[]')
            for index, chapter_id in enumerate(chapter_ids):
                Chapter.objects.filter(id=chapter_id, project=project).update(order=index)
            return JsonResponse({'status': 'success'})
        
        elif action == 'save_chapter':
            # Handle chapter content saving
            chapter_id = request.POST.get('chapter_id')
            title = request.POST.get('title', 'Untitled Chapter')
            content = request.POST.get('content', '')
            
            if chapter_id:
                # Update existing chapter
                chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                chapter.title = title
                chapter.content = content
                chapter.save()
                # After save, fix duplicate orders if any
                chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
                for idx, ch in enumerate(chapters):
                    if ch.order != idx:
                        ch.order = idx
                        ch.save(update_fields=['order'])
                return JsonResponse({
                    'status': 'success',
                    'word_count': chapter.word_count,
                    'message': 'Chapter saved successfully'
                })
            else:
                # Create new chapter if none exists
                # Always assign next available order
                used_orders = set(Chapter.objects.filter(project=project).values_list('order', flat=True))
                order = 0
                while order in used_orders:
                    order += 1
                chapter = Chapter.objects.create(
                    title=title,
                    content=content,
                    project=project,
                    order=order
                )
                # After create, fix duplicate orders if any
                chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
                for idx, ch in enumerate(chapters):
                    if ch.order != idx:
                        ch.order = idx
                        ch.save(update_fields=['order'])
                return JsonResponse({
                    'status': 'success',
                    'chapter_id': chapter.id,
                    'word_count': chapter.word_count,
                    'message': 'New chapter created and saved successfully',
                    'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={chapter.id}'
                })
        
        elif action == 'create_chapter':
            # Handle creating new chapter
            title = request.POST.get('title', 'Untitled Chapter')
            heading_level = request.POST.get('heading_level', 'h2')
            
            # Get the next order number
            max_order = Chapter.objects.filter(project=project).aggregate(
                max_order=models.Max('order')
            )['max_order'] or -1
            
            # Create heading content based on level
            heading_tag = f"<{heading_level}>{title}</{heading_level}>"
            content = f"{heading_tag}<p></p>"
            
            chapter = Chapter.objects.create(
                title=title,
                content=content,
                project=project,
                order=max_order + 1
            )
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': chapter.id,
                'chapter_title': chapter.title,
                'chapter_order': chapter.order + 1,
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={chapter.id}'
            })
        
        elif action == 'delete_chapter':
            # Handle chapter deletion
            chapter_id = request.POST.get('chapter_id')
            chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
            chapter.delete()
            
            # Reorder remaining chapters
            remaining_chapters = Chapter.objects.filter(project=project).order_by('order')
            for index, ch in enumerate(remaining_chapters):
                ch.order = index
                ch.save()
            
            return JsonResponse({'status': 'success'})
        
        elif action == 'publish_chapter':
            # Handle publishing chapter to library
            chapter_id = request.POST.get('chapter_id')
            title = request.POST.get('title', 'Untitled Chapter')
            content = request.POST.get('content', '')
            if chapter_id:
                chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                chapter.title = title
                chapter.content = content
                chapter.is_published = True
                chapter.save()
                return JsonResponse({
                    'status': 'success',
                    'message': 'Chapter published successfully',
                    'word_count': chapter.word_count
                })
            else:
                return JsonResponse({'status': 'error', 'message': 'No chapter selected to publish.'})
        elif action == 'split_long_chapter':
            # Handle automatic chapter splitting when content is too long
            try:
                import json
                data = json.loads(request.body)
                chapter_id = data.get('chapter_id')
                first_part = data.get('first_part')
                remaining_part = data.get('remaining_part')
                
                current_chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
                
                # Update current chapter with first part
                current_chapter.content = first_part
                current_chapter.save()
                
                # Create new chapter with remaining content
                new_chapter = Chapter.objects.create(
                    title="Untitled Chapter (Auto-split)",
                    content=remaining_part,
                    project=project,
                    order=current_chapter.order + 1
                )
                
                # Reorder subsequent chapters
                subsequent_chapters = Chapter.objects.filter(
                    project=project, 
                    order__gt=current_chapter.order
                ).exclude(id=new_chapter.id)
                
                for chapter in subsequent_chapters:
                    chapter.order += 1
                    chapter.save()
                
                return JsonResponse({
                    'success': True,
                    'new_chapter_title': new_chapter.title,
                    'new_chapter_id': new_chapter.id
                })
                
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)})
        
        elif action == 'create_chapter_from_heading':
            # Handle creating new chapter from H2 heading
            title = request.POST.get('title', 'Untitled Chapter')
            current_chapter_id = request.POST.get('current_chapter_id')
            content_before = request.POST.get('content_before', '')
            content_after = request.POST.get('content_after', '')
            
            # Update current chapter content (remove the heading and content after)
            if current_chapter_id:
                current_chapter = get_object_or_404(Chapter, id=current_chapter_id, project=project)
                current_chapter.content = content_before
                current_chapter.save()
            
            # Get the next order number (insert after current chapter)
            current_order = current_chapter.order if current_chapter_id else -1
            
            # Shift all chapters after current chapter
            chapters_to_shift = Chapter.objects.filter(project=project, order__gt=current_order)
            for ch in chapters_to_shift:
                ch.order += 1
                ch.save()
            
            # Create new chapter
            new_chapter = Chapter.objects.create(
                title=title,
                content=f"<h2>{title}</h2>{content_after}",
                project=project,
                order=current_order + 1
            )
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': new_chapter.id,
                'chapter_title': new_chapter.title,
                'chapter_order': new_chapter.order + 1,
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={new_chapter.id}'
            })
    
    # Handle create_from_selection action
    if request.method == 'POST' and request.POST.get('action') == 'create_from_selection':
        try:
            selected_text = request.POST.get('selected_text', '').strip()
            new_title = request.POST.get('new_title', '').strip()
            remove_from_current = request.POST.get('remove_from_current', '').lower() == 'true'
            
            if not selected_text:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No text selected'
                })
            
            if not new_title:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Chapter title is required'
                })
            
            # Get the highest order number for new chapter positioning
            current_order = 0
            if current_chapter:
                current_order = current_chapter.order
            
            # Create new chapter with selected text
            new_chapter = Chapter.objects.create(
                title=new_title,
                content=selected_text,
                project=project,
                order=current_order + 1
            )
            
            # Update orders of subsequent chapters
            Chapter.objects.filter(
                project=project,
                order__gt=current_order
            ).exclude(id=new_chapter.id).update(order=F('order') + 1)
            
            # Optionally remove selected text from current chapter
            if remove_from_current and current_chapter:
                # Remove the selected text from current chapter's content
                current_content = current_chapter.content or ''
                if selected_text in current_content:
                    updated_content = current_content.replace(selected_text, '', 1)
                    current_chapter.content = updated_content
                    current_chapter.save()
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': new_chapter.id,
                'chapter_title': new_chapter.title,
                'chapter_order': new_chapter.order,
                'message': f'Chapter "{new_title}" created successfully',
                'redirect_url': f'/writer/projects/{project_id}/editor/?chapter_id={new_chapter.id}'
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Error creating chapter: {str(e)}'
            })
    
    # Handle direct_upload_to_editor action
    if request.method == 'POST' and request.POST.get('action') == 'direct_upload_to_editor':
        try:
            if 'file' not in request.FILES:
                return JsonResponse({
                    'status': 'error',
                    'message': 'No file selected'
                })
            
            uploaded_file = request.FILES['file']
            insert_position = request.POST.get('insert_position', 'end')  # 'start', 'end', 'cursor'
            
            # Save the uploaded file temporarily
            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_uploads')
            os.makedirs(temp_dir, exist_ok=True)
            
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            # Save the file
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Extract text content
            content = extract_text_from_file(temp_file_path)
            
            if not content:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                
                return JsonResponse({
                    'status': 'error',
                    'message': 'Could not extract text from the uploaded file. Please try a different format.'
                })
            
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            
            # Return the extracted content for insertion
            return JsonResponse({
                'status': 'success',
                'content': content,
                'filename': uploaded_file.name,
                'message': f'File "{uploaded_file.name}" uploaded successfully'
            })
            
        except Exception as e:
            # Clean up temporary file on error
            if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing file: {str(e)}'
            })
    
    context = {
        'project': project,
        'all_chapters': all_chapters,
        'current_chapter': current_chapter,
        'user': request.user,
        'document': current_chapter,  # Pass chapter as document for compatibility
    }
    return render(request, 'writer/google_docs_editor.html', context)

@login_required
def upload_test(request):
    """Simple upload test page."""
    return render(request, 'writer/upload_test.html')



# Document Views (Updated)
class DocumentListView(LoginRequiredMixin, ListView):
    model = Document
    template_name = 'writer/document_list.html'
    context_object_name = 'documents'
    paginate_by = 10
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentDetailView(LoginRequiredMixin, DetailView):
    model = Document
    template_name = 'writer/document_detail.html'
    context_object_name = 'document'
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentCreateView(LoginRequiredMixin, CreateView):
    model = Document
    form_class = DocumentForm
    template_name = 'writer/document_form.html'
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs
    
    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)


class DocumentUpdateView(LoginRequiredMixin, UpdateView):
    model = Document
    form_class = DocumentForm
    template_name = 'writer/document_form.html'
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


class DocumentDeleteView(LoginRequiredMixin, DeleteView):
    model = Document
    template_name = 'writer/document_confirm_delete.html'
    success_url = reverse_lazy('writer:document_list')
    
    def get_queryset(self):
        return Document.objects.filter(author=self.request.user)


# Document Import Views
@login_required
def import_document(request):
    if request.method == 'POST':
        form = ImportDocumentForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            imported_doc = form.save(commit=False)
            imported_doc.user = request.user
            
            # Handle Google Docs URL
            google_docs_url = form.cleaned_data.get('google_docs_url')
            if google_docs_url:
                # Extract content from Google Docs
                imported_doc.google_docs_url = google_docs_url
                extracted_content = extract_text_from_file(None, google_docs_url=google_docs_url)
                imported_doc.import_type = 'google_docs'
                imported_doc.file_size = 0  # No file uploaded
                
                # Save without file
                imported_doc.save()
                imported_doc.extracted_content = extracted_content
            else:
                # Handle file upload
                imported_doc.file_size = imported_doc.original_file.size
                
                # Determine file type
                file_name = imported_doc.original_file.name.lower()
                if file_name.endswith('.pdf'):
                    imported_doc.import_type = 'pdf'
                elif file_name.endswith(('.docx', '.doc')):
                    imported_doc.import_type = 'docx'
                elif file_name.endswith('.txt'):
                    imported_doc.import_type = 'txt'
                elif file_name.endswith('.rtf'):
                    imported_doc.import_type = 'rtf'
                elif file_name.endswith('.odt'):
                    imported_doc.import_type = 'odt'
                elif file_name.endswith(('.html', '.htm')):
                    imported_doc.import_type = 'html'
                
                # Save the file first to get the path
                imported_doc.save()
                
                # Extract text content from the saved file
                temp_file_path = imported_doc.original_file.path
                extracted_content = extract_text_from_file(temp_file_path)
            
            # Set target project if selected, otherwise create a new project
            target_project = form.cleaned_data.get('target_project')
            if not target_project:
                # Auto-create a new project based on the imported document
                project_title = imported_doc.title
                if len(project_title) > 50:  # Ensure title isn't too long
                    project_title = project_title[:47] + "..."
                
                # Create new project
                target_project = Project.objects.create(
                    title=project_title,
                    description=f"Project created from imported document: {imported_doc.title}",
                    author=request.user,
                    target_word_count=10000,  # Default word count
                    genre="General",  # Default genre
                )
                
                imported_doc.project = target_project
            
            # Apply formatting options
            if not form.cleaned_data.get('preserve_formatting', True):
                # Strip HTML formatting if user doesn't want it preserved
                import re
                extracted_content = re.sub(r'<[^>]+>', '', extracted_content)
                # Convert back to simple paragraphs
                paragraphs = [p.strip() for p in extracted_content.split('\n') if p.strip()]
                extracted_content = ''.join(f'<p>{p}</p>' for p in paragraphs)
            
            imported_doc.extracted_content = extracted_content
            
            # Save again with extracted content
            imported_doc.save()
            
            # Auto-create chapters if requested, or create a single chapter if no target project was selected
            if form.cleaned_data.get('auto_create_chapters', False) and target_project:
                create_chapters_from_content(imported_doc.extracted_content, target_project, imported_doc.title)
                messages.success(request, f"Document '{imported_doc.title}' imported and chapters created successfully!")
            elif target_project and not form.cleaned_data.get('auto_create_chapters', False):
                # Create a single chapter with all the content
                Chapter.objects.create(
                    title=imported_doc.title,
                    content=imported_doc.extracted_content,
                    project=target_project,
                    order=1
                )
                messages.success(request, f"Document '{imported_doc.title}' imported successfully and converted to a new project!")
            else:
                messages.success(request, f"Document '{imported_doc.title}' imported successfully!")
            
            return redirect('writer:import_detail', pk=imported_doc.pk)
    else:
        form = ImportDocumentForm(user=request.user)
    
    return render(request, 'writer/import_form.html', {'form': form})


def create_chapters_from_content(content, project, base_title):
    """Create chapters from content based on HTML headings"""
    import re
    from bs4 import BeautifulSoup
    
    try:
        soup = BeautifulSoup(content, 'html.parser')
        
        # Find all headings (h1, h2, h3)
        headings = soup.find_all(['h1', 'h2', 'h3'])
        
        if not headings:
            # No headings found, create a single chapter
            Chapter.objects.create(
                title=base_title,
                content=content,
                project=project,
                order=1
            )
            return
        
        current_content = ""
        chapter_order = 1
        
        for i, heading in enumerate(headings):
            # Get content between this heading and the next
            current_element = heading.next_sibling
            chapter_content = f"<{heading.name}>{heading.get_text()}</{heading.name}>"
            
            while current_element and current_element != headings[i + 1] if i + 1 < len(headings) else True:
                if hasattr(current_element, 'name') and current_element.name:
                    chapter_content += str(current_element)
                elif isinstance(current_element, str) and current_element.strip():
                    chapter_content += current_element
                current_element = current_element.next_sibling
                if i + 1 < len(headings) and current_element == headings[i + 1]:
                    break
            
            # Create chapter
            Chapter.objects.create(
                title=heading.get_text().strip() or f"{base_title} - Chapter {chapter_order}",
                content=chapter_content,
                project=project,
                order=chapter_order
            )
            chapter_order += 1
            
    except Exception as e:
        # Fallback: create single chapter
        Chapter.objects.create(
            title=base_title,
            content=content,
            project=project,
            order=1
        )


@login_required
def import_detail(request, pk):
    imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
    return render(request, 'writer/import_detail.html', {'imported_doc': imported_doc})


@login_required
def convert_import_to_project(request, pk):
    """Convert an imported document to a new project with chapters"""
    if request.method == 'POST':
        imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
        
        # Create a new project
        project = Project.objects.create(
            title=imported_doc.title,
            description=f"Project created from imported document: {imported_doc.title}",
            author=request.user,
            genre='Other'  # Default genre
        )
        
        # Create a chapter with the imported content
        chapter = Chapter.objects.create(
            title=f"Chapter 1: {imported_doc.title}",
            content=imported_doc.extracted_content or "No content available",
            project=project,
            order=1
        )
        
        messages.success(request, f"Successfully converted '{imported_doc.title}' to project '{project.title}'!")
        return JsonResponse({
            'status': 'success',
            'project_id': project.id,
            'redirect_url': reverse('writer:chapter_editor', kwargs={'project_id': project.id})
        })
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


@login_required
def add_import_to_project(request, pk, project_id):
    """Add imported document content as a new chapter to an existing project"""
    if request.method == 'POST':
        imported_doc = get_object_or_404(ImportedDocument, pk=pk, user=request.user)
        project = get_object_or_404(Project, id=project_id, author=request.user)
        
        # Get the next order for the new chapter
        last_chapter = Chapter.objects.filter(project=project).order_by('-order').first()
        next_order = (last_chapter.order + 1) if last_chapter else 1
        
        # Create a new chapter with the imported content
        chapter = Chapter.objects.create(
            title=f"Chapter {next_order}: {imported_doc.title}",
            content=imported_doc.extracted_content or "No content available",
            project=project,
            order=next_order
        )
        
        messages.success(request, f"Successfully added '{imported_doc.title}' as a new chapter to '{project.title}'!")
        return JsonResponse({
            'status': 'success',
            'project_id': project.id,
            'chapter_id': chapter.id,
            'redirect_url': reverse('writer:chapter_editor', kwargs={'project_id': project.id}) + f'?chapter_id={chapter.id}'
        })
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def projects_api_list(request):
    """API endpoint to list user's projects"""
    # Check if we should only return bookshelf projects
    bookshelf_only = request.GET.get('bookshelf_only', 'false').lower() == 'true'
    
    if bookshelf_only:
        projects = Project.objects.filter(author=request.user, show_on_dashboard=True).order_by('-updated_at')
    else:
        projects = Project.objects.filter(author=request.user).order_by('-updated_at')
    
    projects_data = [
        {
            'id': project.id,
            'title': project.title,
            'description': project.description or '',
            'updated_at': project.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            'created_at': project.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'chapter_count': project.chapters.count(),
            'word_count': project.total_word_count,
            'genre': project.genre or '',
            'target_word_count': project.target_word_count,
            'show_on_dashboard': project.show_on_dashboard,
            'progress_percentage': round(project.progress_percentage, 1),
        }
        for project in projects
    ]
    return Response({'projects': projects_data})


@login_required
@csrf_exempt
def chapters_api(request, project_id):
    """API endpoint for chapter operations"""
    try:
        project = Project.objects.get(id=project_id)
        # Check if user has access to this project
        if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({'error': 'Permission denied'}, status=403)
    except Project.DoesNotExist:
        return JsonResponse({'error': 'Project not found'}, status=404)
    
    if request.method == 'GET':
        chapters = Chapter.objects.filter(project=project).order_by('order', 'id')
        chapters_data = [
            {
                'id': chapter.id,
                'title': chapter.title,
                'content': chapter.content or '',
                'word_count': chapter.word_count,
                'order': chapter.order,
                'created_at': chapter.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': chapter.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            }
            for chapter in chapters
        ]
        return JsonResponse({'chapters': chapters_data})
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            chapter = Chapter.objects.create(
                title=data.get('title', 'New Chapter'),
                content=data.get('content', ''),
                project=project,
                order=data.get('order', 0)
            )
            return JsonResponse({
                'success': True,
                'chapter': {
                    'id': chapter.id,
                    'title': chapter.title,
                    'content': chapter.content,
                    'word_count': chapter.word_count,
                    'order': chapter.order,
                }
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)


@api_view(['GET', 'PUT'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def chapter_api(request, chapter_id):
    """API endpoint for individual chapter operations"""
    try:
        chapter = Chapter.objects.get(id=chapter_id)
        # Check if user has access to this chapter's project
        if not (chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists()):
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)
    except Chapter.DoesNotExist:
        return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)
    
    if request.method == 'GET':
        return Response({
            'id': chapter.id,
            'title': chapter.title,
            'content': chapter.content or '',
            'word_count': chapter.word_count,
            'order': chapter.order,
            'project_id': chapter.project.id,
            'project_title': chapter.project.title,
            'created_at': chapter.created_at.isoformat(),
            'updated_at': chapter.updated_at.isoformat(),
        })
    
    elif request.method == 'PUT':
        try:
            # Handle both JSON data from DRF and direct request.data
            data = request.data if hasattr(request, 'data') else json.loads(request.body)
            if 'title' in data:
                chapter.title = data['title']
            if 'content' in data:
                chapter.content = data['content']
            if 'order' in data:
                chapter.order = data['order']
            
            chapter.save()
            
            return Response({
                'success': True,
                'chapter': {
                    'id': chapter.id,
                    'title': chapter.title,
                    'content': chapter.content,
                    'word_count': chapter.word_count,
                }
            })
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class ImportedDocumentDeleteView(LoginRequiredMixin, DeleteView):
    model = ImportedDocument
    template_name = 'writer/import_confirm_delete.html'
    success_url = reverse_lazy('writer:personal_library')
    
    def get_queryset(self):
        return ImportedDocument.objects.filter(user=self.request.user)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['object_type'] = 'Imported Document'
        return context


# AJAX Views
@login_required
@require_http_methods(["POST"])
def ajax_create_chapter(request):
    try:
        import json
        data = json.loads(request.body)
        
        project_id = data.get('project_id')
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        
        if not all([project_id, title, content]):
            return JsonResponse({'success': False, 'error': 'Missing required fields'})
        
        # Get the project and verify ownership
        try:
            project = Project.objects.get(id=project_id)
            # Check if user has access to this project
            if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'success': False, 'error': 'Permission denied'})
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        
        # Create the chapter
        chapter = Chapter.objects.create(
            project=project,
            title=title,
            content=content,
            order=project.chapters.count() + 1
        )
        
        # Update project word count
        project.update_word_count()
        
        return JsonResponse({
            'success': True, 
            'chapter_id': chapter.id,
            'message': f'Chapter "{title}" created successfully!'
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def dashboard(request):
    recent_documents = Document.objects.filter(author=request.user)[:5]
    
    total_documents = Document.objects.filter(author=request.user).count()
    total_projects = Project.objects.filter(
        Q(author=request.user) | Q(collaborators=request.user)
    ).distinct().count()
    
    # Calculate comprehensive writing statistics
    user_projects = Project.objects.filter(author=request.user).select_related('author')
    
    # Total words from both documents and projects
    doc_words = sum(doc.word_count for doc in Document.objects.filter(author=request.user))
    project_words = sum(project.total_word_count for project in user_projects)
    total_words = doc_words + project_words
    
    # Get user's characters
    user_characters = Character.objects.filter(
        project__in=Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct()
    ).count()
    
    # Create shelves with max 7 books each
    shelves = []
    projects_list = list(user_projects)
    for i in range(0, len(projects_list), 7):
        shelf = projects_list[i:i+7]
        shelves.append(shelf)
    
    # Get user profile for theme preferences
    profile = UserProfile.get_or_create_for_user(request.user)
    
    context = {
        'recent_documents': recent_documents,
        'total_documents': total_documents,
        'total_projects': total_projects,
        'total_words': total_words,
        'total_characters': user_characters,
        'user_projects': user_projects,
        'shelves': shelves,
        'show_bookshelf': True,
        'profile': profile,
    }
    return render(request, 'writer/dashboard.html', context)


# Format & Export Views
@login_required
def format_page(request):
    """Format page for styling and export options"""
    user_projects = Project.objects.filter(
        Q(author=request.user) | Q(collaborators=request.user)
    ).distinct()
    
    context = {
        'user_projects': user_projects,
    }
    return render(request, 'writer/format_page.html', context)


@login_required
def export_project(request, project_id, format_type):
    """Export project in various formats"""
    project = get_object_or_404(Project, id=project_id)
    
    # Check permissions
    if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
        messages.error(request, "You don't have permission to export this project.")
        return redirect('writer:format_page')
    
    # Get all chapters in order
    chapters = Chapter.objects.filter(project=project).order_by('order')
    
    if format_type == 'pdf':
        return export_as_pdf(project, chapters)
    elif format_type == 'epub':
        return export_as_epub(project, chapters)
    elif format_type == 'docx':
        return export_as_docx(project, chapters)
    elif format_type == 'gdoc':
        return export_as_google_doc(project, chapters)
    else:
        messages.error(request, "Unsupported export format.")
        return redirect('writer:format_page')


def export_as_pdf(project, chapters):
    """Export project as PDF"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    import io
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title page
    title = Paragraph(f"<para align=center><font size=24><b>{project.title}</b></font></para>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 72))
    
    if project.description:
        desc = Paragraph(f"<para align=center><i>{project.description}</i></para>", styles['Normal'])
        story.append(desc)
        story.append(Spacer(1, 36))
    
    author = Paragraph(f"<para align=center>by {project.author.get_full_name() or project.author.username}</para>", styles['Normal'])
    story.append(author)
    story.append(Spacer(1, 72))
    
    # Chapters
    for chapter in chapters:
        if chapter.title:
            chapter_title = Paragraph(f"<para><font size=18><b>{chapter.title}</b></font></para>", styles['Heading1'])
            story.append(chapter_title)
            story.append(Spacer(1, 12))
        
        if chapter.content:
            # Clean HTML content for PDF
            import re
            clean_content = re.sub('<[^<]+?>', '', chapter.content)
            content = Paragraph(clean_content, styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{project.title}.pdf"'
    return response


def export_as_epub(project, chapters):
    """Export project as EPUB"""
    from ebooklib import epub
    import io
    
    book = epub.EpubBook()
    book.set_identifier(f'project_{project.id}')
    book.set_title(project.title)
    book.set_language('en')
    book.add_author(project.author.get_full_name() or project.author.username)
    
    if project.description:
        book.add_metadata('DC', 'description', project.description)
    
    # Create chapters
    epub_chapters = []
    for i, chapter in enumerate(chapters):
        epub_chapter = epub.EpubHtml(
            title=chapter.title or f'Chapter {i+1}',
            file_name=f'chapter_{i+1}.xhtml',
            lang='en'
        )
        epub_chapter.content = f'''
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head><title>{chapter.title or f'Chapter {i+1}'}</title></head>
        <body>
        <h1>{chapter.title or f'Chapter {i+1}'}</h1>
        {chapter.content or ''}
        </body>
        </html>
        '''
        book.add_item(epub_chapter)
        epub_chapters.append(epub_chapter)
    
    # Create table of contents
    book.toc = epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Create spine
    book.spine = ['nav'] + epub_chapters
    
    # Generate EPUB
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/epub+zip')
    response['Content-Disposition'] = f'attachment; filename="{project.title}.epub"'
    return response


def export_as_docx(project, chapters):
    """Export project as Word document"""
    from docx import Document
    from docx.shared import Inches
    import io
    
    doc = Document()
    
    # Title page
    title = doc.add_heading(project.title, 0)
    title.alignment = 1  # Center alignment
    
    if project.description:
        desc_para = doc.add_paragraph(project.description)
        desc_para.alignment = 1
    
    author_para = doc.add_paragraph(f"by {project.author.get_full_name() or project.author.username}")
    author_para.alignment = 1
    
    doc.add_page_break()
    
    # Chapters
    for chapter in chapters:
        if chapter.title:
            doc.add_heading(chapter.title, 1)
        
        if chapter.content:
            # Clean HTML content
            import re
            clean_content = re.sub('<[^<]+?>', '', chapter.content)
            doc.add_paragraph(clean_content)
        
        doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{project.title}.docx"'
    return response


def export_as_google_doc(project, chapters):
    """Export project as Google Doc (HTML format for import)"""
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{project.title}</title>
        <meta charset="utf-8">
        <style>
            body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 1in; }}
            h1 {{ text-align: center; page-break-before: always; }}
            h2 {{ margin-top: 2em; }}
            .title-page {{ text-align: center; margin-bottom: 3em; }}
            .chapter {{ page-break-before: always; }}
        </style>
    </head>
    <body>
        <div class="title-page">
            <h1>{project.title}</h1>
            {f'<p><i>{project.description}</i></p>' if project.description else ''}
            <p>by {project.author.get_full_name() or project.author.username}</p>
        </div>
    """
    
    for chapter in chapters:
        html_content += f'<div class="chapter">'
        if chapter.title:
            html_content += f'<h2>{chapter.title}</h2>'
        if chapter.content:
            html_content += chapter.content
        html_content += '</div>'
    
    html_content += """
    </body>
    </html>
    """
    
    response = HttpResponse(html_content, content_type='text/html')
    response['Content-Disposition'] = f'attachment; filename="{project.title}_for_google_docs.html"'
    return response


# AI Assistant Views (Enhanced)
@login_required
def ai_assistant(request):
    if request.method == 'POST':
        form = AIAssistanceForm(request.POST)
        if form.is_valid():
            assistance = form.save(commit=False)
            assistance.user = request.user
            
            # Generate AI response
            assistance.response = generate_ai_response(
                assistance.content, 
                assistance.assistance_type, 
                assistance.prompt
            )
            assistance.save()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'response': assistance.response
                })
            else:
                messages.success(request, 'AI assistance generated successfully!')
                return redirect('writer:ai_assistant')
    else:
        form = AIAssistanceForm()
    
    recent_requests = AIAssistanceRequest.objects.filter(user=request.user)[:10]
    context = {
        'form': form,
        'recent_requests': recent_requests
    }
    return render(request, 'writer/ai_assistant.html', context)


def generate_ai_response(content, assistance_type, prompt):
    """Enhanced AI response generator with character and plot development"""
    responses = {
        'brainstorm': [
            "Here are some creative ideas to expand your story:\n\n• What if your protagonist discovers a hidden secret about their past?\n• Consider adding a subplot involving a secondary character\n• Explore the consequences of your main character's decisions\n• Add unexpected obstacles or complications\n• Introduce a mentor figure who challenges your protagonist's beliefs",
            "To develop this further, consider:\n\n• Character motivations and backstory\n• Setting details that affect the plot\n• Potential conflicts between characters\n• Themes you want to explore\n• How the environment shapes your characters",
        ],
        'character': [
            "Character development suggestions:\n\n• Give your character a specific fear or weakness\n• Create a contradiction in their personality\n• Develop their unique voice and speech patterns\n• Consider their relationships with family and friends\n• What drives them? What do they want most?\n• Add physical quirks or habits that reveal personality",
            "To make this character more compelling:\n\n• Create internal conflict that mirrors external conflict\n• Give them a secret they're hiding\n• Develop their backstory - what shaped them?\n• Consider how they've changed throughout the story\n• What would they never do? What would force them to do it?",
        ],
        'dialogue': [
            "Dialogue enhancement tips:\n\n• Each character should have a distinct voice\n• Use subtext - characters rarely say exactly what they mean\n• Add action beats between dialogue\n• Show character relationships through how they speak to each other\n• Cut unnecessary dialogue tags\n• Make every line serve a purpose",
            "To improve this dialogue:\n\n• Add more conflict or tension\n• Use contractions for natural speech\n• Show character emotions through word choice\n• Consider what characters aren't saying\n• Add interruptions and overlapping speech for realism",
        ],
        'plot': [
            "Plot development ideas:\n\n• Increase the stakes - what happens if your protagonist fails?\n• Add a ticking clock element\n• Create obstacles that force character growth\n• Consider the three-act structure\n• Plant seeds early that pay off later\n• Make sure each scene advances plot or character",
            "To strengthen your plot:\n\n• Give your antagonist clear, understandable motivations\n• Create moments where victory seems impossible\n• Add plot twists that feel inevitable in hindsight\n• Make sure cause and effect are clear\n• Consider multiple storylines that intersect",
        ],
        'edit': [
            "Here are some suggestions to improve your writing:\n\n• Consider varying your sentence structure for better flow\n• Some paragraphs could be shortened for better pacing\n• Look for opportunities to show rather than tell\n• Strong dialogue! Consider adding more action beats between speeches\n• Watch for repeated words or phrases",
            "Editorial suggestions:\n\n• The opening could be stronger with more immediate action\n• Character descriptions are vivid and engaging\n• Consider tightening some exposition\n• The dialogue feels natural and authentic\n• Look for places to add sensory details",
        ],
        'continue': [
            "Here's a possible continuation:\n\nAs the door creaked open, Sarah held her breath. The hallway beyond was darker than she'd expected, and the musty smell of old books filled her nostrils. Each step forward felt like crossing into another world, where the familiar rules no longer applied...",
            "To continue this scene, you might consider:\n\nHaving your character notice something unexpected, introduce a new character, or reveal important information that changes everything. The tension you've built is perfect for a major revelation or plot twist.",
        ],
        'rewrite': [
            "Here's a rewritten version with improved flow:\n\n[Your original text would be revised here with better sentence structure, clearer descriptions, and enhanced dialogue while maintaining your unique voice and style.]",
            "I've rewritten this section to be more concise and impactful:\n\n[The rewritten version would focus on stronger verbs, clearer imagery, and better pacing while preserving the essence of your original writing.]",
        ],
        'grammar': [
            "Grammar and style check complete:\n\n• Fixed several comma splices\n• Corrected subject-verb agreement in paragraph 2\n• Suggested alternatives for repetitive word usage\n• Overall, your writing is clear and well-structured!\n• Consider breaking up some longer sentences",
            "Grammar review:\n\n✓ Sentence structure is generally good\n• Watch for run-on sentences in longer paragraphs\n• Consider breaking up some complex sentences\n• Punctuation is mostly correct\n• Strong vocabulary choices throughout",
        ],
        'style': [
            "Style improvements:\n\n• Vary your sentence beginnings for better rhythm\n• Consider using more active voice\n• Your descriptive language is engaging\n• Look for opportunities to eliminate redundant words\n• Add more specific, concrete details",
            "To enhance your writing style:\n\n• Strong character voice comes through clearly\n• Consider adding more sensory details\n• Good use of dialogue to advance the plot\n• Pacing works well for this genre\n• Consider the emotional arc of your scenes",
        ]
    }
    
    return random.choice(responses.get(assistance_type, ["I'm here to help! Please provide more specific guidance for your request."]))


@csrf_exempt
@login_required
def auto_save(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        document_id = data.get('document_id')
        chapter_id = data.get('chapter_id')
        content = data.get('content')
        title = data.get('title', 'Untitled')
        
        if chapter_id:
            # Auto-save chapter
            chapter = get_object_or_404(Chapter, id=chapter_id)
            # Check permissions
            if not (chapter.project.author == request.user or 
                   chapter.project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'status': 'error', 'message': 'Permission denied'})
            
            chapter.content = content
            chapter.title = title
            chapter.last_edited_by = request.user
            chapter.save()
            
            return JsonResponse({
                'status': 'success',
                'chapter_id': chapter.id,
                'word_count': chapter.word_count
            })
        elif document_id:
            # Auto-save document
            document = get_object_or_404(Document, id=document_id, author=request.user)
            document.content = content
            document.title = title
            document.save()
            return JsonResponse({
                'status': 'success',
                'document_id': document.id,
                'word_count': document.word_count
            })
        else:
            # Create new document
            document = Document.objects.create(
                title=title,
                content=content,
                author=request.user
            )
            return JsonResponse({
                'status': 'success',
                'document_id': document.id,
                'word_count': document.word_count
            })
    
    return JsonResponse({'status': 'error'})


@login_required
def upload_file(request):
    """Simple file upload for chapter editor."""
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        try:
            # Get file path for processing
            import tempfile
            import os
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            # Extract text content
            from .document_parser import extract_text_from_file
            content = extract_text_from_file(temp_file_path)
            
            # Clean up temporary file
            os.unlink(temp_file_path)
            
            # Return just the text content for the editor
            return JsonResponse({
                'success': True,
                'content': content
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    return JsonResponse({
        'success': False,
        'error': 'No file uploaded'
    })


@login_required
@require_http_methods(["POST"])
def reorder_chapters(request, project_id):
    """Handle drag-and-drop reordering of chapters"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    
    try:
        data = json.loads(request.body)
        chapter_orders = data.get('chapter_orders', [])
        
        # Update chapter orders
        for item in chapter_orders:
            chapter_id = item.get('id')
            new_order = item.get('order')
            Chapter.objects.filter(id=chapter_id, project=project).update(order=new_order)
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


@login_required
@require_http_methods(["POST"])
def update_chapter_order(request, project_id, chapter_id):
    """Manually update chapter order by number input"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
    
    try:
        new_order = int(request.POST.get('order', 0))
        old_order = chapter.order
        
        if new_order != old_order:
            # Adjust other chapters' orders
            if new_order > old_order:
                # Moving down: decrease order of chapters between old and new position
                Chapter.objects.filter(
                    project=project,
                    order__gt=old_order,
                    order__lte=new_order
                ).update(order=F('order') - 1)
            else:
                # Moving up: increase order of chapters between new and old position
                Chapter.objects.filter(
                    project=project,
                    order__gte=new_order,
                    order__lt=old_order
                ).update(order=F('order') + 1)
            
            # Set the new order for the moved chapter
            chapter.order = new_order
            chapter.save()
        
        return JsonResponse({'status': 'success'})
    except (ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Invalid order number'})


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_chapter_list(request, project_id):
    """Get paginated list of chapters for the sidebar"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapters = Chapter.objects.filter(project=project).order_by('order')
    
    chapter_data = []
    for chapter in chapters:
        chapter_data.append({
            'id': chapter.id,
            'title': chapter.title,
            'order': chapter.order,
            'word_count': chapter.word_count,
            'content': chapter.content[:200] if chapter.content else '',  # Add content preview for mobile
            'created_at': chapter.created_at.isoformat(),  # Use ISO format for proper date parsing
            'updated_at': chapter.updated_at.isoformat()  # Use ISO format for proper date parsing
        })
    
    return Response({
        'chapters': chapter_data,
        'total_chapters': len(chapter_data),
        'project_word_count': project.total_word_count
    })


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def create_new_chapter(request, project_id):
    """Create a new chapter"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    
    try:
        # Handle both POST form data and JSON data
        if request.content_type == 'application/json':
            import json
            data = json.loads(request.body)
            title = data.get('title', 'Untitled Chapter')
        else:
            title = request.POST.get('title', 'Untitled Chapter')
        
        # Get the next order number safely
        with transaction.atomic():
            max_order = Chapter.objects.filter(project=project).aggregate(
                max_order=models.Max('order')
            )['max_order']
            
            # Fix: Handle when max_order is 0 (which is falsy in Python)
            next_order = 0 if max_order is None else max_order + 1
            
            chapter = Chapter.objects.create(
                title=title,
                content='',
                project=project,
                order=next_order,
                last_edited_by=request.user
            )
        
        return Response({
            'success': True,
            'chapter': {
                'id': chapter.id,
                'title': chapter.title,
                'order': chapter.order,
                'word_count': chapter.word_count,
                'content': chapter.content,
                'created_at': chapter.created_at.isoformat(),
                'updated_at': chapter.updated_at.isoformat()
            }
        })
    except Exception as e:
        return Response({'success': False, 'error': str(e)})


@login_required
@require_http_methods(["POST"])
def delete_chapter(request, project_id, chapter_id):
    """Delete a chapter and reorder remaining chapters"""
    project = get_object_or_404(Project, id=project_id, author=request.user)
    chapter = get_object_or_404(Chapter, id=chapter_id, project=project)
    
    try:
        deleted_order = chapter.order
        chapter.delete()
        
        # Reorder remaining chapters
        Chapter.objects.filter(
            project=project,
            order__gt=deleted_order
        ).update(order=F('order') - 1)
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


# Chapter Management API Endpoints

@login_required
@csrf_exempt 
def create_chapter_api(request):
    """Create a new chapter via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST
            
        project_id = data.get('project')
        title = data.get('title', '').strip()
        content = data.get('content', '')
        
        if not project_id or not title:
            return JsonResponse({'success': False, 'error': 'Project ID and title are required'})
        
        # Get the project and verify access
        try:
            project = Project.objects.get(id=project_id)
            if not (project.author == request.user or project.collaborators.filter(id=request.user.id).exists()):
                return JsonResponse({'success': False, 'error': 'Permission denied'})
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        
        # Create the chapter
        chapter = Chapter.objects.create(
            project=project,
            title=title,
            content=content,
            order=Chapter.objects.filter(project=project).count() + 1
        )
        
        # Update project word count
        project.update_word_count()
        
        return JsonResponse({
            'success': True, 
            'chapter_id': chapter.id,
            'message': f'Chapter "{title}" created successfully!'
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@csrf_exempt
def delete_chapter_api(request, chapter_id):
    """Delete a chapter via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        chapter = get_object_or_404(Chapter, id=chapter_id)
        # Check if user has access to this chapter's project
        if not (chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists()):
            return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
        chapter_title = chapter.title
        deleted_order = chapter.order
        project = chapter.project
        
        chapter.delete()
        
        # Reorder remaining chapters
        Chapter.objects.filter(
            project=project,
            order__gt=deleted_order
        ).update(order=F('order') - 1)
        
        return JsonResponse({
            'success': True,
            'message': f'Chapter "{chapter_title}" has been deleted.'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Error deleting chapter: {str(e)}'
        }, status=400)


@login_required
@csrf_exempt
def reorder_chapters_api(request):
    """Reorder chapters via API"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        chapters_data = data.get('chapters', [])
        
        # Update chapter orders
        for item in chapters_data:
            chapter_id = item.get('id')
            new_order = item.get('order')
            if chapter_id and new_order is not None:
                try:
                    chapter = Chapter.objects.get(id=chapter_id)
                    # Check if user has access to this chapter's project
                    if chapter.project.author == request.user or chapter.project.collaborators.filter(id=request.user.id).exists():
                        chapter.order = new_order
                        chapter.save()
                except Chapter.DoesNotExist:
                    pass
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import Document, User
import json

@csrf_exempt
@login_required
def share_document_api(request, document_id):
    """Share a document with other users"""
    if request.method == 'POST':
        try:
            doc = Document.objects.get(id=document_id, author=request.user)
            data = json.loads(request.body)
            user_ids = data.get('user_ids', [])
            users = User.objects.filter(id__in=user_ids)
            doc.shared_with.set(users)
            doc.save()
            return JsonResponse({'success': True, 'shared_with': [u.username for u in users]})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
    return JsonResponse({'error': 'POST required'}, status=405)

@login_required
def get_document_collaborators_api(request, document_id):
    """Get list of collaborators for a document"""
    try:
        doc = Document.objects.get(id=document_id)
        collaborators = doc.shared_with.all()
        return JsonResponse({'collaborators': [u.username for u in collaborators]})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def users_api(request):
    users = User.objects.all()
    data = [
        {'id': user.id, 'username': user.username}
        for user in users
    ]
    return JsonResponse(data, safe=False)


# New Modern Editor Views
@login_required
def modern_editor(request):
    """Modern Google Docs-like editor"""
    context = {
        'user': request.user,
    }
    
    # Get current chapter or document if specified
    chapter_id = request.GET.get('chapter')
    project_id = request.GET.get('project')
    
    if chapter_id:
        try:
            chapter = Chapter.objects.get(id=chapter_id, project__author=request.user)
            context['chapter'] = chapter
            context['project'] = chapter.project
        except Chapter.DoesNotExist:
            messages.error(request, 'Chapter not found.')
    elif project_id:
        try:
            project = Project.objects.get(id=project_id, author=request.user)
            context['project'] = project
            # Get the latest chapter or create a new one
            latest_chapter = project.chapters.order_by('-created_at').first()
            if latest_chapter:
                context['chapter'] = latest_chapter
        except Project.DoesNotExist:
            messages.error(request, 'Project not found.')
    
    return render(request, 'writer/modern_editor.html', context)


@login_required  
def integrated_editor(request):
    """Integrated editor with all features"""
    context = {
        'user': request.user,
    }
    
    # Get current chapter or document if specified
    chapter_id = request.GET.get('chapter')
    project_id = request.GET.get('project')
    
    if chapter_id:
        try:
            chapter = Chapter.objects.get(id=chapter_id, project__author=request.user)
            context['chapter'] = chapter
            context['project'] = chapter.project
        except Chapter.DoesNotExist:
            messages.error(request, 'Chapter not found.')
    elif project_id:
        try:
            project = Project.objects.get(id=project_id, author=request.user)
            context['project'] = project
            # Get the latest chapter or create a new one
            latest_chapter = project.chapters.order_by('-created_at').first()
            if latest_chapter:
                context['chapter'] = latest_chapter
        except Project.DoesNotExist:
            messages.error(request, 'Project not found.')
    
    return render(request, 'writer/integrated_editor.html', context)


@login_required
def creativity_workshop(request):
    """AI-powered creativity workshop for writers"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creativity_workshop.html', context)


@login_required
def clean_dashboard(request):
    """Clean, writer-focused dashboard"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')[:6]
        
        # Get recent chapters
        recent_chapters = Chapter.objects.filter(
            project__author=request.user
        ).order_by('-updated_at')[:5]
        
        # Calculate stats
        total_words = sum(
            chapter.word_count or 0 
            for project in user_projects 
            for chapter in project.chapters.all()
        )
        
        projects_count = user_projects.count()
        chapters_count = Chapter.objects.filter(project__author=request.user).count()
        
        # Get personal library counts (if models exist)
        try:
            library = PersonalLibrary.objects.get(user=request.user)
            research_notes_count = library.imported_documents.count()
            characters_count = sum(project.characters.count() for project in user_projects)
        except:
            research_notes_count = 0
            characters_count = 0
        
        context.update({
            'recent_projects': user_projects,
            'recent_chapters': recent_chapters,
            'total_words': total_words,
            'projects_count': projects_count,
            'chapters_count': chapters_count,
            'research_notes_count': research_notes_count,
            'characters_count': characters_count,
            'worldbuilding_notes_count': 0,  # Placeholder
            'resources_count': research_notes_count,
        })
        
    except Exception as e:
        logging.error(f"Error loading clean dashboard for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your dashboard.'
    
    return render(request, 'writer/clean_dashboard.html', context)


# API Endpoints for Modern Editor
@csrf_exempt
@login_required
def api_save_document(request):
    """API endpoint to save document content from modern editor"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            title = data.get('title', '')
            content = data.get('content', '')
            chapters_data = data.get('chapters', [])
            
            # For now, save to the user's first project or create a new one
            project = Project.objects.filter(author=request.user).first()
            if not project:
                project = Project.objects.create(
                    title="My Writing Project",
                    description="Created from editor",
                    author=request.user
                )
            
            # Save or update the chapter
            if title and content:
                chapter, created = Chapter.objects.get_or_create(
                    project=project,
                    title=title,
                    defaults={'content': content, 'author': request.user}
                )
                if not created:
                    chapter.content = content
                    chapter.save()
                
                return JsonResponse({
                    'success': True,
                    'chapter_id': chapter.id,
                    'message': 'Document saved successfully'
                })
            
            return JsonResponse({'success': False, 'error': 'Title and content required'})
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})


@csrf_exempt
@login_required  
def api_ai_assistance(request):
    """API endpoint for AI writing assistance"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            action = data.get('action', '')
            text = data.get('text', '')
            
            # Simple mock responses for now (you can integrate real AI later)
            mock_responses = {
                'continue': f"Continuing from your text: '{text[:50]}...', here's a possible continuation: The story unfolds with unexpected twists as the protagonist discovers hidden truths about their world.",
                'improve': f"Here's an improved version: {text.replace('good', 'excellent').replace('bad', 'terrible').replace('nice', 'wonderful')}" if text else "Please select some text to improve.",
                'grammar': f"Grammar-corrected version: {text.replace(' i ', ' I ').replace("dont", "don't").replace("cant", "can't")}" if text else "No grammar issues found.",
                'character': "Character development suggestion: Consider adding internal conflict to make your character more relatable and three-dimensional.",
                'plot': "Plot suggestion: What if the protagonist's greatest strength becomes their weakness at the climax?"
            }
            
            suggestion = mock_responses.get(action, "AI assistance is not available for this action.")
            
            return JsonResponse({
                'success': True,
                'suggestion': suggestion,
                'action': action
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})


@login_required
def compact_library(request):
    """Compact, readable version of the personal library"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Get latest project for continue writing
        latest_project = user_projects.first()
        
        # Get imported documents if available
        try:
            library = PersonalLibrary.objects.get(user=request.user)
            imported_documents = library.imported_documents.all().order_by('-created_at')[:10]
        except PersonalLibrary.DoesNotExist:
            imported_documents = []
        
        context.update({
            'projects': user_projects,
            'latest_project': latest_project,
            'imported_documents': imported_documents,
        })
        
    except Exception as e:
        logging.error(f"Error loading compact library for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your library.'
    
    return render(request, 'writer/compact_library.html', context)


# Ultimate Templates - The Most Beautiful Writing Platform


@login_required 
def ultimate_dashboard(request):
    """The most beautiful, organized, and mobile-friendly dashboard"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects (with visibility control)
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Calculate statistics
        total_words = sum(project.total_word_count for project in user_projects)
        active_projects = user_projects.filter(is_public=True, author=request.user).count()
        
        context.update({
            'user_projects': user_projects,
            'total_words': total_words,
            'active_projects': active_projects,
        })
        
    except Exception as e:
        logging.error(f"Error loading ultimate dashboard for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your dashboard.'
    
    return render(request, 'writer/ultimate_dashboard.html', context)


@login_required
def ultimate_library(request):
    """The most beautifully organized library with perfect text visibility"""
    context = {
        'user': request.user,
    }
    
    try:
        # Get user's projects with detailed information
        user_projects = Project.objects.filter(
            Q(author=request.user) | Q(collaborators=request.user)
        ).distinct().order_by('-updated_at')
        
        # Add computed fields for library display
        projects_data = []
        for project in user_projects:
            project_data = {
                'id': project.id,
                'title': project.title,
                'description': project.description or 'No description available.',
                'genre': project.genre or 'Unspecified',
                'word_count': project.total_word_count,
                'target_words': project.target_word_count,
                'chapter_count': project.chapter_count,
                'progress': project.progress_percentage,
                'updated_at': project.updated_at,
                'created_at': project.created_at,
                'show_on_dashboard': getattr(project, 'show_on_dashboard', True),
            }
            projects_data.append(project_data)
        
        context.update({
            'projects': projects_data,
            'total_projects': len(projects_data),
            'total_words': sum(p['word_count'] for p in projects_data),
        })
        
    except Exception as e:
        logging.error(f"Error loading ultimate library for user {request.user}: {e}")
        context['error_message'] = 'There was an issue loading your library.'
    
    return render(request, 'writer/ultimate_library.html', context)


@login_required
def ultimate_workshop(request):
    """The ultimate AI-powered creativity workshop for writers"""
    context = {
        'user': request.user,
    }
    
    # In a real implementation, this would load writing exercises from the database
    # For now, we'll use the JavaScript-based exercises in the template
    
    return render(request, 'writer/ultimate_workshop.html', context)


@login_required
def ai_playground(request):
    """AI Creative Playground - Where writers play with AI tools"""
    context = {
        'user': request.user,
    }
    
    return render(request, 'writer/ai_playground.html', context)


@login_required
def creativity_workshop(request):
    """Creativity workshop for brainstorming and connecting ideas"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creativity_workshop.html', context)


@login_required
def creative_notebook(request):
    """Creative notebook for mind mapping and connecting ideas"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/creative_notebook.html', context)


@login_required
def bookshelf_dashboard(request):
    """Bookshelf dashboard with visual book display"""
    user_projects = Project.objects.filter(author=request.user).order_by('-updated_at')
    user_documents = Document.objects.filter(author=request.user).order_by('-updated_at')
    
    # Calculate writing statistics
    total_words = sum(project.total_word_count for project in user_projects)
    total_projects = user_projects.count()
    total_documents = user_documents.count()
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
        'user_documents': user_documents,
        'total_words': total_words,
        'total_projects': total_projects,
        'total_documents': total_documents,
        'show_bookshelf': True,
    }
    return render(request, 'writer/dashboard_with_bookshelf.html', context)


# Legacy Creative Notebook models remain for data migration if needed


# Google Docs-like Editor View
@login_required
def google_docs_editor(request, document_id=None):
    """Google Docs-like editor for documents"""
    document = None
    
    if document_id:
        document = get_object_or_404(Document, id=document_id, author=request.user)
    
    context = {
        'user': request.user,
        'document': document,
    }
    return render(request, 'writer/google_docs_editor.html', context)


@login_required
def academic_editor(request):
    """Academic writing editor"""
    context = {
        'user': request.user,
    }
    return render(request, 'writer/academic_editor.html', context)


@login_required
def book_reader(request, project_id):
    """Book reader view - displays project content as a formatted ebook"""
    project = get_object_or_404(Project, pk=project_id)
    
    # Check if user can access this project
    if project.author != request.user and not project.collaborators.filter(id=request.user.id).exists():
        return render(request, '403.html')
    
    chapters = project.chapters.all().order_by('order')
    
    context = {
        'user': request.user,
        'project': project,
        'chapters': chapters,
    }
    return render(request, 'writer/book_reader.html', context)


@login_required
def user_preferences(request):
    """User preferences page"""
    profile = UserProfile.get_or_create_for_user(request.user)
    context = {
        'user': request.user,
        'profile': profile,
    }
    return render(request, 'writer/user_preferences.html', context)


@csrf_exempt
@login_required
def api_update_theme(request):
    """Update user theme preference"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            theme = data.get('theme', 'ethereal')
            
            # Get or create user profile
            profile = UserProfile.get_or_create_for_user(request.user)
            profile.theme = theme
            profile.save()
            
            return JsonResponse({'success': True, 'theme': theme})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required
def my_library(request):
    """My Library - Personal collection of saved books"""
    user_projects = Project.objects.filter(author=request.user).order_by('-created_at')
    user_documents = Document.objects.filter(author=request.user).order_by('-created_at')
    
    # Separate projects by bookshelf visibility
    bookshelf_projects = user_projects.filter(show_on_dashboard=True)
    library_projects = user_projects.filter(show_on_dashboard=False)
    
    # Calculate statistics
    total_books = user_projects.count() + user_documents.count()
    total_words = sum(p.total_word_count or 0 for p in user_projects) + sum(d.word_count or 0 for d in user_documents)
    
    context = {
        'user': request.user,
        'user_projects': user_projects,
        'bookshelf_projects': bookshelf_projects,
        'library_projects': library_projects,
        'user_documents': user_documents,
        'total_books': total_books,
        'total_words': total_words,
    }
    
    return render(request, 'writer/my_library.html', context)


@csrf_exempt
@login_required
def toggle_bookshelf_visibility(request):
    """Toggle project visibility on dashboard bookshelf"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_id = data.get('project_id')
            
            project = Project.objects.get(id=project_id, author=request.user)
            project.show_on_dashboard = not project.show_on_dashboard
            project.save()
            
            return JsonResponse({
                'success': True,
                'show_on_dashboard': project.show_on_dashboard,
                'message': f"Project {'added to' if project.show_on_dashboard else 'removed from'} bookshelf"
            })
            
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Project not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST method required'})

@csrf_exempt
def api_register(request):
    """API endpoint for user registration"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')
            
            if not username or not email or not password:
                return JsonResponse({
                    'success': False, 
                    'error': 'Username, email, and password are required'
                }, status=400)
            
            # Check if username already exists
            if User.objects.filter(username=username).exists():
                return JsonResponse({
                    'success': False, 
                    'error': 'Username already exists'
                }, status=400)
                
            # Check if email already exists
            if User.objects.filter(email=email).exists():
                return JsonResponse({
                    'success': False, 
                    'error': 'Email already exists'
                }, status=400)
            
            # Create user
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password
            )
            
            # Create auth token
            from rest_framework.authtoken.models import Token
            token, created = Token.objects.get_or_create(user=user)
            
            return JsonResponse({
                'success': True,
                'message': 'Account created successfully',
                'token': token.key,
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': user.email
                }
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'error': str(e)
            }, status=500)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)
